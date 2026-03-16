"""
Generate PhishAware Frontend Design Documentation (.docx)
Run: python generate_frontend_doc.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Helper colours ───────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1E, 0x3A, 0x8A)
MID_BLUE   = RGBColor(0x3B, 0x82, 0xF6)
LIGHT_BLUE = RGBColor(0xDB, 0xEA, 0xFE)
GRAY_700   = RGBColor(0x37, 0x41, 0x51)
GRAY_100   = RGBColor(0xF3, 0xF4, 0xF6)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: add heading ───────────────────────────────────────────────────────
def add_heading(text, level=1, color=DARK_BLUE):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
    return h

# ── Helper: add body paragraph ────────────────────────────────────────────────
def add_para(text='', bold=False, color=None, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold  = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p

# ── Helper: figure placeholder ────────────────────────────────────────────────
def add_figure(fig_id: str, caption: str):
    """Add a labelled placeholder box + Caption-styled caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[{fig_id}  –  insert screenshot here]')
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = MID_BLUE

    # Use built-in Caption style so Word's Table of Figures picks it up
    cap = doc.add_paragraph(style='Caption')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f'{fig_id}: {caption}')
    r.font.size   = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GRAY_700
    doc.add_paragraph()          # small vertical space

# ── Helper: Table of Figures field (Word updates it on right-click) ───────────
def add_tof():
    """Insert a TOC field that collects all Caption-style paragraphs."""
    p = doc.add_paragraph()

    def _run_with(elem):
        r = OxmlElement('w:r')
        r.append(elem)
        p._p.append(r)

    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    _run_with(begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\h \\z \\t "Caption,1" '
    _run_with(instr)

    sep = OxmlElement('w:fldChar')
    sep.set(qn('w:fldCharType'), 'separate')
    _run_with(sep)

    # Placeholder text shown before Word updates the field
    placeholder = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '9CA3AF')
    rPr.append(color)
    italic = OxmlElement('w:i')
    rPr.append(italic)
    placeholder.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = '[Right-click here → Update Field to generate Table of Figures]'
    placeholder.append(t)
    p._p.append(placeholder)

    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    _run_with(end)

# ── Helper: build a table ─────────────────────────────────────────────────────
def add_table(headers, rows, caption=''):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.font.bold   = True
        r.font.size   = Pt(10)
        r.font.italic = True
        r.font.color.rgb = GRAY_700

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade_cell(cell, '1E3A8A')
        run = cell.paragraphs[0].add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(10)
        run.font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        bg = 'F3F4F6' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            shade_cell(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    return t

# ═══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('PhishAware Platform')
title_run.font.size  = Pt(28)
title_run.font.bold  = True
title_run.font.color.rgb = DARK_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('Frontend Design Documentation')
sub_run.font.size  = Pt(20)
sub_run.font.color.rgb = MID_BLUE

doc.add_paragraph()
doc.add_paragraph()

for label, value in [
    ('Project:',  'PhishAware – Cybersecurity Awareness Platform for Saudi Arabia'),
    ('Author:',   '[Your Name]'),
    ('Date:',     'March 2026'),
    ('Version:',  '1.0'),
    ('Document:', 'Frontend Design Documentation'),
]:
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bold_run = meta.add_run(f'{label} ')
    bold_run.font.bold = True
    bold_run.font.size = Pt(12)
    val_run = meta.add_run(value)
    val_run.font.size  = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual, static)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('Table of Contents', level=1)
toc_entries = [
    ('1.', 'Introduction'),
    ('   1.1', 'Purpose of This Document'),
    ('   1.2', 'Platform Overview'),
    ('   1.3', 'Development Timeline'),
    ('   1.4', 'Technology Stack'),
    ('2.', 'Design Overview & Principles'),
    ('   2.1', 'Design Philosophy'),
    ('   2.2', 'Color System'),
    ('   2.3', 'Typography'),
    ('   2.4', 'Logo & Branding'),
    ('   2.5', 'Layout Principles'),
    ('3.', 'Week 9: Foundation & Authentication'),
    ('   3.1', 'React.js Environment Setup'),
    ('   3.2', 'Bilingual Configuration (i18next)'),
    ('   3.3', 'Axios API Service Layer'),
    ('   3.4', 'Login Page Design'),
    ('   3.5', 'Company Registration Page'),
    ('   3.6', 'Protected Route System'),
    ('4.', 'Week 10: Admin Portal & Campaign Management'),
    ('   4.1', 'Admin Main Layout'),
    ('   4.2', 'Campaign Management Interface'),
    ('   4.3', 'Template Library Interface'),
    ('   4.4', 'Analytics Dashboard'),
    ('   4.5', 'Employee Management Interface'),
    ('   4.6', 'AI Email Generation Feature'),
    ('5.', 'Week 11: Employee Portal, Public Portal & Gamification'),
    ('   5.1', 'Employee Dashboard'),
    ('   5.2', 'Quiz Interface'),
    ('   5.3', 'Training Module Interface'),
    ('   5.4', 'Gamification Features'),
    ('   5.5', 'Public Portal Homepage'),
    ('   5.6', 'Awareness Content Pages'),
    ('   5.7', 'Public Quiz Interface'),
    ('6.', 'Responsive Design & Accessibility'),
    ('7.', 'Conclusion'),
    ('8.', 'Appendix: Screenshot Index'),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0 if num.strip().endswith('.') and len(num.strip()) <= 2 else 0.3)
    r1 = p.add_run(f'{num}  ')
    r1.font.bold = num.strip().endswith('.') and len(num.strip()) <= 2
    r1.font.size = Pt(11)
    r2 = p.add_run(title)
    r2.font.size = Pt(11)

doc.add_paragraph()
add_heading('Table of Figures', level=1)
add_tof()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 – INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('1. Introduction', level=1)

add_heading('1.1 Purpose of This Document', level=2)
add_para(
    'This document provides comprehensive visual documentation of the PhishAware platform '
    'frontend design and implementation. It captures all user interfaces developed over a '
    'three-week frontend development sprint, organised chronologically to demonstrate the '
    'systematic build-up of the platform\'s capabilities.'
)

add_heading('1.2 Platform Overview', level=2)
add_para(
    'PhishAware is a bilingual (Arabic/English) cybersecurity awareness platform designed '
    'specifically for Saudi Arabian and MENA region organisations. The platform provides '
    'phishing simulation campaigns, training modules, and awareness content with cultural '
    'relevance. Users operate across three distinct roles: Public (no login required), '
    'Employee, and Administrator.'
)

add_heading('1.3 Development Timeline', level=2)
add_para('Frontend development was completed in three structured weeks:')
for week, desc in [
    ('Week 9:',  'Foundation setup, authentication, and core infrastructure'),
    ('Week 10:', 'Admin portal, campaign management, and analytics'),
    ('Week 11:', 'Employee portal, public portal, and gamification features'),
]:
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(week + '  ')
    r1.font.bold  = True
    r1.font.color.rgb = DARK_BLUE
    p.add_run(desc)

add_heading('1.4 Technology Stack', level=2)
add_table(
    headers=['Technology', 'Version', 'Purpose'],
    rows=[
        ['React.js',       '18.2+',   'UI component framework'],
        ['React Router',   '6.x',     'Client-side routing'],
        ['Tailwind CSS',   '3.x',     'Utility-first styling'],
        ['i18next',        '23.x',    'Internationalisation (Arabic / English)'],
        ['Axios',          '1.x',     'HTTP client for API communication'],
        ['Recharts',       '2.x',     'Data visualisation charts'],
        ['Lucide React',   'Latest',  'Icon library'],
    ],
    caption='Table 1.1: Frontend Technology Stack',
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 – DESIGN OVERVIEW & PRINCIPLES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('2. Design Overview & Principles', level=1)

add_heading('2.1 Design Philosophy', level=2)
add_para('PhishAware\'s interface design follows four core principles:')
add_table(
    headers=['Principle', 'Description', 'Implementation Examples'],
    rows=[
        ['Clarity',              'Clear visual hierarchy and intuitive navigation',           'Consistent spacing, typography, colour-coded status indicators'],
        ['Accessibility',        'WCAG 2.1 AA compliance for inclusive design',               'High contrast ratios, keyboard navigation, screen reader support'],
        ['Bilingual Excellence', 'Seamless Arabic/English experience',                        'RTL/LTR layouts, culturally appropriate translations, language toggle'],
        ['Cultural Relevance',   'Saudi/MENA-specific context',                               'Examples using Nafath, Absher, Saudi banks, local government services'],
    ],
    caption='Table 2.1: Core Design Principles',
)

add_heading('2.2 Color System', level=2)
add_figure('Figure 2.1', 'PhishAware Color Palette – screenshot of colour swatches with hex codes and labels')
add_para(
    'Caption: PhishAware platform colour system. Primary blue (#3B82F6) for branding and '
    'actions, green (#10B981) for success states, yellow (#F59E0B) for warnings, red '
    '(#EF4444) for errors, and neutral greys for text and borders.',
    color=GRAY_700
)
add_table(
    headers=['Color', 'Hex Code', 'Primary Usage', 'UI Examples'],
    rows=[
        ['Primary Blue',   '#3B82F6', 'Buttons, links, active states',      'Login button, sidebar active item, primary CTAs'],
        ['Dark Blue',      '#1E3A8A', 'Headers, navigation',                  'Top bar, section headings'],
        ['Success Green',  '#10B981', 'Completed, passed, success',           'Training completed badge, quiz passed'],
        ['Warning Yellow', '#F59E0B', 'Pending, in-progress',                 'Assignment pending, campaign active'],
        ['Danger Red',     '#EF4444', 'Failed, errors, deletions',            'Quiz failed, error messages, delete buttons'],
        ['Gray 600',       '#6B7280', 'Secondary text, labels',               'Descriptions, metadata'],
        ['Gray 200',       '#E5E7EB', 'Borders, dividers',                    'Card borders, table lines'],
    ],
    caption='Table 2.2: Color Usage Guide',
)

add_heading('2.3 Typography', level=2)
add_figure('Figure 2.2', 'Typography Hierarchy – sample page showing all heading levels in English and Arabic')
add_para(
    'Caption: Typography hierarchy demonstration. (Left) English using Inter font family. '
    '(Right) Arabic using Tajawal font family. Both maintain consistent size ratios and '
    'visual weight for optimal readability.',
    color=GRAY_700
)
add_table(
    headers=['Element', 'English Font', 'Arabic Font', 'Size', 'Weight', 'Usage'],
    rows=[
        ['H1 (Page Title)',  'Inter', 'Tajawal', '32px', 'Bold 700',      'Dashboard titles, page headings'],
        ['H2 (Section)',     'Inter', 'Tajawal', '24px', 'Semibold 600',  'Card titles, section headers'],
        ['H3 (Subsection)',  'Inter', 'Tajawal', '20px', 'Semibold 600',  'Card subtitles'],
        ['Body Text',        'Inter', 'Tajawal', '16px', 'Regular 400',   'Main content, descriptions'],
        ['Small Text',       'Inter', 'Tajawal', '14px', 'Regular 400',   'Metadata, timestamps'],
        ['Labels',           'Inter', 'Tajawal', '12px', 'Medium 500',    'Form labels, tags'],
    ],
    caption='Table 2.3: Typography Scale',
)

add_heading('2.4 Logo & Branding', level=2)
add_figure('Figure 2.3', 'PhishAware Logo System – icon only, horizontal, vertical, and light/dark variants')
add_para(
    'Caption: PhishAware logo system: (a) Icon only (shield with fishing hook), '
    '(b) Horizontal layout with wordmark, (c) Vertical layout for square spaces, '
    '(d) White version for dark backgrounds. The shield represents security while '
    'the hook symbolises phishing threats.',
    color=GRAY_700
)

add_heading('2.5 Layout Principles', level=2)
add_figure('Figure 2.4', 'Grid System & Spacing – wireframe overlay showing grid structure and spacing units')
add_para(
    'Caption: PhishAware employs a consistent 8 px base spacing unit. Layout uses a '
    '12-column grid for desktop (1280 px max-width), collapsing to single column on '
    'mobile. Consistent padding (24 px) and margins maintain visual rhythm.',
    color=GRAY_700
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 – WEEK 9
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('3. Week 9: Foundation & Authentication', level=1)
add_para(
    'Week 9 focused on establishing the technical foundation and core authentication system, '
    'enabling secure user access across three user roles (Public, Employee, Admin).'
)
add_table(
    headers=['Task', 'Component', 'Status', 'Description'],
    rows=[
        ['React.js Setup',        'Project scaffold',  'Complete', 'Initialised with Vite, configured routing'],
        ['i18next Integration',   'Language system',   'Complete', 'Arabic/English with RTL support'],
        ['Axios Configuration',   'API layer',         'Complete', 'Base URL, interceptors, error handling'],
        ['Authentication Pages',  'Login/Register UI', 'Complete', 'Bilingual auth flow'],
        ['Route Protection',      'Guard system',      'Complete', 'Role-based access control'],
    ],
    caption='Table 3.1: Week 9 Deliverables',
)

add_heading('3.1 React.js Environment Setup', level=2)
add_para(
    'The project was initialised using Vite (faster build tooling than Create React App). '
    'React Router v6 was configured for client-side routing, with a feature-module '
    'folder structure for maintainability.'
)
add_figure('Figure 3.1', 'Project Folder Structure – VS Code file explorer showing feature-module organisation')
add_para(
    'Caption: PhishAware frontend folder structure. Organised by feature modules '
    '(pages, components, api, utils) for maintainability. Separate folders for public, '
    'employee, and admin portals.',
    color=GRAY_700
)

add_heading('3.2 Bilingual Configuration (i18next)', level=2)
add_figure('Figure 3.2', 'Language Switcher Component – AR/EN toggle button in navigation bar')
add_para(
    'Caption: Language switcher in the navigation bar. Users can toggle between '
    'Arabic (العربية) and English at any time. Current language is highlighted in blue.',
    color=GRAY_700
)
add_figure('Figure 3.3', 'RTL Layout Transformation – side-by-side English (LTR) and Arabic (RTL)')
add_para(
    'Caption: Automatic RTL transformation when Arabic is selected. (Left) English '
    'interface with left-aligned navigation. (Right) Arabic interface with mirrored '
    'layout, right-aligned text, and reversed element order.',
    color=GRAY_700
)
add_table(
    headers=['Feature', 'Implementation', 'Example'],
    rows=[
        ['Language Toggle',   'Dropdown in header',          'AR / EN selector'],
        ['RTL Layout',        'CSS direction attribute',     'dir="rtl" for Arabic'],
        ['Translation Files', 'JSON key-value pairs',        'en.json, ar.json'],
        ['Dynamic Loading',   'Language-specific assets',    'Date formats, number formats'],
        ['Persistent Choice', 'LocalStorage',                'Remembers user preference'],
    ],
    caption='Table 3.2: Language Support Features',
)
add_figure('Figure 3.4', 'Translation File Structure – en.json and ar.json key-value files in code editor')
add_para(
    'Caption: Translation file structure. All interface text is externalised into JSON '
    'files with key-value pairs. Nested objects organise translations by feature '
    '(auth, dashboard, training, etc.).',
    color=GRAY_700
)

add_heading('3.3 Axios API Service Layer', level=2)
add_figure('Figure 3.5', 'API Configuration Code – Axios instance creation with request/response interceptors')
add_para(
    'Caption: Axios API client configuration. Base URL points to Django backend, '
    'request interceptor adds JWT token, response interceptor handles authentication '
    'errors and token refresh.',
    color=GRAY_700
)
add_table(
    headers=['Service', 'File', 'Endpoints', 'Purpose'],
    rows=[
        ['Authentication', 'api/auth.js',      'login, register, verify',  'User authentication'],
        ['Training',       'api/training.js',  'getModules, assign, complete', 'Training management'],
        ['Campaigns',      'api/campaigns.js', 'create, send, track',      'Phishing simulations'],
        ['Analytics',      'api/analytics.js', 'getStats, getReports',     'Data visualisation'],
        ['Employees',      'api/employees.js', 'list, invite, remove',     'Employee management'],
    ],
    caption='Table 3.3: API Endpoints Organisation',
)

add_heading('3.4 Login Page Design', level=2)
add_figure('Figure 3.6', 'Login Page – English version with email, password, and sign-in button')
add_para(
    'Caption: PhishAware login page (English). Features include email and password '
    'fields with validation, "Remember Me" checkbox, "Forgot Password" link, and '
    'prominent "Sign In" button. PhishAware logo displayed at top with tagline '
    '"Cybersecurity Awareness Platform".',
    color=GRAY_700
)
add_figure('Figure 3.7', 'Login Page – Arabic version demonstrating complete RTL layout')
add_para(
    'Caption: Login page in Arabic demonstrating complete RTL transformation. Text is '
    'right-aligned, form fields are mirrored, and the layout flows from right to left '
    'naturally for Arabic readers.',
    color=GRAY_700
)
add_table(
    headers=['Field', 'Validation Rules', 'Error Messages'],
    rows=[
        ['Email',    'Required, valid email format',     '"Please enter a valid email address"'],
        ['Password', 'Required, min 8 characters',      '"Password must be at least 8 characters"'],
        ['Form',     'Both fields valid on submission',  '"Please check your credentials" (on auth error)'],
    ],
    caption='Table 3.4: Login Form Validation Rules',
)
add_figure('Figure 3.8', 'Login Error States – validation errors with red borders and inline messages')
add_para(
    'Caption: Login form validation errors. Red border indicates invalid fields, error '
    'messages appear below inputs in red text. Email shows "Invalid email format" and '
    'password shows "Password required".',
    color=GRAY_700
)

add_heading('3.5 Company Registration Page', level=2)
add_figure('Figure 3.9', 'Company Registration – Section 1: Company Details fields')
add_para(
    'Caption: Company registration page – Section 1: Company Details. Fields include '
    'company name, industry dropdown, company size, and contact phone. All required '
    'fields are marked with asterisks (*).',
    color=GRAY_700
)
add_figure('Figure 3.10', 'Company Registration – Section 2: Administrator Account fields')
add_para(
    'Caption: Company registration page – Section 2: Administrator Account. Creates the '
    'first admin user with first name, last name, email, and password. Password strength '
    'indicator updates as the user types.',
    color=GRAY_700
)
add_table(
    headers=['Section', 'Field', 'Type', 'Validation', 'Example'],
    rows=[
        ['Company Info',  'Company Name', 'Text',     'Required, 3–100 chars', '"Saudi Tech Solutions"'],
        ['Company Info',  'Industry',     'Dropdown', 'Required',              'Finance, Healthcare, Education'],
        ['Company Info',  'Company Size', 'Dropdown', 'Required',              '"51–200 employees"'],
        ['Company Info',  'Phone',        'Tel',      'Required, valid format', '+966 XX XXX XXXX'],
        ['Admin Details', 'First Name',   'Text',     'Required, 2–50 chars',  '"Ahmed"'],
        ['Admin Details', 'Last Name',    'Text',     'Required, 2–50 chars',  '"Al-Harbi"'],
        ['Admin Details', 'Email',        'Email',    'Required, unique',      'admin@company.sa'],
        ['Admin Details', 'Password',     'Password', 'Min 8 chars, complex',  '••••••••'],
    ],
    caption='Table 3.5: Registration Form Fields',
)
add_figure('Figure 3.11', 'Password Strength Indicator – Weak (red) / Medium (yellow) / Strong (green)')
add_para(
    'Caption: Password strength indicator. (Top) Weak password shown in red. '
    '(Middle) Medium strength in yellow. (Bottom) Strong password in green. '
    'Indicator updates in real time as the user types.',
    color=GRAY_700
)
add_figure('Figure 3.12', 'Email Verification Flow – success page after clicking verification link')
add_para(
    'Caption: Email verification confirmation page. Displayed after user clicks '
    'verification link from email. Shows success message with checkmark icon and '
    '"Login" button to proceed to the platform.',
    color=GRAY_700
)

add_heading('3.6 Protected Route System', level=2)
add_figure('Figure 3.13', 'Route Guard Flow Diagram – flowchart of authentication and role checks')
add_para(
    'Caption: Route protection flow. When a user attempts to access a protected route: '
    '(1) Check if authenticated, (2) Verify user role matches route requirement, '
    '(3) Allow access or redirect to login.',
    color=GRAY_700
)
add_table(
    headers=['Route Path', 'Required Role', 'Redirect If Unauthorised'],
    rows=[
        ['/login',              'None (Public)',  'Dashboard (if already logged in)'],
        ['/register/company',   'None (Public)',  'Dashboard (if already logged in)'],
        ['/training',           'None (Public)',  'N/A'],
        ['/employee/*',         'EMPLOYEE',       '/login'],
        ['/admin/*',            'ADMIN',          '/login'],
        ['/super-admin/*',      'SUPER_ADMIN',    '/login'],
    ],
    caption='Table 3.6: Route Protection Matrix',
)
add_figure('Figure 3.14', 'Unauthorised Access Redirect – unauthenticated user redirected to login')
add_para(
    'Caption: Protected route behaviour. When an unauthenticated user attempts to '
    'access /admin/dashboard, they are automatically redirected to /login with the '
    'return URL preserved for post-login redirect.',
    color=GRAY_700
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 – WEEK 10
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('4. Week 10: Admin Portal & Campaign Management', level=1)
add_para(
    'Week 10 focused on building the complete admin portal, including campaign creation, '
    'template management, employee oversight, and analytics dashboards.'
)
add_table(
    headers=['Task', 'Component', 'Status', 'Description'],
    rows=[
        ['Admin Layout',        'Sidebar + Header',  'Complete', 'Persistent navigation with role indicator'],
        ['Campaign Management', 'CRUD interface',    'Complete', 'Create, send, track simulation campaigns'],
        ['Template Library',    'Upload system',     'Complete', 'Manage phishing email templates'],
        ['Analytics Dashboard', 'Charts & stats',    'Complete', 'Click rates, department stats, risk visualisation'],
        ['Employee Management', 'List & assign',     'Complete', 'View employees, assign campaigns, track risk'],
        ['AI Email Generation', 'Generate button',   'Complete', 'AI-powered template creation'],
    ],
    caption='Table 4.1: Week 10 Deliverables',
)

add_heading('4.1 Admin Main Layout', level=2)
add_figure('Figure 4.1', 'Admin Dashboard Layout – full desktop view showing sidebar, header, and content area')
add_para(
    'Caption: Admin portal main layout. (Left) Collapsible sidebar with navigation menu. '
    '(Top) Header with search, notifications, and user profile. (Centre) Main content area '
    'displaying dashboard widgets. Fixed sidebar with scrollable content area.',
    color=GRAY_700
)
add_figure('Figure 4.2', 'Admin Sidebar Navigation – expanded menu with icons and section labels')
add_para(
    'Caption: Admin sidebar navigation menu. Hierarchical structure with icons and labels. '
    'Active route highlighted in blue. Includes Dashboard, Campaigns, Simulations, '
    'Employees, Training, Analytics, and Profile sections.',
    color=GRAY_700
)
add_table(
    headers=['Section', 'Icon', 'Sub-items', 'Route'],
    rows=[
        ['Dashboard',  '📊', '—',                                    '/admin/dashboard'],
        ['Campaigns',  '📧', '—',                                    '/admin/campaigns'],
        ['Simulations','🎯', '—',                                    '/admin/simulations'],
        ['Employees',  '👥', 'List, Invite, Risk Scores',            '/admin/employees/*'],
        ['Training',   '📚', 'Modules, Assignments, Progress',       '/admin/training/*'],
        ['Analytics',  '📈', 'Overview, Reports',                    '/admin/analytics/*'],
        ['Profile',    '👤', 'Settings, Security',                   '/admin/profile'],
    ],
    caption='Table 4.2: Admin Navigation Structure',
)
add_figure('Figure 4.3', 'Admin Header Components – search bar, notification bell, language toggle, profile dropdown')
add_para(
    'Caption: Admin header components. (Left to Right) Search bar, notification bell with '
    'badge, language toggle (AR/EN), user profile dropdown showing "Admin" role badge and '
    'logout option.',
    color=GRAY_700
)

add_heading('4.2 Campaign Management Interface', level=2)
add_figure('Figure 4.4', 'Campaigns List View – table of all campaigns with status, dates, and metrics')
add_para(
    'Caption: Campaigns management page showing all simulation campaigns. Columns include '
    'campaign name, status badge, recipients count, clicked count, reported count, start '
    'date, and action buttons. Colour-coded status badges.',
    color=GRAY_700
)
add_table(
    headers=['Status', 'Color', 'Badge', 'Description'],
    rows=[
        ['Draft',     'Gray',   '⚪ Draft',     'Campaign created but not sent'],
        ['Scheduled', 'Blue',   '📅 Scheduled', 'Set to send at future date'],
        ['Active',    'Yellow', '⚡ Active',    'Currently running'],
        ['Completed', 'Green',  '✅ Completed', 'Campaign finished'],
        ['Paused',    'Orange', '⏸ Paused',    'Temporarily stopped'],
    ],
    caption='Table 4.3: Campaign Status Indicators',
)
add_figure('Figure 4.5', 'Create Campaign Modal – form with name, description, template, and recipients')
add_para(
    'Caption: Create Campaign modal. Fields include campaign name, description, template '
    'selection dropdown, recipient selection (individuals or departments), and send options '
    '(immediate or scheduled).',
    color=GRAY_700
)
add_figure('Figure 4.6', 'Campaign Creation – Template Selection grid with difficulty badges')
add_para(
    'Caption: Template selection interface. Grid of email template cards with preview '
    'thumbnails, names, and difficulty levels (Easy/Medium/Hard). Selected template '
    'highlighted with blue border.',
    color=GRAY_700
)
add_figure('Figure 4.7', 'Campaign Creation – Recipient Selection with checkboxes and department filter')
add_para(
    'Caption: Recipient selection. Employee list with checkboxes, department filter, and '
    '"Select All" option. Selected count shown at top ("15 employees selected").',
    color=GRAY_700
)
add_figure('Figure 4.8', 'Campaign Send Confirmation – summary dialog before sending')
add_para(
    'Caption: Send campaign confirmation modal. Shows summary (name, template, 15 recipients, '
    'scheduled time) with warning message and "Cancel" / "Send Campaign" buttons.',
    color=GRAY_700
)

add_heading('4.3 Template Library Interface', level=2)
add_figure('Figure 4.9', 'Template Library Grid – all templates with thumbnails, categories, and difficulty badges')
add_para(
    'Caption: Template library in grid layout. Each card shows thumbnail, name, difficulty '
    'badge, category tag, and Edit/Delete buttons. "+ Upload Template" button at top-right.',
    color=GRAY_700
)
add_table(
    headers=['Category', 'Icon', 'Example Templates', 'Count'],
    rows=[
        ['Finance',    '💰', 'Fake invoice, payment verification',       '8 templates'],
        ['HR',         '👥', 'Benefits update, policy change',           '6 templates'],
        ['IT Support', '💻', 'Password reset, system upgrade',           '7 templates'],
        ['Delivery',   '📦', 'Package notification, shipment update',    '5 templates'],
        ['Government', '🏛', 'Nafath verification, Absher update',       '4 templates'],
    ],
    caption='Table 4.4: Template Categories',
)
add_figure('Figure 4.10', 'Upload Template Modal – fields for name, category, difficulty, and HTML body')
add_para(
    'Caption: Upload new template modal with name, category dropdown, difficulty selector, '
    'subject line, HTML body editor, and file attachment option.',
    color=GRAY_700
)
add_figure('Figure 4.11', 'Template Edit Interface – WYSIWYG editor with formatting toolbar')
add_para(
    'Caption: Template editor with WYSIWYG editor. Toolbar includes text formatting, link '
    'insertion, image upload, and HTML source view. Preview pane shows final rendering.',
    color=GRAY_700
)
add_figure('Figure 4.12', 'Template Preview – full email preview as seen by recipient')
add_para(
    'Caption: Template preview modal. Shows complete email as it will appear to the recipient '
    'with sender, subject, body, and attachments. "Use This Template" button to select.',
    color=GRAY_700
)

add_heading('4.4 Analytics Dashboard', level=2)
add_figure('Figure 4.13', 'Admin Analytics Dashboard Overview – KPIs, charts, and risk distribution')
add_para(
    'Caption: Admin analytics dashboard showing KPI cards followed by charts for campaign '
    'performance over time, department comparison, and employee risk distribution.',
    color=GRAY_700
)
add_figure('Figure 4.14', 'KPI Summary Cards – total campaigns, employees, click rate, high-risk count')
add_para(
    'Caption: KPI summary cards with trend arrows. Cards cover Total Campaigns, Total '
    'Employees, Average Click Rate, and High Risk Employees. Colour-coded green/red arrows '
    'indicate improvement or worsening.',
    color=GRAY_700
)
add_table(
    headers=['Metric', 'Visualisation', 'Time Range', 'Benchmark'],
    rows=[
        ['Click Rates',          'Line chart',    'Last 6 months',  'Industry avg: 30%'],
        ['Department Stats',     'Bar chart',     'Current',        'By department'],
        ['Risk Scores',          'Pie chart',     'Current',        'Low/Medium/High distribution'],
        ['Employee Performance', 'Table',         'Current',        'Individual scores'],
        ['Campaign Success',     'Trend line',    'Last 12 months', 'Target: <20% click rate'],
    ],
    caption='Table 4.5: Analytics Metrics Displayed',
)
add_figure('Figure 4.15', 'Campaign Click Rates Chart – line graph showing declining trend over 6 months')
add_para(
    'Caption: Click rates trend over six months. Declining line from 35 % in October to '
    '23 % in March indicates improving employee awareness.',
    color=GRAY_700
)
add_figure('Figure 4.16', 'Department Comparison Chart – horizontal bar chart colour-graded by risk')
add_para(
    'Caption: Department performance comparison. Horizontal bar chart showing click rates '
    'by department with colour gradient (green=low risk, red=high risk).',
    color=GRAY_700
)
add_figure('Figure 4.17', 'Employee Risk Distribution – pie chart with Low/Medium/High segments')
add_para(
    'Caption: Risk level pie chart. Green (62 % – 28 employees), Yellow (24 % – 11 employees), '
    'Red (14 % – 6 employees). Legend shows counts and percentages.',
    color=GRAY_700
)
add_figure('Figure 4.18', 'Export Analytics Modal – CSV/PDF export with date range and field selection')
add_para(
    'Caption: Export analytics data dialog. Options for CSV or PDF, date range selector, '
    'and checkboxes for data selection (Campaigns, Employee Stats, Department Stats).',
    color=GRAY_700
)

add_heading('4.5 Employee Management Interface', level=2)
add_figure('Figure 4.19', 'Employee List View – sortable table with risk badges, campaigns, and actions')
add_para(
    'Caption: Employee management page. Table with name, email, department, risk score badge, '
    'campaigns assigned, last activity, and action buttons. Search bar and department filter '
    'at top.',
    color=GRAY_700
)
add_table(
    headers=['Column', 'Data Type', 'Sorting', 'Filter'],
    rows=[
        ['Name',        'Text',         'A–Z',           'Search'],
        ['Email',       'Text',         'A–Z',           'Search'],
        ['Department',  'Badge',        'A–Z',           'Dropdown'],
        ['Risk Score',  'Colour badge', 'Low→High',      'Dropdown'],
        ['Campaigns',   'Number',       '0–99',          '—'],
        ['Last Active', 'Date',         'Recent→Old',    'Date range'],
        ['Actions',     'Buttons',      '—',             '—'],
    ],
    caption='Table 4.6: Employee Table Columns',
)
add_figure('Figure 4.20', 'Invite Employee Modal – email, name, and department fields with send button')
add_para(
    'Caption: Invite employee modal. Email, first name, last name, and department dropdown. '
    '"Send Invitation" sends email with registration link.',
    color=GRAY_700
)
add_figure('Figure 4.21', 'Employee Detail View – profile with training history, campaigns, and quiz tabs')
add_para(
    'Caption: Employee detail page with personal info and tabbed sections: Training History, '
    'Campaign Results, Quiz Performance, and Achievements.',
    color=GRAY_700
)
add_figure('Figure 4.22', 'Assign Campaign to Employee – select campaign and employees with note field')
add_para(
    'Caption: Assign campaign modal. Dropdown to select campaign, employee checklist, optional '
    'note, and "Assign" button.',
    color=GRAY_700
)
add_figure('Figure 4.23', 'Employee Risk Score Breakdown – contributing factors and recommendations')
add_para(
    'Caption: Risk score breakdown. Contributing factors listed with point values. Progress '
    'bar visualisation. Recommendations section suggests next steps to reduce score.',
    color=GRAY_700
)

add_heading('4.6 AI Email Generation Feature', level=2)
add_figure('Figure 4.24', 'AI Email Generation Button – "Generate with AI" in campaign creation')
add_para(
    'Caption: AI-powered email generation feature. "Generate with AI" button beside template '
    'selection triggers realistic phishing email creation based on selected parameters.',
    color=GRAY_700
)
add_figure('Figure 4.25', 'AI Generation Parameters Modal – category, difficulty, tone, and keywords')
add_para(
    'Caption: AI generation parameters. Select category, difficulty, tone, and optional '
    'keywords. "Generate Email" button triggers creation with loading spinner.',
    color=GRAY_700
)
add_figure('Figure 4.26', 'AI-Generated Email Preview – quality score, red flags, and action buttons')
add_para(
    'Caption: AI-generated email preview with realism score and detected red flags. Options '
    'to "Use This Email", "Regenerate", or "Edit & Save".',
    color=GRAY_700
)
add_table(
    headers=['Parameter', 'Options', 'Effect on Output'],
    rows=[
        ['Category',  'Finance, HR, IT, Delivery, Government', 'Determines email topic and sender type'],
        ['Difficulty','Easy, Medium, Hard',                     'Controls obvious vs subtle red flags'],
        ['Tone',      'Urgent, Friendly, Official',             'Adjusts language formality and pressure tactics'],
        ['Keywords',  'Optional custom terms',                  'Includes specific phrases (e.g., "Nafath", "salary")'],
    ],
    caption='Table 4.7: AI Generation Parameters',
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 – WEEK 11
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('5. Week 11: Employee Portal, Public Portal & Gamification', level=1)
add_para(
    'Week 11 completed the platform with employee-facing features, a public awareness portal, '
    'and gamification elements to encourage sustained engagement.'
)
add_table(
    headers=['Task', 'Component', 'Status', 'Description'],
    rows=[
        ['Employee Dashboard', 'Personal stats UI',  'Complete', 'Risk score, assigned training, badges'],
        ['Quiz Interface',     'Email flagging game', 'Complete', 'Safe/phishing identification with feedback'],
        ['Training Module UI', 'Video + quiz system', 'Complete', 'Watch training, complete assessment'],
        ['Gamification',       'Badges & leaderboard','Complete', 'Achievement showcase, ranking table'],
        ['Public Portal',      'No-login homepage',   'Complete', 'Awareness content accessible to all'],
        ['Awareness Pages',    'Educational content', 'Complete', 'Phishing/Vishing/Smishing guides (AR/EN)'],
        ['Public Quiz',        'Knowledge test',      'Complete', 'Anonymous quiz for public users'],
        ['Final Polish',       'UX refinement',       'Complete', 'Consistent styling, responsive, RTL verified'],
    ],
    caption='Table 5.1: Week 11 Deliverables',
)

add_heading('5.1 Employee Dashboard', level=2)
add_figure('Figure 5.1', 'Employee Dashboard Overview – risk gauge, training progress, and badges')
add_para(
    'Caption: Employee dashboard homepage. Top: personal risk score metre, training progress '
    '(3/5 completed), and upcoming deadlines. Bottom: assigned trainings, notifications, '
    'earned badges, and quick-access buttons.',
    color=GRAY_700
)
add_figure('Figure 5.2', 'Employee Risk Score Widget – circular gauge with colour gradient (Low/Medium/High)')
add_para(
    'Caption: Circular risk gauge showing score out of 100. Colour gradient: 0–33 green '
    '(Low), 34–66 yellow (Medium), 67–100 red (High). Breakdown link explains contributing '
    'factors.',
    color=GRAY_700
)
add_figure('Figure 5.3', 'Assigned Training Cards – due dates, completion status, and start buttons')
add_para(
    'Caption: Assigned training cards with module name, category badge, due date (red if '
    'overdue), completion status, and "Start Training" / "Continue" button. Completed '
    'trainings show green checkmark and score.',
    color=GRAY_700
)
add_table(
    headers=['Widget', 'Content', 'Action'],
    rows=[
        ['Risk Score',          'Current score /100, trend arrow',         'View breakdown'],
        ['Training Progress',   'X/Y completed, progress bar',             'View all trainings'],
        ['Upcoming Deadlines',  'Next 3 due trainings and dates',          'Start training'],
        ['Recent Badges',       'Last 3 earned badges',                    'View all achievements'],
        ['Notifications',       'Unread count, recent 5',                  'View all notifications'],
        ['Leaderboard Rank',    'Current position, +/- change',            'View full leaderboard'],
    ],
    caption='Table 5.2: Employee Dashboard Widgets',
)
add_figure('Figure 5.4', 'Employee Dashboard – Mobile View (375 px) with stacked widgets')
add_para(
    'Caption: Employee dashboard on mobile. Widgets stack vertically, sidebar collapses to '
    'hamburger, touch-friendly buttons (min 44 px height). All functionality preserved.',
    color=GRAY_700
)

add_heading('5.2 Quiz Interface', level=2)
add_figure('Figure 5.5', 'Quiz Start Screen – title, description, difficulty badge, and Start button')
add_para(
    'Caption: Quiz start screen with title "Phishing Identification Quiz", description, '
    'difficulty badge (Medium), estimated time (5–10 minutes), and "Start Quiz" button.',
    color=GRAY_700
)
add_figure('Figure 5.6', 'Quiz Question – email mockup with Safe and Phishing answer buttons')
add_para(
    'Caption: Quiz question interface. Full email mockup with sender, subject, and body. '
    'Two large buttons: "Safe" (green) and "Phishing" (red). Progress indicator at top.',
    color=GRAY_700
)
add_figure('Figure 5.7', 'Quiz Instant Feedback – Correct Answer (green banner with explanation)')
add_para(
    'Caption: Correct answer feedback. Green banner with checkmark icon, explanation of '
    'red flags detected. "Next Question" button to continue.',
    color=GRAY_700
)
add_figure('Figure 5.8', 'Quiz Instant Feedback – Incorrect Answer (red banner with learning tip)')
add_para(
    'Caption: Incorrect answer feedback. Red banner explaining missed red flags: suspicious '
    'domain, pressure tactics, sensitive information request. Learning opportunity provided.',
    color=GRAY_700
)
add_figure('Figure 5.9', 'Quiz Results Screen – score, breakdown by category, and recommendations')
add_para(
    'Caption: Quiz results page. Large score "8/10 – 80% Correct", breakdown of correct '
    'identifications, recommendations, and buttons: "Retake Quiz", "View Training", '
    '"Back to Dashboard".',
    color=GRAY_700
)
add_table(
    headers=['Element', 'Correct Answer', 'Incorrect Answer'],
    rows=[
        ['Banner Color', 'Green (#10B981)',          'Red (#EF4444)'],
        ['Icon',         '✅ Checkmark',              '❌ Cross'],
        ['Message',      'Correct! This is phishing/safe', 'Incorrect. This was phishing/safe'],
        ['Explanation',  'Brief confirmation of red flags', 'Detailed explanation of missed red flags'],
        ['Learning Tip', 'Optional additional info',  'Always provided'],
        ['Next Action',  '"Next Question" button',    '"Next Question" button'],
    ],
    caption='Table 5.3: Quiz Feedback Components',
)

add_heading('5.3 Training Module Interface', level=2)
add_figure('Figure 5.10', 'Training Module Page – video player, learning objectives, and start button')
add_para(
    'Caption: Training module page. 16:9 video player at top with play controls. Below: '
    'module description, learning objectives, estimated duration, and "Start Training" button.',
    color=GRAY_700
)
add_figure('Figure 5.11', 'Training Mode Selection – Video Training vs Interactive Learning cards')
add_para(
    'Caption: Training delivery method selection. "Video Training" card (watch video then '
    'quiz) and "Interactive Learning" card (hands-on scenarios). User chooses preferred '
    'method.',
    color=GRAY_700
)
add_figure('Figure 5.12', 'Video Training Player – standard controls, closed captions, progress tracking')
add_para(
    'Caption: Video player with play/pause, scrubber, volume, fullscreen, and playback '
    'speed. Closed captions available in Arabic and English.',
    color=GRAY_700
)
add_figure('Figure 5.13', 'Interactive Learning – Saudi Examples with Nafath scene and click-to-reveal')
add_para(
    'Caption: Interactive learning (dark theme). Scene 2: comparing fake vs real Nafath '
    'message. Click-to-reveal red flags feature. Navigation pills show "Scene 2/7".',
    color=GRAY_700
)
add_figure('Figure 5.14', 'Post-Training Quiz – multiple choice questions unlocked after video')
add_para(
    'Caption: Post-training assessment quiz. Appears after video completion. Passing score: '
    '80 %. If failed, options to "Review Training" or "Retake Quiz".',
    color=GRAY_700
)
add_figure('Figure 5.15', 'Training Completion Certificate – downloadable PDF with score and date')
add_para(
    'Caption: Completion certificate design. PhishAware logo, employee name, training title, '
    'completion date, final score. "Download Certificate" and "Share Achievement" buttons.',
    color=GRAY_700
)
add_table(
    headers=['Component', 'Function', 'Completion Criteria'],
    rows=[
        ['Video Player',        'Primary content delivery',       'Watch 90 %+ of video'],
        ['Interactive Lesson',  'Alternative delivery method',    'Complete all scenes + quiz'],
        ['Comprehension Quiz',  'Knowledge assessment',           'Score ≥ 80 % (4/5 correct)'],
        ['Certificate',         'Completion proof',               'Auto-generated on pass'],
        ['Progress Tracker',    'Resume capability',              'Saves position every 30 seconds'],
    ],
    caption='Table 5.4: Training Module Components',
)

add_heading('5.4 Gamification Features', level=2)
add_figure('Figure 5.16', 'Badges Showcase – earned and locked badge grid with unlock requirements')
add_para(
    'Caption: Employee achievements page. Earned badges in full colour; locked badges in '
    'greyscale with unlock requirements shown on hover.',
    color=GRAY_700
)
add_figure('Figure 5.17', 'Leaderboard View – ranked table with medal icons and department column')
add_para(
    'Caption: Company leaderboard ranking employees by points. Medal icons for top 3. '
    '"This Week / Month / All Time" tabs. Current user highlighted in blue.',
    color=GRAY_700
)
add_figure('Figure 5.18', 'Badge Unlock Animation – modal with shine effect and confetti')
add_para(
    'Caption: Badge unlocked notification. Shine animation on badge icon, name, description, '
    'points earned (+50), and "Awesome!" button. Confetti in background.',
    color=GRAY_700
)
add_table(
    headers=['Badge', 'Icon', 'Requirement', 'Points'],
    rows=[
        ['First Steps',          '🎯', 'Complete first training',                 '50'],
        ['Quick Learner',        '⚡', 'Complete training in < 10 min',           '25'],
        ['Perfect Score',        '💯', 'Score 100 % on quiz',                     '50'],
        ['Phish Detector',       '🎣', 'Identify 10 phishing emails',             '100'],
        ['Reporter',             '📢', 'Report first real phishing attempt',      '75'],
        ['Streak Master',        '🔥', 'Complete training 5 weeks in a row',      '150'],
        ['Department Champion',  '🏆', 'Highest score in department',             '200'],
    ],
    caption='Table 5.5: Available Badges',
)
add_figure('Figure 5.19', 'Points Breakdown – table of point values for all actions')
add_para(
    'Caption: Points system explanation. Actions and their values: Complete Training (+50), '
    'Pass Quiz (+25), Perfect Score (+50), Report Phishing (+10), Identify Phishing (+5 each).',
    color=GRAY_700
)
add_figure('Figure 5.20', 'Leaderboard Top 3 Podium – gold/silver/bronze with confetti for first place')
add_para(
    'Caption: Podium display for top 3 employees with avatars, names, and point totals. '
    'Confetti animation for #1. "View Full Leaderboard" button below.',
    color=GRAY_700
)

add_heading('5.5 Public Portal Homepage', level=2)
add_figure('Figure 5.21', 'Public Homepage – Hero Section with headline, CTA, and language toggle')
add_para(
    'Caption: PhishAware public hero section. Heading "Protect Yourself from Cyber Threats", '
    'subheading highlighting Saudi Arabia context, "Start Free Training" CTA button, and '
    'AR/EN language toggle.',
    color=GRAY_700
)
add_figure('Figure 5.22', 'Public Homepage – Features Section with three benefit columns')
add_para(
    'Caption: Features section: (1) Free Training, (2) Saudi Context (Nafath/Absher examples), '
    '(3) Bilingual support. Each column has icon, heading, and description.',
    color=GRAY_700
)
add_figure('Figure 5.23', 'Public Homepage – Training Topic Cards (Phishing, Vishing, Smishing)')
add_para(
    'Caption: Training topics section. Three large cards with icon, bilingual title, '
    'description, lesson count, and "Learn More" button.',
    color=GRAY_700
)
add_figure('Figure 5.24', 'Public Homepage – Statistics Banner with social-proof counters')
add_para(
    'Caption: Statistics section: 10,000+ Users Trained, 50+ Saudi Companies, 95% '
    'Satisfaction Rate, 2 Languages. Counter animation on scroll.',
    color=GRAY_700
)
add_table(
    headers=['Section', 'Content', 'Purpose'],
    rows=[
        ['Hero',            'Headline, CTA, background image',    'Capture attention, drive action'],
        ['Features',        '3 key benefits with icons',          'Communicate value proposition'],
        ['Training Topics', 'Phishing, Vishing, Smishing cards',  'Navigate to specific content'],
        ['Statistics',      'User count, company count, satisfaction', 'Build trust and credibility'],
        ['How It Works',    '3-step process (Choose→Learn→Test)', 'Explain user journey'],
        ['Testimonials',    'User quotes and avatars',            'Social proof'],
        ['Footer',          'Links, contact, social media',       'Navigation and information'],
    ],
    caption='Table 5.6: Public Homepage Sections',
)

add_heading('5.6 Awareness Content Pages', level=2)
add_figure('Figure 5.25', 'Phishing Awareness Page – English version with Saudi examples and embedded quiz')
add_para(
    'Caption: Phishing awareness page. Sections: What is Phishing?, Common Examples, How to '
    'Identify, What to Do, Test Your Knowledge. Real Saudi examples (Nafath phishing). '
    '"Take Quiz" button at bottom.',
    color=GRAY_700
)
add_figure('Figure 5.26', 'Phishing Awareness Page – Arabic (RTL) with mirrored layout')
add_para(
    'Caption: Arabic version with complete RTL layout. Same structure and visual hierarchy '
    'as English version but culturally adapted terminology.',
    color=GRAY_700
)
add_figure('Figure 5.27', 'Interactive Example – Fake vs Real Nafath SMS with click-to-reveal red flags')
add_para(
    'Caption: Side-by-side comparison of fake and real Nafath SMS. Click-to-reveal red '
    'flags on fake message (suspicious link, urgency, grammar errors).',
    color=GRAY_700
)
add_figure('Figure 5.28', 'Vishing Awareness Page – audio player with scam call transcript')
add_para(
    'Caption: Vishing page with audio example of scam call, transcript, Saudi-specific '
    'scenarios (fake bank calls, lottery scams), and prevention tips.',
    color=GRAY_700
)
add_figure('Figure 5.29', 'Smishing Awareness Page – SMS screenshots with annotated red flags')
add_para(
    'Caption: Smishing page showing fraudulent text message examples targeting Saudi users. '
    'Screenshots with annotations and step-by-step response guide.',
    color=GRAY_700
)
add_table(
    headers=['Page', 'Sections', 'Interactive Elements', 'Saudi Examples'],
    rows=[
        ['Phishing', 'Definition, Examples, Identification, Prevention, Quiz',
         'Click-to-reveal, embedded quiz', 'Nafath SMS, Absher email, bank notifications'],
        ['Vishing',  'Definition, Common Scams, Audio Examples, Prevention',
         'Audio player, Spot Red Flags game', 'Fake ministry calls, lottery scams, family emergency'],
        ['Smishing', 'Definition, Examples, Identification, Prevention',
         'Screenshot comparison, embedded quiz', 'Package delivery SMS, payment alerts, account verification'],
    ],
    caption='Table 5.7: Awareness Content Structure',
)

add_heading('5.7 Public Quiz Interface', level=2)
add_figure('Figure 5.30', 'Public Quiz Landing – no login required, topic and difficulty selection')
add_para(
    'Caption: Public knowledge quiz landing page. No login required. Three difficulty options '
    '(Beginner/Intermediate/Advanced) and "Start Quiz" button.',
    color=GRAY_700
)
add_figure('Figure 5.31', 'Public Quiz – Multiple Choice Question with radio buttons and progress bar')
add_para(
    'Caption: Multiple choice quiz question with four options, radio buttons, "Submit Answer" '
    'button, and "Question 4/10" progress bar.',
    color=GRAY_700
)
add_figure('Figure 5.32', 'Public Quiz Results – score, category breakdown, and registration CTA')
add_para(
    'Caption: Public quiz results "7/10 – 70% Correct". Breakdown by category '
    '(Phishing/Vishing/Smishing). Recommendations and CTA: "Sign up for free training".',
    color=GRAY_700
)
add_table(
    headers=['Feature', 'Public Users', 'Registered Employees'],
    rows=[
        ['Authentication',  'Not required',        'JWT token required'],
        ['Score Saving',    'Session only',         'Persistent in database'],
        ['Progress Tracking','No',                  'Full history'],
        ['Recommendations', 'Generic',              'Personalised to risk profile'],
        ['Gamification',    'No badges/leaderboard','Full badges, leaderboard, points'],
        ['Training Access', 'Public modules only',  'All assigned + public modules'],
        ['Certificate',     'No',                   'Yes, downloadable PDF'],
    ],
    caption='Table 5.8: Public vs Registered Employee Quiz Experience',
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 – RESPONSIVE DESIGN & ACCESSIBILITY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('6. Responsive Design & Accessibility', level=1)

add_heading('6.1 Breakpoint System', level=2)
add_table(
    headers=['Breakpoint', 'Screen Width', 'Target Devices', 'Layout Changes'],
    rows=[
        ['Mobile (sm)',   '< 640 px',   'Smartphones',             'Single column, hamburger menu, stacked cards'],
        ['Tablet (md)',   '640–1024 px','Tablets, landscape phones','2-column grid, collapsible sidebar'],
        ['Desktop (lg)',  '1024–1280 px','Laptops',                 '3-column grid, full sidebar visible'],
        ['Wide (xl)',     '> 1280 px',  'Large monitors',          '4-column grid, expanded dashboard'],
    ],
    caption='Table 6.1: Responsive Breakpoints',
)
add_figure('Figure 6.1', 'Responsive Grid – same page at mobile, tablet, and desktop widths')
add_para(
    'Caption: Responsive layout demonstration across three screen sizes. Dashboard widgets '
    'adapt from single column (mobile) to four columns (desktop) using Tailwind CSS grid.',
    color=GRAY_700
)

add_heading('6.2 Accessibility Features', level=2)
add_table(
    headers=['Feature', 'Implementation', 'Standard Met'],
    rows=[
        ['Colour Contrast',  'Min 4.5:1 ratio for normal text',       'WCAG 2.1 AA'],
        ['Keyboard Nav',     'All interactive elements tab-accessible','WCAG 2.1 A'],
        ['ARIA Labels',      'All icons and buttons labelled',         'WCAG 2.1 A'],
        ['Focus Indicators', 'Visible focus ring on all elements',     'WCAG 2.1 AA'],
        ['Alt Text',         'All images have descriptive alt text',   'WCAG 2.1 A'],
        ['Form Labels',      'Every input paired with visible label',  'WCAG 2.1 A'],
        ['Error Messages',   'Not colour-only; icon + text used',      'WCAG 2.1 AA'],
    ],
    caption='Table 6.2: Accessibility Features',
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 – CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('7. Conclusion', level=1)
add_para(
    'The PhishAware frontend was designed and implemented over three focused development '
    'weeks, progressing from foundational authentication in Week 9, through a comprehensive '
    'admin portal in Week 10, to employee-facing features, public awareness content, and '
    'gamification in Week 11.'
)
add_para(
    'The resulting platform achieves its primary goals: a bilingual (Arabic/English) '
    'experience with authentic RTL support, culturally relevant Saudi examples, and a '
    'progressive disclosure design that guides users from public awareness content through '
    'to full organisational security training.'
)
add_para(
    'Key design achievements include: a coherent colour and typography system, a robust '
    'route protection model, interactive phishing simulation campaigns with analytics, '
    'AI-assisted template generation, and a gamification layer (badges, leaderboard, '
    'points) that incentivises continued engagement.'
)
add_para(
    'Future work may include dark mode support, native mobile applications, and integration '
    'with additional Saudi government identity platforms such as Nafath for verified user '
    'authentication.'
)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 – APPENDIX: SCREENSHOT INDEX
# ═══════════════════════════════════════════════════════════════════════════════
add_heading('8. Appendix: Screenshot Index', level=1)
add_para(
    'The table below lists all figures referenced in this document, organised by section. '
    'Replace each placeholder with the actual screenshot when preparing the final submission.'
)

screenshot_rows = [
    # Week 9
    ['3.1',  'Project Folder Structure',                            'Week 9 – Setup',              '3.1'],
    ['3.2',  'Language Switcher Component',                         'Week 9 – i18n',               '3.2'],
    ['3.3',  'RTL Layout Transformation',                           'Week 9 – i18n',               '3.3'],
    ['3.4',  'Translation File Structure',                          'Week 9 – i18n',               '3.4'],
    ['3.5',  'API Configuration Code',                              'Week 9 – Axios',              '3.5'],
    ['3.6',  'Login Page – English',                                'Week 9 – Auth',               '3.6'],
    ['3.7',  'Login Page – Arabic',                                 'Week 9 – Auth',               '3.7'],
    ['3.8',  'Login Error States',                                  'Week 9 – Auth',               '3.8'],
    ['3.9',  'Company Registration – Company Details',              'Week 9 – Auth',               '3.9'],
    ['3.10', 'Company Registration – Admin Details',                'Week 9 – Auth',               '3.10'],
    ['3.11', 'Password Strength Indicator',                         'Week 9 – Auth',               '3.11'],
    ['3.12', 'Email Verification Flow',                             'Week 9 – Auth',               '3.12'],
    ['3.13', 'Route Guard Flow Diagram',                            'Week 9 – Routes',             '3.13'],
    ['3.14', 'Unauthorised Access Redirect',                        'Week 9 – Routes',             '3.14'],
    # Week 10
    ['4.1',  'Admin Dashboard Layout',                              'Week 10 – Admin',             '4.1'],
    ['4.2',  'Admin Sidebar Navigation',                            'Week 10 – Admin',             '4.2'],
    ['4.3',  'Admin Header Components',                             'Week 10 – Admin',             '4.3'],
    ['4.4',  'Campaigns List View',                                 'Week 10 – Campaigns',         '4.4'],
    ['4.5',  'Create Campaign Modal',                               'Week 10 – Campaigns',         '4.5'],
    ['4.6',  'Campaign – Template Selection',                       'Week 10 – Campaigns',         '4.6'],
    ['4.7',  'Campaign – Recipient Selection',                      'Week 10 – Campaigns',         '4.7'],
    ['4.8',  'Campaign Send Confirmation',                          'Week 10 – Campaigns',         '4.8'],
    ['4.9',  'Template Library Grid',                               'Week 10 – Templates',         '4.9'],
    ['4.10', 'Upload Template Modal',                               'Week 10 – Templates',         '4.10'],
    ['4.11', 'Template Edit Interface',                             'Week 10 – Templates',         '4.11'],
    ['4.12', 'Template Preview',                                    'Week 10 – Templates',         '4.12'],
    ['4.13', 'Analytics Dashboard Overview',                        'Week 10 – Analytics',         '4.13'],
    ['4.14', 'KPI Summary Cards',                                   'Week 10 – Analytics',         '4.14'],
    ['4.15', 'Campaign Click Rates Chart',                          'Week 10 – Analytics',         '4.15'],
    ['4.16', 'Department Comparison Chart',                         'Week 10 – Analytics',         '4.16'],
    ['4.17', 'Employee Risk Distribution',                          'Week 10 – Analytics',         '4.17'],
    ['4.18', 'Export Analytics Modal',                              'Week 10 – Analytics',         '4.18'],
    ['4.19', 'Employee List View',                                  'Week 10 – Employees',         '4.19'],
    ['4.20', 'Invite Employee Modal',                               'Week 10 – Employees',         '4.20'],
    ['4.21', 'Employee Detail View',                                'Week 10 – Employees',         '4.21'],
    ['4.22', 'Assign Campaign to Employee',                         'Week 10 – Employees',         '4.22'],
    ['4.23', 'Employee Risk Score Breakdown',                       'Week 10 – Employees',         '4.23'],
    ['4.24', 'AI Email Generation Button',                          'Week 10 – AI',                '4.24'],
    ['4.25', 'AI Generation Parameters Modal',                      'Week 10 – AI',                '4.25'],
    ['4.26', 'AI-Generated Email Preview',                          'Week 10 – AI',                '4.26'],
    # Week 11
    ['5.1',  'Employee Dashboard Overview',                         'Week 11 – Employee Portal',   '5.1'],
    ['5.2',  'Employee Risk Score Widget',                          'Week 11 – Employee Portal',   '5.2'],
    ['5.3',  'Assigned Training Cards',                             'Week 11 – Employee Portal',   '5.3'],
    ['5.4',  'Employee Dashboard – Mobile View',                    'Week 11 – Employee Portal',   '5.4'],
    ['5.5',  'Quiz Start Screen',                                   'Week 11 – Quiz',              '5.5'],
    ['5.6',  'Quiz Question – Email Display',                       'Week 11 – Quiz',              '5.6'],
    ['5.7',  'Quiz Feedback – Correct',                             'Week 11 – Quiz',              '5.7'],
    ['5.8',  'Quiz Feedback – Incorrect',                           'Week 11 – Quiz',              '5.8'],
    ['5.9',  'Quiz Results Screen',                                 'Week 11 – Quiz',              '5.9'],
    ['5.10', 'Training Module Page',                                'Week 11 – Training',          '5.10'],
    ['5.11', 'Training Mode Selection',                             'Week 11 – Training',          '5.11'],
    ['5.12', 'Video Training Player',                               'Week 11 – Training',          '5.12'],
    ['5.13', 'Interactive Learning – Saudi Examples',               'Week 11 – Training',          '5.13'],
    ['5.14', 'Post-Training Quiz',                                  'Week 11 – Training',          '5.14'],
    ['5.15', 'Training Completion Certificate',                     'Week 11 – Training',          '5.15'],
    ['5.16', 'Badges Showcase',                                     'Week 11 – Gamification',      '5.16'],
    ['5.17', 'Leaderboard View',                                    'Week 11 – Gamification',      '5.17'],
    ['5.18', 'Badge Unlock Animation',                              'Week 11 – Gamification',      '5.18'],
    ['5.19', 'Points Breakdown',                                    'Week 11 – Gamification',      '5.19'],
    ['5.20', 'Leaderboard Top 3 Podium',                            'Week 11 – Gamification',      '5.20'],
    ['5.21', 'Public Homepage – Hero Section',                      'Week 11 – Public Portal',     '5.21'],
    ['5.22', 'Public Homepage – Features Section',                  'Week 11 – Public Portal',     '5.22'],
    ['5.23', 'Public Homepage – Training Topics',                   'Week 11 – Public Portal',     '5.23'],
    ['5.24', 'Public Homepage – Statistics',                        'Week 11 – Public Portal',     '5.24'],
    ['5.25', 'Phishing Awareness – English',                        'Week 11 – Awareness',         '5.25'],
    ['5.26', 'Phishing Awareness – Arabic RTL',                     'Week 11 – Awareness',         '5.26'],
    ['5.27', 'Interactive Example – Nafath SMS Comparison',         'Week 11 – Awareness',         '5.27'],
    ['5.28', 'Vishing Awareness Page',                              'Week 11 – Awareness',         '5.28'],
    ['5.29', 'Smishing Awareness Page',                             'Week 11 – Awareness',         '5.29'],
    ['5.30', 'Public Quiz Landing',                                 'Week 11 – Public Quiz',       '5.30'],
    ['5.31', 'Public Quiz – Question',                              'Week 11 – Public Quiz',       '5.31'],
    ['5.32', 'Public Quiz Results',                                 'Week 11 – Public Quiz',       '5.32'],
    # Design System
    ['2.1',  'Color Palette',                                       'Design System',               '2.1'],
    ['2.2',  'Typography Hierarchy',                                'Design System',               '2.2'],
    ['2.3',  'Logo System',                                         'Design System',               '2.3'],
    ['2.4',  'Grid & Spacing',                                      'Design System',               '2.4'],
    ['6.1',  'Responsive Grid Comparison',                          'Responsive Design',           '6.1'],
]

add_table(
    headers=['Figure #', 'Description', 'Section', 'Document Reference'],
    rows=screenshot_rows,
    caption='Table 8.1: Complete Screenshot Index',
)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(__file__), 'Frontend_Design_Documentation.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
