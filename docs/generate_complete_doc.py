"""
Generate PhishAware_Complete_Documentation.docx
Complete unified thesis: Backend (Weeks 1-5 + Testing) + Frontend (Weeks 9-11)
Run: python generate_complete_doc.py
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

HERE       = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DOC = os.path.join(HERE, 'PhishAware_Complete_Documentation.docx')

doc = Document()

# ── Page setup (matches frontend doc) ───────────────────────────────────────
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Colours ──────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1E, 0x3A, 0x8A)
MID_BLUE   = RGBColor(0x3B, 0x82, 0xF6)
LIGHT_BLUE = RGBColor(0xDB, 0xEA, 0xFE)
GRAY_700   = RGBColor(0x37, 0x41, 0x51)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREEN      = RGBColor(0x05, 0x96, 0x69)
YELLOW_DK  = RGBColor(0x92, 0x40, 0x09)

# ── Helpers ───────────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_heading(text, level=1, color=DARK_BLUE):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_para(text='', bold=False, color=None, size=11,
             align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold        = bold
        run.italic      = italic
        run.font.size   = Pt(size)
        if color:
            run.font.color.rgb = color
    return p

def add_bullet(text, bold_prefix='', color=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r1 = p.add_run(bold_prefix + '  ')
        r1.font.bold = True
        r1.font.size = Pt(11)
        if color:
            r1.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def add_figure(fig_id: str, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[{fig_id}  -  insert screenshot here]')
    run.font.size      = Pt(11)
    run.font.bold      = True
    run.font.color.rgb = MID_BLUE
    cap = doc.add_paragraph(style='Caption')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f'{fig_id}: {caption}')
    r.font.size      = Pt(10)
    r.font.italic    = True
    r.font.color.rgb = GRAY_700
    doc.add_paragraph()

def add_table(headers, rows, caption=''):
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.font.bold      = True
        r.font.size      = Pt(10)
        r.font.italic    = True
        r.font.color.rgb = GRAY_700
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style     = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        shade_cell(cell, '1E3A8A')
        run = cell.paragraphs[0].add_run(h)
        run.font.bold      = True
        run.font.size      = Pt(10)
        run.font.color.rgb = WHITE
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        bg = 'F3F4F6' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            shade_cell(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()
    return t

def add_tof_field(instr_text):
    """Insert a Word auto-update field (for ToC/ToF/ToT)."""
    p = doc.add_paragraph()
    def _run_with(elem):
        r = OxmlElement('w:r'); r.append(elem); p._p.append(r)
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin'); _run_with(begin)
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve'); instr.text = instr_text; _run_with(instr)
    sep = OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'), 'separate'); _run_with(sep)
    ph = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    col = OxmlElement('w:color'); col.set(qn('w:val'), '9CA3AF'); rPr.append(col)
    it = OxmlElement('w:i'); rPr.append(it); ph.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve')
    t.text = '[Right-click here and select Update Field to populate this list]'
    ph.append(t); p._p.append(ph)
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end'); _run_with(end)

def add_deliverables_table(rows, caption):
    add_table(
        headers=['Task', 'Type', 'Status', 'Description'],
        rows=rows, caption=caption
    )

# =============================================================================
#  TITLE PAGE
# =============================================================================
for _ in range(3):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run('PhishAware Platform')
tr.font.size = Pt(32); tr.font.bold = True; tr.font.color.rgb = DARK_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run('Complete System Documentation')
sr.font.size = Pt(22); sr.font.color.rgb = MID_BLUE

doc.add_paragraph()
tag_p = doc.add_paragraph()
tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tg = tag_p.add_run('"Click Smarter, Stay Safer"')
tg.font.size = Pt(14); tg.font.italic = True; tg.font.color.rgb = GRAY_700

for _ in range(2):
    doc.add_paragraph()

for label, value in [
    ('Project:',    'PhishAware - Cybersecurity Awareness Platform for Saudi Arabia'),
    ('Supervisor:', 'Dr. Saim Rasheed'),
    ('Date:',       'March 2026'),
    ('Version:',    '1.0'),
    ('Document:',   'Backend & Frontend Complete Documentation'),
]:
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = meta.add_run(f'{label} '); br.font.bold = True; br.font.size = Pt(12)
    vr = meta.add_run(value); vr.font.size = Pt(12)

doc.add_paragraph()
add_table(
    headers=['Student Name', 'Student ID'],
    rows=[
        ['Emad Saeed Alzahrani',   '2236570'],
        ['Talal Abid Alharbi',     '2236482'],
        ['Thamer Musaad Alkatheri','2236532'],
    ],
    caption='Project Team'
)
doc.add_page_break()

# =============================================================================
#  TABLE OF CONTENTS  (static)
# =============================================================================
add_heading('Table of Contents', level=1)
toc_entries = [
    # num, title, is_top_level
    ('1.',     'Introduction',                                              True),
    ('   1.1', 'Project Overview',                                         False),
    ('   1.2', 'Platform Overview',                                        False),
    ('   1.3', 'Development Timeline Overview',                            False),
    ('   1.4', 'Technology Stack',                                         False),
    ('',       '',                                                          False),
    ('',       'BACKEND FOUNDATION',                                        False),
    ('2.',     'Week 1: Project Initialization & Core Setup',              True),
    ('   2.1', 'Overview',                                                 False),
    ('   2.2', 'Backend Foundation Tasks',                                 False),
    ('3.',     'Week 2: Authentication, Users & Role-Based Access Control',True),
    ('   3.1', 'Overview',                                                 False),
    ('   3.2', 'Task 1 - User, Role & Company Models',                    False),
    ('   3.3', 'Task 2 - JWT Authentication',                             False),
    ('   3.4', 'Task 3 - RBAC System',                                    False),
    ('   3.5', 'Task 4 - User Management APIs',                           False),
    ('4.',     'Week 3: Campaigns, Quizzes & Simulation Logic',           True),
    ('   4.1', 'Overview',                                                 False),
    ('   4.2', 'Database Schema Implementation',                          False),
    ('   4.3', 'Simulation System Implementation',                        False),
    ('   4.4', 'Campaign Management APIs',                                False),
    ('   4.5', 'Quiz Assignment Logic',                                   False),
    ('   4.6', 'Quiz Taking & Submission APIs',                           False),
    ('5.',     'Week 4: Email Tracking, Risk Scoring & Remediation',      True),
    ('   5.1', 'Overview',                                                 False),
    ('   5.2', 'Email Template Library',                                  False),
    ('   5.3', 'Tracking Pixel Logic',                                    False),
    ('   5.4', 'Lure Link Tracking',                                      False),
    ('   5.5', '"You\'ve Been Caught" Landing Page',                      False),
    ('   5.6', 'Risk Score & Remediation Models',                         False),
    ('   5.7', 'Risk Scoring Algorithm',                                  False),
    ('6.',     'Week 5: Gamification, Dashboards & Analytics',            True),
    ('   6.1', 'Overview',                                                 False),
    ('   6.2', 'Task 1 - Auto-Remediation Assignment Logic',              False),
    ('   6.3', 'Task 2 - LMS APIs',                                       False),
    ('   6.4', 'Task 3 - Leaderboard, Badge & Achievement Tables',       False),
    ('   6.5', 'Task 4 - Badge Award System',                             False),
    ('   6.6', 'Task 5 - Leaderboard APIs',                               False),
    ('   6.7', 'Task 6 - Employee Dashboard APIs',                        False),
    ('   6.8', 'Task 7 - Admin Analytics APIs',                           False),
    ('   6.9', 'Task 8 - Public Content APIs',                            False),
    ('   6.10','Week 5 Deliverables Summary',                             False),
    ('',       '',                                                          False),
    ('',       'FRONTEND DEVELOPMENT',                                      False),
    ('7.',     'Frontend Design Overview & Principles',                   True),
    ('   7.1', 'Design Philosophy',                                        False),
    ('   7.2', 'Color System',                                             False),
    ('   7.3', 'Typography',                                               False),
    ('   7.4', 'Logo & Branding',                                         False),
    ('8.',     'Week 9: Foundation & Authentication (Frontend)',           True),
    ('   8.1', 'React.js Environment Setup',                              False),
    ('   8.2', 'Bilingual Configuration (i18next)',                       False),
    ('   8.3', 'Axios API Service Layer',                                  False),
    ('   8.4', 'Login Page Design',                                        False),
    ('   8.5', 'Company Registration Page',                               False),
    ('   8.6', 'Protected Route System',                                  False),
    ('9.',     'Week 10: Admin Portal & Campaign Management',             True),
    ('   9.1', 'Admin Main Layout',                                        False),
    ('   9.2', 'Campaign Management Interface',                           False),
    ('   9.3', 'Template Library Interface',                              False),
    ('   9.4', 'Analytics Dashboard',                                     False),
    ('   9.5', 'Employee Management Interface',                           False),
    ('   9.6', 'AI Email Generation Feature',                             False),
    ('10.',    'Week 11: Employee Portal, Public Portal & Gamification',  True),
    ('   10.1','Employee Dashboard',                                      False),
    ('   10.2','Quiz Interface',                                          False),
    ('   10.3','Training Module Interface',                               False),
    ('   10.4','Gamification Features',                                   False),
    ('   10.5','Public Portal Homepage',                                  False),
    ('   10.6','Awareness Content Pages',                                 False),
    ('   10.7','Public Quiz Interface',                                   False),
    ('11.',    'Responsive Design & Accessibility',                       True),
    ('',       '',                                                          False),
    ('',       'TESTING & CONCLUSION',                                      False),
    ('12.',    'Backend Testing Documentation',                            True),
    ('   12.1','Authentication & Authorization',                          False),
    ('   12.2','Company Management',                                      False),
    ('   12.3','Employee Management',                                     False),
    ('   12.4','Training Module System',                                  False),
    ('   12.5','Phishing Simulation System',                              False),
    ('   12.6','Security Testing',                                        False),
    ('   12.7','Results Summary',                                         False),
    ('13.',    'Conclusion',                                               True),
    ('',       '',                                                          False),
    ('Appendix A.', 'Screenshot Index',                                   False),
]

for num, title, is_top in toc_entries:
    if not num and not title:
        doc.add_paragraph()
        continue
    if not num and title:
        p = doc.add_paragraph()
        r = p.add_run(f'  {title}')
        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = MID_BLUE
        continue
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0 if is_top else 0.3)
    r1 = p.add_run(f'{num}  ')
    r1.font.bold = is_top; r1.font.size = Pt(11)
    r2 = p.add_run(title)
    r2.font.size = Pt(11)

doc.add_page_break()

# ── Table of Figures (Word auto-update field) ─────────────────────────────────
add_heading('Table of Figures', level=1)
add_para('Right-click the field below and select "Update Field" to populate.',
         italic=True, color=GRAY_700)
add_tof_field(' TOC \\h \\z \\t "Caption,1" ')
doc.add_paragraph()

# ── Table of Tables (static list) ─────────────────────────────────────────────
add_heading('Table of Tables', level=1)
table_list = [
    ('Table 1.1',  'Development Timeline Summary'),
    ('Table 1.2',  'Complete Technology Stack'),
    ('Table 2.1',  'Week 1 Backend Deliverables'),
    ('Table 2.2',  'Core Configuration Files'),
    ('Table 3.1',  'Week 2 Deliverables Summary'),
    ('Table 3.2',  'JWT Authentication Endpoints'),
    ('Table 3.3',  'RBAC Roles & Permission Classes'),
    ('Table 4.1',  'Week 3 Deliverables'),
    ('Table 4.2',  'Campaign Model Fields'),
    ('Table 4.3',  'Quiz & Assessment Models'),
    ('Table 4.4',  'Campaign Management API Endpoints'),
    ('Table 4.5',  'Quiz Submission API Endpoints'),
    ('Table 5.1',  'Week 4 Deliverables'),
    ('Table 5.2',  'Risk Level to Remediation Mapping'),
    ('Table 5.3',  'Email Tracking Serializers'),
    ('Table 6.1',  'Week 5 Deliverables Summary'),
    ('Table 6.2',  'Badge Award Conditions'),
    ('Table 6.3',  'Platform Event Signal Sources'),
    ('Table 6.4',  'Points Award Schedule'),
    ('Table 6.5',  'Leaderboard & Badge API Endpoints'),
    ('Table 6.6',  'Employee Dashboard API Sections'),
    ('Table 6.7',  'Admin Analytics API Endpoints'),
    ('Table 6.8',  'Public Content API Endpoints'),
    ('Table 7.1',  'Core Design Principles'),
    ('Table 7.2',  'Color Usage Guide'),
    ('Table 7.3',  'Typography Scale'),
    ('Table 8.1',  'Week 9 Frontend Deliverables'),
    ('Table 8.2',  'Language Support Features'),
    ('Table 8.3',  'API Service Files'),
    ('Table 8.4',  'Login Form Validation Rules'),
    ('Table 8.5',  'Registration Form Fields'),
    ('Table 8.6',  'Route Protection Matrix'),
    ('Table 9.1',  'Week 10 Frontend Deliverables'),
    ('Table 9.2',  'Admin Navigation Structure'),
    ('Table 9.3',  'Campaign Status Indicators'),
    ('Table 9.4',  'Template Categories'),
    ('Table 9.5',  'Analytics Metrics Displayed'),
    ('Table 9.6',  'Employee Table Columns'),
    ('Table 9.7',  'AI Generation Parameters'),
    ('Table 10.1', 'Week 11 Frontend Deliverables'),
    ('Table 10.2', 'Employee Dashboard Widgets'),
    ('Table 10.3', 'Quiz Feedback Components'),
    ('Table 10.4', 'Training Module Components'),
    ('Table 10.5', 'Available Badges'),
    ('Table 10.6', 'Public Homepage Sections'),
    ('Table 10.7', 'Awareness Content Structure'),
    ('Table 10.8', 'Public vs Employee Quiz Experience'),
    ('Table 11.1', 'Responsive Breakpoints'),
    ('Table 11.2', 'Accessibility Features'),
    ('Table 12.1', 'Backend Testing Results Summary'),
    ('Table A.1',  'Complete Screenshot Index'),
]
for tnum, ttitle in table_list:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r1 = p.add_run(f'{tnum}  ')
    r1.font.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(ttitle); r2.font.size = Pt(10)
doc.add_page_break()

# =============================================================================
#  1. INTRODUCTION
# =============================================================================
add_heading('1. Introduction', level=1)

add_heading('1.1 Project Overview', level=2)
add_para(
    'PhishAware is a comprehensive cybersecurity awareness platform designed for Saudi '
    'Arabian and MENA region organisations. The system enables administrators to conduct '
    'realistic phishing simulation campaigns, assign targeted training modules, track '
    'employee risk scores, and gamify the learning experience to drive sustained engagement. '
    'The platform was developed as a final-year engineering project at university, '
    'supervised by Dr. Saim Rasheed.'
)

add_heading('1.2 Platform Overview', level=2)
add_para(
    'PhishAware operates across three distinct portals serving different user roles: a '
    'Public portal accessible without login that delivers general cybersecurity awareness '
    'content; an Employee portal where enrolled staff complete assigned training and '
    'receive personalised risk feedback; and a Company Admin portal where administrators '
    'manage campaigns, templates, employees, and analytics. A Super-Admin role provides '
    'platform-level oversight across all organisations.'
)

add_heading('1.3 Development Timeline Overview', level=2)
add_para('The project was executed across two parallel tracks:')
add_bullet('Backend Foundation (Weeks 1-5): Django REST API, database models, business logic',
           bold_prefix='Backend:', color=DARK_BLUE)
add_bullet('Frontend Development (Weeks 9-11): React.js interfaces across all three portals',
           bold_prefix='Frontend:', color=MID_BLUE)
add_table(
    headers=['Phase', 'Weeks', 'Key Deliverables'],
    rows=[
        ['Backend Foundation',   'Weeks 1-5',  'Django setup, auth, campaigns, tracking, gamification, testing'],
        ['Frontend Development', 'Weeks 9-11', 'React setup, admin portal, employee portal, public portal'],
    ],
    caption='Table 1.1: Development Timeline Summary',
)

add_heading('1.4 Technology Stack', level=2)
add_table(
    headers=['Layer', 'Technology', 'Version', 'Purpose'],
    rows=[
        ['Backend',  'Python / Django',     '4.x',    'Server-side framework'],
        ['Backend',  'Django REST Framework','3.14+',  'RESTful API layer'],
        ['Backend',  'PostgreSQL',           '15+',    'Production database'],
        ['Backend',  'Simple JWT',           '5.x',    'JWT authentication tokens'],
        ['Backend',  'SendGrid',             'SMTP',   'Transactional email delivery'],
        ['Frontend', 'React.js',             '18.2+',  'UI component framework'],
        ['Frontend', 'React Router',         '6.x',    'Client-side routing'],
        ['Frontend', 'Tailwind CSS',         '3.x',    'Utility-first styling'],
        ['Frontend', 'i18next',              '23.x',   'Arabic/English internationalisation'],
        ['Frontend', 'Axios',               '1.x',    'HTTP client and API layer'],
        ['Frontend', 'Recharts',            '2.x',    'Data visualisation charts'],
    ],
    caption='Table 1.2: Complete Technology Stack',
)

# =============================================================================
#  SECTION DIVIDER – BACKEND
# =============================================================================
doc.add_page_break()
div = doc.add_paragraph()
div.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = div.add_run('BACKEND FOUNDATION')
dr.font.size = Pt(18); dr.font.bold = True; dr.font.color.rgb = DARK_BLUE
doc.add_paragraph()

# =============================================================================
#  2. WEEK 1 – PROJECT INITIALIZATION & CORE SETUP
# =============================================================================
add_heading('2. Week 1: Project Initialization & Core Setup', level=1)

add_heading('2.1 Overview', level=2)
add_para(
    'Week 1 focused on setting up the backend development environment and establishing '
    'the core Django project structure. The goal was to create a maintainable, scalable '
    'foundation that all subsequent weeks could build upon.'
)

add_heading('2.2 Backend Foundation Tasks', level=2)
add_deliverables_table(
    rows=[
        ['Design complete ER diagram (all modules)', 'Database', 'Complete', 'Full entity-relationship diagram covering all apps'],
        ['Setup SQLite for development',             'Database', 'Complete', 'Local development database configuration'],
        ['Setup Django + DRF environment',           'Backend',  'Complete', 'Django project with Django REST Framework installed'],
        ['Initialize Django project structure',      'Backend',  'Complete', 'Root configuration files, URL routing'],
        ['Create Django apps structure',             'Backend',  'Complete', 'Separate apps for logical separation of concerns'],
    ],
    caption='Table 2.1: Week 1 Backend Deliverables',
)

add_para('The Django project was structured into five dedicated apps:', bold=False)
for app, desc in [
    ('simulations', 'Phishing email templates, simulation campaigns, tracking events'),
    ('training',    'Training modules, assignments, quiz questions and results'),
    ('campaigns',   'Campaign orchestration linking simulations and training'),
    ('companies',   'Multi-tenant company and employee management'),
    ('core',        'Shared utilities, permissions, email helpers, signals'),
]:
    add_bullet(f'{desc}', bold_prefix=app)

add_table(
    headers=['Configuration File', 'Purpose'],
    rows=[
        ['settings.py',   'Database configuration, installed apps, middleware, email settings'],
        ['urls.py',       'Root URL routing to each app\'s endpoint namespace'],
        ['wsgi.py',       'WSGI server entry point for production deployment'],
        ['asgi.py',       'ASGI server entry point for async support'],
        ['.env',          'Environment variables: SECRET_KEY, DATABASE_URL, SENDGRID_KEY, FRONTEND_URL'],
    ],
    caption='Table 2.2: Core Configuration Files',
)

# =============================================================================
#  3. WEEK 2 – AUTHENTICATION, USERS & RBAC
# =============================================================================
doc.add_page_break()
add_heading('3. Week 2: Authentication, Users & Role-Based Access Control', level=1)

add_heading('3.1 Overview', level=2)
add_para(
    'Week 2 implemented the complete security and identity layer. This included the '
    'multi-tenant Company model, a custom User model with four distinct roles, stateless '
    'JWT authentication, and a role-based access control (RBAC) system enforced on '
    'every API endpoint.'
)
add_deliverables_table(
    rows=[
        ['Create User, Role & Company tables', 'Database', 'Complete', 'Custom User model + Company multi-tenant root entity'],
        ['Implement JWT authentication',        'Backend',  'Complete', 'Login, register, refresh, logout, profile endpoints'],
        ['Implement RBAC system',               'Backend',  'Complete', '4 roles with dedicated permission classes'],
        ['Build user management APIs',          'Backend',  'Complete', 'Full CRUD for users with role-filtered access'],
    ],
    caption='Table 3.1: Week 2 Deliverables Summary',
)

add_heading('3.2 Task 1 - User, Role & Company Models', level=2)
add_para(
    'The Company model is the root multi-tenant entity. Every user, campaign, and '
    'training record is scoped to a company. The custom User model extends '
    'AbstractBaseUser and PermissionsMixin, storing role, company FK, department, '
    'email verification status, and invitation token fields.'
)

add_heading('3.3 Task 2 - JWT Authentication', level=2)
add_para(
    'Authentication uses djangorestframework-simplejwt. A custom serializer '
    '(CustomTokenObtainPairSerializer) injects the user\'s role, company_id, and '
    'is_verified flag directly into the JWT payload, eliminating a second database '
    'lookup on every request.'
)
add_table(
    headers=['Method', 'Endpoint', 'Access', 'Description'],
    rows=[
        ['POST', '/api/v1/auth/register/',        'Public',  'Create new company + admin account'],
        ['POST', '/api/v1/auth/login/',            'Public',  'Returns access + refresh JWT pair'],
        ['POST', '/api/v1/auth/token/refresh/',    'Public',  'Exchange refresh token for new access token'],
        ['POST', '/api/v1/auth/logout/',           'Auth',    'Blacklists refresh token'],
        ['GET/PATCH', '/api/v1/auth/profile/',     'Auth',    'Retrieve or update own profile'],
        ['POST', '/api/v1/auth/change-password/',  'Auth',    'Validates old password, sets new one'],
        ['POST', '/api/v1/auth/verify-email/<uuid>/', 'Public', 'Verify email address via token link'],
        ['POST', '/api/v1/auth/resend-verification/','Auth', 'Re-send verification email'],
    ],
    caption='Table 3.2: JWT Authentication Endpoints',
)

add_heading('3.4 Task 3 - RBAC System', level=2)
add_para(
    'The RBAC system is implemented in apps/core/permissions.py. Four role-specific '
    'permission classes are supplemented by three composite classes covering cross-role '
    'scenarios.'
)
add_table(
    headers=['Role', 'Permission Class', 'Scope'],
    rows=[
        ['SUPER_ADMIN',   'IsSuperAdmin',          'Platform-wide, no company restrictions'],
        ['COMPANY_ADMIN', 'IsCompanyAdmin',         'Full control within own company only'],
        ['EMPLOYEE',      'IsEmployee',             'Own assignments, quiz results, training'],
        ['PUBLIC_USER',   'IsPublicUser',           'Public content, no personal data'],
        ['—',             'IsAdminOrSuperAdmin',    'Combined admin access'],
        ['—',             'IsEmployeeOrAdmin',      'Any authenticated company user'],
        ['—',             'IsOwnerOrAdmin',         'Own records or admin override'],
    ],
    caption='Table 3.3: RBAC Roles & Permission Classes',
)

add_heading('3.5 Task 4 - User Management APIs', level=2)
add_para(
    'Full CRUD endpoints for user management are exposed through dedicated ViewSets '
    'with role-filtered querysets. Admins can only see users within their own company; '
    'Super-Admins see all. Serializer selection adapts based on the requesting role '
    '(full detail vs. summary view).'
)

# =============================================================================
#  4. WEEK 3 – CAMPAIGNS, QUIZZES & SIMULATION LOGIC
# =============================================================================
doc.add_page_break()
add_heading('4. Week 3: Campaigns, Quizzes & Simulation Logic', level=1)

add_heading('4.1 Overview', level=2)
add_para(
    'Week 3 introduced the core business logic: phishing simulation models, campaign '
    'orchestration, quiz delivery, and automated quiz assignment triggered by simulation '
    'interactions.'
)
add_deliverables_table(
    rows=[
        ['Create Campaign, Quiz, QuizResult tables',              'Database', 'Complete', 'Core assessment models'],
        ['Create EmailTemplate, SimulationLog, TrackingEvent tables','Database','Complete','Simulation tracking models'],
        ['Build campaign creation & management APIs',             'Backend',  'Complete', 'CRUD endpoints for campaigns'],
        ['Implement quiz assignment logic',                       'Backend',  'Complete', 'Automated remediation triggers'],
        ['Build quiz taking & submission APIs',                   'Backend',  'Complete', 'Start, submit, results endpoints'],
    ],
    caption='Table 4.1: Week 3 Deliverables',
)

add_heading('4.2 Database Schema Implementation', level=2)
add_heading('4.2.1 Campaign Module', level=3)
add_para(
    'Location: Backend/apps/campaigns/  - '
    'The Campaign model stores metadata (title, description, status), lifecycle dates, '
    'and foreign keys to the company and creating admin. It links outward to '
    'SimulationCampaign records and training assignments.'
)
add_table(
    headers=['Field', 'Type', 'Purpose'],
    rows=[
        ['title',        'CharField',     'Campaign name displayed in admin UI'],
        ['description',  'TextField',     'Optional campaign notes'],
        ['status',       'CharField',     'DRAFT / ACTIVE / COMPLETED / PAUSED'],
        ['company',      'ForeignKey',    'Owning company (multi-tenant scope)'],
        ['created_by',   'ForeignKey',    'Admin user who created the campaign'],
        ['start_date',   'DateTimeField', 'Campaign activation date'],
        ['end_date',     'DateTimeField', 'Campaign close date'],
    ],
    caption='Table 4.2: Campaign Model Fields',
)

add_heading('4.2.2 Assessment Module (Quizzes)', level=3)
add_para(
    'Location: Backend/apps/assessments/  - '
    'Three interconnected models support the full quiz lifecycle: Quiz (general info '
    'and pass threshold), QuizQuestion (individual questions with multiple-choice '
    'options), and QuizResult (employee attempt records with score and pass/fail).'
)
add_table(
    headers=['Model', 'Key Fields', 'Purpose'],
    rows=[
        ['Quiz',         'title, pass_score, time_limit, company',           'Quiz definition and configuration'],
        ['QuizQuestion', 'quiz, question_text, options (JSON), correct_answer','Individual questions per quiz'],
        ['QuizResult',   'quiz, employee, score, passed, completed_at',       'Employee attempt record'],
    ],
    caption='Table 4.3: Quiz & Assessment Models',
)

add_heading('4.3 Simulation System Implementation', level=2)
add_para(
    'Location: Backend/apps/simulations/  - '
    'The simulation system uses four cooperating models. SimulationTemplate stores '
    'reusable phishing email HTML with placeholder tokens ({TRACKING_PIXEL}, '
    '{LURE_LINK}, {EMPLOYEE_NAME}). EmailSimulation records each individual email sent. '
    'TrackingEvent captures every user interaction. SimulationCampaign aggregates '
    'denormalised statistics (total_sent, total_clicked, total_reported) for fast '
    'dashboard queries.'
)

add_heading('4.4 Campaign Management APIs', level=2)
add_table(
    headers=['Method', 'Endpoint', 'Description'],
    rows=[
        ['GET/POST',       '/api/v1/campaigns/',            'List all or create new campaign'],
        ['GET/PUT/DELETE', '/api/v1/campaigns/{id}/',        'Retrieve, update, or delete campaign'],
        ['POST',           '/api/v1/simulations/{id}/send/','Send campaign emails to selected employees'],
        ['GET',            '/api/v1/simulations/{id}/stats/','Live click/open/report statistics'],
    ],
    caption='Table 4.4: Campaign Management API Endpoints',
)

add_heading('4.5 Quiz Assignment Logic', level=2)
add_para(
    'Automated remediation: a Django post_save signal on TrackingEvent fires whenever '
    'an employee clicks a phishing link or opens a suspicious attachment. The signal '
    'calls the assignment service which selects the appropriate quiz module based on '
    'the simulation category and assigns it with a due date.'
)
add_para(
    'Manual assignment: admins can also assign any quiz to any employee or department '
    'group through the bulk assignment API, with optional force_assign flag to override '
    'warnings about incomplete prior assignments.'
)

add_heading('4.6 Quiz Taking & Submission APIs', level=2)
add_table(
    headers=['Method', 'Endpoint', 'Description'],
    rows=[
        ['POST', '/api/v1/quizzes/{id}/start/',    'Create a quiz attempt record, return questions'],
        ['POST', '/api/v1/quizzes/{id}/submit/',   'Score submission, update QuizResult, trigger badge check'],
        ['GET',  '/api/v1/quizzes/{id}/result/',   'Retrieve result for own last attempt'],
        ['GET',  '/api/v1/training/assignments/',  'List all assigned training for current user'],
        ['POST', '/api/v1/training/assignments/{id}/complete/', 'Mark training as completed'],
    ],
    caption='Table 4.5: Quiz Submission API Endpoints',
)

# =============================================================================
#  5. WEEK 4 – EMAIL TRACKING, RISK SCORING & REMEDIATION
# =============================================================================
doc.add_page_break()
add_heading('5. Week 4: Email Tracking, Risk Scoring & Remediation', level=1)

add_heading('5.1 Overview', level=2)
add_para(
    'Week 4 added phishing simulation delivery mechanics: a tracking pixel for '
    'email open detection, unique lure links for click tracking, a "You\'ve Been Caught" '
    'educational landing page, and the risk scoring engine that automatically adjusts '
    'each employee\'s risk score based on their interactions.'
)
add_deliverables_table(
    rows=[
        ['Build email template library CRUD', 'Backend',  'Complete', 'Template upload, edit, delete, preview'],
        ['Implement tracking pixel logic',     'Backend',  'Complete', '1x1 transparent GIF embedded in email'],
        ['Implement lure link tracking',       'Backend',  'Complete', 'Unique per-employee redirect links'],
        ['Create "You\'ve been caught" page',  'Backend',  'Complete', 'Educational landing page after click'],
        ['Create RiskScore, Remediation tables','Database','Complete', 'Score model + remediation training model'],
        ['Implement risk scoring algorithm',   'Backend',  'Complete', 'Signal-driven score recalculation'],
    ],
    caption='Table 5.1: Week 4 Deliverables',
)

add_heading('5.2 Email Template Library', level=2)
add_para(
    'Admins can upload, edit, and delete phishing email templates. Each template stores '
    'an HTML body with placeholder tokens, a plain-text fallback, a sender name/email, '
    'subject line, category (Finance/HR/IT/Delivery/Government), difficulty level '
    '(Easy/Medium/Hard), and an optional company FK (NULL = platform-wide template).'
)

add_heading('5.3 Tracking Pixel Logic', level=2)
add_para(
    'A 1x1 transparent GIF is served at /api/v1/simulations/track/open/{token}/. '
    'When the email client loads this image, a TrackingEvent of type EMAIL_OPENED is '
    'created. The token is unique per EmailSimulation record, preventing cross-employee '
    'attribution errors.'
)

add_heading('5.4 Lure Link Tracking', level=2)
add_para(
    'Each simulation email contains a unique lure link: /api/v1/simulations/track/'
    'click/{link_token}/. When clicked, the view records a LINK_CLICKED TrackingEvent, '
    'then redirects to {FRONTEND_URL}/simulation/caught/{link_token} — the educational '
    'landing page.'
)

add_heading('5.5 "You\'ve Been Caught" Landing Page', level=2)
add_para(
    'The SimulationCaught React page (at /simulation/caught/:token) calls '
    '/api/v1/simulations/feedback/{link_token}/ to retrieve the simulation\'s red flags '
    'and explanation in the user\'s current language. The page displays this as an '
    'educational debrief explaining exactly why the email was a phishing attempt.'
)

add_heading('5.6 Risk Score & Remediation Models', level=2)
add_para(
    'The RiskScore model stores a rolling numeric score (0-100) for each employee. '
    'RemediationTraining links a risk score range to specific training modules. When '
    'an employee\'s score crosses a threshold, the auto-remediation system (Week 5) '
    'automatically assigns the relevant training.'
)
add_table(
    headers=['Risk Level', 'Score Range', 'Modules Assigned'],
    rows=[
        ['Low',    '0-33',  'Optional awareness content'],
        ['Medium', '34-66', 'Email phishing + one category-specific module'],
        ['High',   '67-84', 'Email phishing + mobile security + social engineering'],
        ['Critical','85-100','All available modules, expedited deadline (3 days)'],
    ],
    caption='Table 5.2: Risk Level to Remediation Mapping',
)
add_table(
    headers=['Serializer', 'Purpose'],
    rows=[
        ['RiskScoreSerializer',           'Expose current score, level, and trend to employee dashboard'],
        ['TrackingEventSerializer',       'Log interaction events with simulation FK and event type'],
        ['EmailSimulationSerializer',     'List sent simulations per campaign with status'],
        ['SimulationFeedbackSerializer',  'Return red_flags list and explanation for caught page'],
        ['RemediationTrainingSerializer', 'Map risk thresholds to training module assignments'],
    ],
    caption='Table 5.3: Email Tracking Serializers',
)

add_heading('5.7 Risk Scoring Algorithm', level=2)
add_para(
    'The algorithm runs inside recalculate_score() called from a post_save signal on '
    'TrackingEvent. Score contributors: each LINK_CLICKED event +15 points, each '
    'EMAIL_OPENED +5, each completed training module -10, each passed quiz -8. '
    'Score is clamped to 0-100. If total_simulations_received is 0 the simulation '
    'adjustment step is skipped to avoid division errors.'
)

# =============================================================================
#  6. WEEK 5 – GAMIFICATION, DASHBOARDS & ANALYTICS
# =============================================================================
doc.add_page_break()
add_heading('6. Week 5: Gamification, Dashboards & Analytics', level=1)

add_heading('6.1 Overview', level=2)
add_para(
    'Week 5 completed the backend with the gamification layer (badges, points, '
    'leaderboard), the employee dashboard API, comprehensive admin analytics, and '
    'public content APIs requiring no authentication.'
)

add_heading('6.2 Task 1 - Auto-Remediation Assignment Logic', level=2)
add_para(
    'A post_save signal on TrackingEvent triggers the remediation assignment service. '
    'The service evaluates the employee\'s updated risk score against the '
    'RemediationTraining thresholds and creates TrainingAssignment records for any '
    'unassigned modules in the relevant tier, with a computed due date.'
)

add_heading('6.3 Task 2 - LMS APIs', level=2)
add_para(
    'The Learning Management System APIs provide full lifecycle management for training '
    'modules and assignments.'
)
add_para('Training Module APIs: list modules (with company filter), retrieve detail, '
         'create/update/delete (admin only), seed platform-wide templates.')
add_para('Assignment APIs: list own assignments (employee) or all assignments '
         '(admin), bulk assign to department, mark complete, force_assign override flag.')

add_heading('6.4 Task 3 - Leaderboard, Badge & Achievement Tables', level=2)
add_table(
    headers=['Model', 'Key Fields', 'Purpose'],
    rows=[
        ['Badge',              'name, icon, description, points_required, trigger_event', 'Badge definition'],
        ['EmployeeBadge',      'employee, badge, awarded_at',                             'Employee badge award record'],
        ['PointsTransaction',  'employee, points, reason, created_at',                    'Individual point event log'],
        ['EmployeePoints',     'employee, total_points, weekly_points, monthly_points',   'Aggregated point totals'],
    ],
    caption='Table 6.1: Week 5 Deliverables Summary',
)

add_heading('6.5 Task 4 - Badge Award System', level=2)
add_table(
    headers=['Badge Type', 'Award Condition'],
    rows=[
        ['First Steps',         'Complete first training module'],
        ['Quick Learner',       'Complete any training in under 10 minutes'],
        ['Perfect Score',       'Score 100% on any quiz'],
        ['Phish Detector',      'Correctly identify 10 phishing emails in quiz'],
        ['Reporter',            'Report a real phishing attempt'],
        ['Streak Master',       'Complete training 5 consecutive weeks'],
        ['Department Champion', 'Achieve highest score in own department'],
        ['Fast Responder',      'Complete remediation training within 24 hours of assignment'],
    ],
    caption='Table 6.2: Badge Award Conditions',
)
add_table(
    headers=['Platform Event', 'Signal Source', 'Points Awarded'],
    rows=[
        ['Training module completed', 'post_save TrainingAssignment', '+50'],
        ['Quiz passed',               'post_save QuizResult',         '+25'],
        ['Quiz perfect score',        'post_save QuizResult (score=100)', '+50 bonus'],
        ['Phishing email reported',   'POST /report/ endpoint',       '+10'],
        ['Badge earned',              'BadgeAwardService',            'Varies by badge'],
    ],
    caption='Table 6.3: Platform Event Signal Sources',
)
add_table(
    headers=['Action', 'Points'],
    rows=[
        ['Complete training module',          '+50'],
        ['Pass quiz (>=80%)',                  '+25'],
        ['Achieve 100% quiz score',            '+50 bonus'],
        ['Report phishing attempt',           '+10'],
        ['Log in 5 days in a row',            '+20'],
        ['First simulation received',         '+5'],
        ['Complete overdue assignment',       '+15'],
    ],
    caption='Table 6.4: Points Award Schedule',
)

add_heading('6.6 Task 5 - Leaderboard APIs', level=2)
add_table(
    headers=['Endpoint', 'Auth', 'Description'],
    rows=[
        ['GET /api/v1/gamification/leaderboard/',          'Employee+', 'Company leaderboard (week/month/all)'],
        ['GET /api/v1/gamification/leaderboard/my-rank/',  'Employee+', 'Current user rank and nearby peers'],
        ['GET /api/v1/gamification/badges/',               'Employee+', 'All badges with earned status'],
        ['GET /api/v1/gamification/badges/my-badges/',     'Employee+', 'Own earned badges with dates'],
        ['GET /api/v1/gamification/points/',               'Employee+', 'Own point total and transaction log'],
    ],
    caption='Table 6.5: Leaderboard & Badge API Endpoints',
)

add_heading('6.7 Task 6 - Employee Dashboard APIs', level=2)
add_table(
    headers=['Dashboard Section', 'Source Model', 'Endpoint'],
    rows=[
        ['Risk Score',         'RiskScore',           'GET /api/v1/risk/my-score/'],
        ['Training Progress',  'TrainingAssignment',  'GET /api/v1/training/assignments/'],
        ['Upcoming Deadlines', 'TrainingAssignment',  'GET /api/v1/training/assignments/?status=pending'],
        ['Recent Badges',      'EmployeeBadge',       'GET /api/v1/gamification/badges/my-badges/'],
        ['Leaderboard Rank',   'EmployeePoints',      'GET /api/v1/gamification/leaderboard/my-rank/'],
        ['Notifications',      'Notification',        'GET /api/v1/notifications/'],
    ],
    caption='Table 6.6: Employee Dashboard API Sections',
)

add_heading('6.8 Task 7 - Admin Analytics APIs', level=2)
add_table(
    headers=['Endpoint', 'Description'],
    rows=[
        ['GET /api/v1/analytics/overview/',          'KPI summary: campaigns, click rate, high-risk count'],
        ['GET /api/v1/analytics/campaigns/',         'Per-campaign click/open/report rates with trend'],
        ['GET /api/v1/analytics/departments/',       'Click rates and risk scores grouped by department'],
        ['GET /api/v1/analytics/risk-distribution/', 'Count of Low/Medium/High/Critical employees'],
        ['GET /api/v1/analytics/employees/',         'Individual employee risk and training stats'],
        ['GET /api/v1/analytics/export/',            'CSV/PDF export of selected metrics and date range'],
        ['GET /api/v1/super-admin/companies/',       'Super-admin: cross-company overview'],
    ],
    caption='Table 6.7: Admin Analytics API Endpoints',
)

add_heading('6.9 Task 8 - Public Content APIs', level=2)
add_table(
    headers=['Method', 'Endpoint', 'Description'],
    rows=[
        ['GET',  '/api/v1/public/awareness/',          'List all published awareness articles'],
        ['GET',  '/api/v1/public/awareness/{slug}/',   'Retrieve article by slug (AR/EN bilingual)'],
        ['GET',  '/api/v1/public/quiz/',               'List public quiz sets (no auth)'],
        ['POST', '/api/v1/public/quiz/{id}/submit/',   'Submit anonymous quiz answers'],
        ['GET',  '/api/v1/public/stats/',              'Platform statistics (user count, company count)'],
    ],
    caption='Table 6.8: Public Content API Endpoints',
)

add_heading('6.10 Week 5 Deliverables Summary', level=2)
add_deliverables_table(
    rows=[
        ['Build auto-remediation assignment logic', 'Backend',  'Complete', 'Signal-driven training assignment on risk threshold breach'],
        ['Create LMS APIs',                         'Backend',  'Complete', 'Module CRUD, bulk assign, complete, force_assign'],
        ['Create Leaderboard, Badge tables',        'Database', 'Complete', 'Badge, EmployeeBadge, PointsTransaction, EmployeePoints'],
        ['Implement badge award system',            'Backend',  'Complete', 'Event-driven badge checks on every qualifying action'],
        ['Build leaderboard APIs',                  'Backend',  'Complete', 'Weekly/monthly/all-time with my-rank endpoint'],
        ['Create employee dashboard APIs',          'Backend',  'Complete', 'Aggregated risk, training, badges, rank in one call'],
        ['Build admin analytics APIs',              'Backend',  'Complete', 'KPIs, department stats, risk distribution, CSV export'],
        ['Create public content APIs',              'Backend',  'Complete', 'Awareness articles, public quiz, platform stats'],
        ['Install PostgreSQL',                      'Database', 'Complete', 'Migrate from SQLite to PostgreSQL for production'],
        ['Run backend testing',                     'Backend',  'Complete', '234 test cases, 104 endpoints covered, 17 defects fixed'],
    ],
    caption='Table 6.9: Week 5 Deliverables (Backend Complete)',
)

# =============================================================================
#  SECTION DIVIDER – FRONTEND
# =============================================================================
doc.add_page_break()
div2 = doc.add_paragraph()
div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr2 = div2.add_run('FRONTEND DEVELOPMENT')
dr2.font.size = Pt(18); dr2.font.bold = True; dr2.font.color.rgb = MID_BLUE
doc.add_paragraph()

# =============================================================================
#  7. FRONTEND DESIGN OVERVIEW & PRINCIPLES
# =============================================================================
add_heading('7. Frontend Design Overview & Principles', level=1)

add_heading('7.1 Design Philosophy', level=2)
add_para('PhishAware\'s interface design follows four core principles:')
add_table(
    headers=['Principle', 'Description', 'Implementation Examples'],
    rows=[
        ['Clarity',              'Clear visual hierarchy and intuitive navigation',         'Consistent spacing, typography, colour-coded status indicators'],
        ['Accessibility',        'WCAG 2.1 AA compliance for inclusive design',             'High contrast ratios, keyboard navigation, screen reader support'],
        ['Bilingual Excellence', 'Seamless Arabic/English experience',                      'RTL/LTR layouts, culturally appropriate translations, language toggle'],
        ['Cultural Relevance',   'Saudi/MENA-specific context',                             'Examples using Nafath, Absher, Saudi banks, local government services'],
    ],
    caption='Table 7.1: Core Design Principles',
)

add_heading('7.2 Color System', level=2)
add_figure('Figure 7.1', 'PhishAware Color Palette - colour swatches with hex codes and labels')
add_table(
    headers=['Color', 'Hex Code', 'Primary Usage'],
    rows=[
        ['Primary Blue',   '#3B82F6', 'Buttons, links, active states, CTAs'],
        ['Dark Blue',      '#1E3A8A', 'Headers, navigation, headings'],
        ['Success Green',  '#10B981', 'Completed, passed, success states'],
        ['Warning Yellow', '#F59E0B', 'Pending, in-progress, due soon'],
        ['Danger Red',     '#EF4444', 'Failed, errors, delete actions'],
        ['Gray 600',       '#6B7280', 'Secondary text, metadata, labels'],
        ['Gray 200',       '#E5E7EB', 'Card borders, table lines, dividers'],
    ],
    caption='Table 7.2: Color Usage Guide',
)

add_heading('7.3 Typography', level=2)
add_figure('Figure 7.2', 'Typography Hierarchy - heading levels in English (Inter) and Arabic (Tajawal)')
add_table(
    headers=['Element', 'English Font', 'Arabic Font', 'Size', 'Weight'],
    rows=[
        ['H1 - Page Title', 'Inter', 'Tajawal', '32px', 'Bold 700'],
        ['H2 - Section',    'Inter', 'Tajawal', '24px', 'Semibold 600'],
        ['H3 - Subsection', 'Inter', 'Tajawal', '20px', 'Semibold 600'],
        ['Body Text',       'Inter', 'Tajawal', '16px', 'Regular 400'],
        ['Small / Meta',    'Inter', 'Tajawal', '14px', 'Regular 400'],
        ['Labels / Tags',   'Inter', 'Tajawal', '12px', 'Medium 500'],
    ],
    caption='Table 7.3: Typography Scale',
)

add_heading('7.4 Logo & Branding', level=2)
add_figure('Figure 7.3', 'PhishAware Logo System - icon-only, horizontal, vertical, and light/dark variants')
add_para(
    'The PhishAware logo combines a security shield (protection) with a fishing hook '
    '(phishing threat awareness). Three variants are provided: icon-only for small '
    'spaces, horizontal with wordmark for headers, and a white version for dark '
    'backgrounds.'
)

# =============================================================================
#  8. WEEK 9 – FRONTEND FOUNDATION & AUTHENTICATION
# =============================================================================
doc.add_page_break()
add_heading('8. Week 9: Foundation & Authentication (Frontend)', level=1)

add_para('Week 9 established the technical foundation and core authentication system, '
         'enabling secure user access across all three roles.')
add_deliverables_table(
    rows=[
        ['Setup React.js environment',       'Frontend', 'Complete', 'Vite initialisation, folder structure, routing'],
        ['Configure i18next for bilingual',  'Frontend', 'Complete', 'Arabic/English with automatic RTL layout'],
        ['Create Axios API service layer',   'Frontend', 'Complete', 'Base URL, JWT interceptors, error handling'],
        ['Build authentication pages',       'Frontend', 'Complete', 'Login and company registration pages'],
        ['Create protected route system',    'Frontend', 'Complete', 'Role-based route guard component'],
    ],
    caption='Table 8.1: Week 9 Frontend Deliverables',
)

add_heading('8.1 React.js Environment Setup', level=2)
add_para(
    'The project was initialised using Vite for fast build tooling. React Router v6 '
    'was configured for client-side routing with a feature-module folder structure '
    'organised into public/, employee/, and admin/ portal sub-trees.'
)
add_figure('Figure 8.1', 'Project Folder Structure - VS Code file explorer showing feature-module layout')

add_heading('8.2 Bilingual Configuration (i18next)', level=2)
add_para(
    'All interface text is externalised into en.json and ar.json translation files. '
    'When Arabic is selected the document direction switches to dir="rtl" and the '
    'entire layout mirrors automatically. The selected language persists in localStorage.'
)
add_figure('Figure 8.2', 'Language Switcher - AR/EN toggle in the navigation bar')
add_figure('Figure 8.3', 'RTL Layout Transformation - side-by-side English (LTR) and Arabic (RTL) views')
add_table(
    headers=['Feature', 'Implementation', 'Example'],
    rows=[
        ['Language Toggle',   'Dropdown in header',          'AR / EN selector button'],
        ['RTL Layout',        'CSS direction attribute',     'dir="rtl" applied to <html> for Arabic'],
        ['Translation Files', 'JSON key-value pairs',        'en.json and ar.json with nested keys'],
        ['Persistent Choice', 'localStorage',                'Remembered across page reloads and sessions'],
    ],
    caption='Table 8.2: Language Support Features',
)

add_heading('8.3 Axios API Service Layer', level=2)
add_para(
    'A centralised Axios instance points to the Django backend. A request interceptor '
    'attaches the JWT access token from localStorage to every outgoing request. '
    'A response interceptor intercepts HTTP 401 errors and silently refreshes the '
    'token before retrying the original request.'
)
add_figure('Figure 8.4', 'API Configuration - Axios instance with request/response interceptors')
add_table(
    headers=['Service File', 'Key Endpoints Covered'],
    rows=[
        ['api/auth.js',         'login, register, verify, refresh, logout, profile'],
        ['api/training.js',     'getModules, assign, complete, quiz start/submit'],
        ['api/campaigns.js',    'create, send, track, getStats'],
        ['api/employees.js',    'list, invite, remove, getRiskScore'],
        ['api/gamification.js', 'badges, leaderboard, points, myRank'],
    ],
    caption='Table 8.3: API Service Files',
)

add_heading('8.4 Login Page Design', level=2)
add_figure('Figure 8.5', 'Login Page - English version with form fields and Sign In button')
add_figure('Figure 8.6', 'Login Page - Arabic version demonstrating complete RTL transformation')
add_figure('Figure 8.7', 'Login Error States - red borders and inline validation messages')
add_table(
    headers=['Field', 'Validation Rules', 'Error Message'],
    rows=[
        ['Email',    'Required, valid email format',   '"Please enter a valid email address"'],
        ['Password', 'Required, minimum 8 characters', '"Password must be at least 8 characters"'],
        ['Form',     'Credentials match backend',       '"Please check your credentials" (401 error)'],
    ],
    caption='Table 8.4: Login Form Validation Rules',
)

add_heading('8.5 Company Registration Page', level=2)
add_para(
    'New organisations register via a two-section progressive form. Section 1 collects '
    'company details; Section 2 creates the first administrator account. A password '
    'strength indicator and real-time validation guide users before submission.'
)
add_figure('Figure 8.8', 'Company Registration - Section 1: Company Details form')
add_figure('Figure 8.9', 'Company Registration - Section 2: Administrator Account setup')
add_figure('Figure 8.10', 'Password Strength Indicator - Weak (red) / Medium (yellow) / Strong (green)')
add_figure('Figure 8.11', 'Email Verification Success Page - after clicking the verification link')
add_table(
    headers=['Section', 'Field', 'Type', 'Validation'],
    rows=[
        ['Company Info',  'Company Name', 'Text',     'Required, 3-100 characters'],
        ['Company Info',  'Industry',     'Dropdown', 'Required - Saudi-relevant sectors'],
        ['Company Info',  'Company Size', 'Dropdown', 'Required - employee count range'],
        ['Company Info',  'Phone',        'Tel',      'Required, Saudi format (+966)'],
        ['Admin Details', 'First Name',   'Text',     'Required, 2-50 characters'],
        ['Admin Details', 'Last Name',    'Text',     'Required, 2-50 characters'],
        ['Admin Details', 'Email',        'Email',    'Required, must be unique'],
        ['Admin Details', 'Password',     'Password', 'Minimum 8 characters, complexity check'],
    ],
    caption='Table 8.5: Registration Form Fields',
)

add_heading('8.6 Protected Route System', level=2)
add_para(
    'A ProtectedRoute wrapper component reads the JWT and user role from the auth '
    'context. If the role does not satisfy the route requirement, the user is '
    'redirected to /login with the original destination URL saved for post-login '
    'redirect.'
)
add_figure('Figure 8.12', 'Route Guard Flow Diagram - authentication and role-check decision flowchart')
add_table(
    headers=['Route Path', 'Required Role', 'Redirect If Unauthorised'],
    rows=[
        ['/login',            'None (Public)',  'Dashboard (if already authenticated)'],
        ['/register/company', 'None (Public)',  'Dashboard (if already authenticated)'],
        ['/training',         'None (Public)',  'N/A - public access'],
        ['/employee/*',       'EMPLOYEE',       '/login with returnUrl preserved'],
        ['/admin/*',          'COMPANY_ADMIN',  '/login with returnUrl preserved'],
        ['/super-admin/*',    'SUPER_ADMIN',    '/login with returnUrl preserved'],
    ],
    caption='Table 8.6: Route Protection Matrix',
)

# =============================================================================
#  9. WEEK 10 – ADMIN PORTAL & CAMPAIGN MANAGEMENT
# =============================================================================
doc.add_page_break()
add_heading('9. Week 10: Admin Portal & Campaign Management', level=1)

add_para('Week 10 built the complete administrator portal covering campaign management, '
         'template library, analytics, employee oversight, and AI-powered email generation.')
add_deliverables_table(
    rows=[
        ['Build admin main layout',         'Frontend', 'Complete', 'Sidebar, header, role indicator, responsive shell'],
        ['Create campaign management UI',   'Frontend', 'Complete', 'Create, send, track simulation campaigns'],
        ['Build template library UI',       'Frontend', 'Complete', 'Upload, preview, manage phishing templates'],
        ['Create admin analytics dashboard','Frontend', 'Complete', 'KPI cards, trend charts, department comparison'],
        ['Build employee management UI',    'Frontend', 'Complete', 'Employee list, invite, risk score, detail view'],
        ['Implement AI email generation UI','Frontend', 'Complete', 'AI-powered phishing template generation'],
    ],
    caption='Table 9.1: Week 10 Frontend Deliverables',
)

add_heading('9.1 Admin Main Layout', level=2)
add_para(
    'The admin portal uses a fixed 240px sidebar that collapses to a 64px icon-only '
    'strip. The header provides search, a notification bell with unread badge, '
    'language toggle, and a user dropdown showing the "Admin" role badge.'
)
add_figure('Figure 9.1', 'Admin Dashboard Layout - full desktop view with sidebar, header, and content area')
add_figure('Figure 9.2', 'Admin Sidebar Navigation - expanded menu with icons, labels, and active highlight')
add_figure('Figure 9.3', 'Admin Header Components - search bar, notification bell, language toggle, profile dropdown')
add_table(
    headers=['Section', 'Icon', 'Route'],
    rows=[
        ['Dashboard',   'Chart icon',    '/admin/dashboard'],
        ['Campaigns',   'Email icon',    '/admin/campaigns'],
        ['Simulations', 'Target icon',   '/admin/simulations'],
        ['Employees',   'Users icon',    '/admin/employees/*'],
        ['Training',    'Book icon',     '/admin/training/*'],
        ['Analytics',   'Trending icon', '/admin/analytics/*'],
        ['Profile',     'User icon',     '/admin/profile'],
    ],
    caption='Table 9.2: Admin Navigation Structure',
)

add_heading('9.2 Campaign Management Interface', level=2)
add_figure('Figure 9.4', 'Campaigns List View - sortable table with status badges, metrics, and actions')
add_table(
    headers=['Status', 'Badge Color', 'Description'],
    rows=[
        ['Draft',     'Gray',   'Campaign created but not yet sent'],
        ['Scheduled', 'Blue',   'Set to send at a future date/time'],
        ['Active',    'Yellow', 'Currently running, tracking interactions'],
        ['Completed', 'Green',  'Campaign window closed, results available'],
        ['Paused',    'Orange', 'Temporarily halted by admin'],
    ],
    caption='Table 9.3: Campaign Status Indicators',
)
add_figure('Figure 9.5', 'Create Campaign Modal - name, description, template selection, recipients, schedule')
add_figure('Figure 9.6', 'Template Selection Grid - difficulty badges, thumbnails, and preview option')
add_figure('Figure 9.7', 'Recipient Selection - employee checkboxes with department filter and search')
add_figure('Figure 9.8', 'Campaign Send Confirmation - summary dialog with recipient count and warning')

add_heading('9.3 Template Library Interface', level=2)
add_figure('Figure 9.9', 'Template Library Grid - cards with category tags, difficulty badges, and actions')
add_table(
    headers=['Category', 'Example Templates'],
    rows=[
        ['Finance',    'Fake invoice approval, payment verification, salary update'],
        ['HR',         'Benefits update, policy change, performance review alert'],
        ['IT Support', 'Password expiry reset, system upgrade, VPN access required'],
        ['Delivery',   'Package failed delivery, shipment tracking, customs clearance'],
        ['Government', 'Nafath verification, Absher account update, GOSI notification'],
    ],
    caption='Table 9.4: Template Categories',
)
add_figure('Figure 9.10', 'Upload Template Modal - name, category, difficulty, subject, HTML body editor')
add_figure('Figure 9.11', 'Template Preview - full email rendering as seen by the recipient')

add_heading('9.4 Analytics Dashboard', level=2)
add_figure('Figure 9.12', 'Admin Analytics Overview - KPI cards, trend charts, and risk distribution')
add_figure('Figure 9.13', 'KPI Summary Cards - campaigns sent, click rate, high-risk employees with trends')
add_table(
    headers=['Metric', 'Visualisation', 'Insight Provided'],
    rows=[
        ['Click Rates over Time',   'Line chart (6 months)',       'Trend of employee susceptibility'],
        ['Department Comparison',   'Horizontal bar chart',        'Which departments need most attention'],
        ['Risk Distribution',       'Pie chart',                   'Low / Medium / High / Critical split'],
        ['Employee Performance',    'Sortable data table',         'Individual scores for targeted follow-up'],
        ['Campaign Results',        'Bar chart per campaign',      'Side-by-side sent vs clicked vs reported'],
    ],
    caption='Table 9.5: Analytics Metrics Displayed',
)
add_figure('Figure 9.14', 'Campaign Click Rates Trend - declining line chart showing improving awareness')
add_figure('Figure 9.15', 'Department Comparison Chart - colour-graded horizontal bars by department')
add_figure('Figure 9.16', 'Risk Distribution Pie Chart - Low 62%, Medium 24%, High 14% segments')

add_heading('9.5 Employee Management Interface', level=2)
add_figure('Figure 9.17', 'Employee List View - sortable table with risk badges and action buttons')
add_table(
    headers=['Column', 'Sorting', 'Filter'],
    rows=[
        ['Name',        'A-Z',        'Search box'],
        ['Email',       'A-Z',        'Search box'],
        ['Department',  'A-Z',        'Dropdown filter'],
        ['Risk Score',  'Low to High','Dropdown filter'],
        ['Campaigns',   '0 to 99',    'None'],
        ['Last Active', 'Recent first','Date range picker'],
    ],
    caption='Table 9.6: Employee Table Columns',
)
add_figure('Figure 9.18', 'Invite Employee Modal - email, first/last name, and department fields')
add_figure('Figure 9.19', 'Employee Detail View - tabbed profile with training history and quiz results')
add_figure('Figure 9.20', 'Risk Score Breakdown - contributing factors with personalised recommendations')

add_heading('9.6 AI Email Generation Feature', level=2)
add_para(
    'Administrators can generate realistic phishing templates using an integrated AI '
    'engine. The generated email includes a realism score and automatically detected '
    'red flags, allowing fine-tuning before saving to the template library.'
)
add_figure('Figure 9.21', 'AI Generation Parameters - category, difficulty, tone, and keyword fields')
add_figure('Figure 9.22', 'AI-Generated Email Preview - realism score, red flags list, and action buttons')
add_table(
    headers=['Parameter', 'Options', 'Effect on Output'],
    rows=[
        ['Category',  'Finance, HR, IT, Delivery, Government', 'Determines email topic and sender type'],
        ['Difficulty','Easy, Medium, Hard',                     'Controls how obvious vs subtle red flags are'],
        ['Tone',      'Urgent, Friendly, Official',             'Adjusts language formality and pressure tactics'],
        ['Keywords',  'Optional custom terms',                  'Includes specific phrases e.g. "Nafath", "GOSI"'],
    ],
    caption='Table 9.7: AI Generation Parameters',
)

# =============================================================================
#  10. WEEK 11 – EMPLOYEE PORTAL, PUBLIC PORTAL & GAMIFICATION
# =============================================================================
doc.add_page_break()
add_heading('10. Week 11: Employee Portal, Public Portal & Gamification', level=1)

add_para('Week 11 completed the platform with all employee-facing screens, the public '
         'awareness portal, and the gamification display layer.')
add_deliverables_table(
    rows=[
        ['Build employee dashboard',        'Frontend', 'Complete', 'Risk gauge, training progress, badges, leaderboard rank'],
        ['Create quiz interface',           'Frontend', 'Complete', 'Email identification quiz with instant feedback'],
        ['Build training module UI',        'Frontend', 'Complete', 'Video player + interactive learning + post-quiz'],
        ['Implement gamification display',  'Frontend', 'Complete', 'Badges showcase, leaderboard, badge unlock animation'],
        ['Build public community portal',   'Frontend', 'Complete', 'No-login homepage with hero, features, statistics'],
        ['Create awareness content pages',  'Frontend', 'Complete', 'Phishing, vishing, smishing guides (AR/EN)'],
        ['Build public quiz interface',     'Frontend', 'Complete', 'Anonymous knowledge test with results and CTA'],
        ['Final UI/UX polish',              'Frontend', 'Complete', 'Consistent styling, responsive verification, RTL check'],
    ],
    caption='Table 10.1: Week 11 Frontend Deliverables',
)

add_heading('10.1 Employee Dashboard', level=2)
add_figure('Figure 10.1', 'Employee Dashboard Overview - risk gauge, training progress, and badge section')
add_figure('Figure 10.2', 'Risk Score Widget - circular gauge with Low/Medium/High colour gradient')
add_figure('Figure 10.3', 'Assigned Training Cards - due dates, status badges, and Start/Continue buttons')
add_figure('Figure 10.4', 'Employee Dashboard - Mobile View (375px) with vertically stacked layout')
add_table(
    headers=['Widget', 'Content', 'Action Available'],
    rows=[
        ['Risk Score',         'Score /100, colour level, trend arrow',    'View score breakdown'],
        ['Training Progress',  'X/Y modules completed, progress bar',      'View all assigned training'],
        ['Upcoming Deadlines', 'Next 3 due training with dates',            'Start training directly'],
        ['Recent Badges',      'Last 3 earned badge icons with names',      'View all achievements'],
        ['Leaderboard Rank',   'Current position and +/- change this week', 'View full company leaderboard'],
    ],
    caption='Table 10.2: Employee Dashboard Widgets',
)

add_heading('10.2 Quiz Interface', level=2)
add_para(
    'The phishing identification quiz presents realistic email mockups one at a time. '
    'Large touch-friendly buttons (Safe / Phishing) make it usable on mobile. '
    'Instant feedback after each question converts wrong answers into learning moments.'
)
add_figure('Figure 10.5', 'Quiz Start Screen - title, description, difficulty badge, and Start Quiz button')
add_figure('Figure 10.6', 'Quiz Question - email mockup with Safe (green) and Phishing (red) answer buttons')
add_figure('Figure 10.7', 'Quiz Feedback - Correct Answer (green banner with red-flag explanation)')
add_figure('Figure 10.8', 'Quiz Feedback - Incorrect Answer (red banner with detailed missed red flags)')
add_figure('Figure 10.9', 'Quiz Results Screen - score, category breakdown, and improvement recommendations')
add_table(
    headers=['Element', 'Correct Answer', 'Incorrect Answer'],
    rows=[
        ['Banner Colour', 'Green (#10B981)',                  'Red (#EF4444)'],
        ['Icon',          'Checkmark',                        'Cross mark'],
        ['Message',       '"Correct! This is phishing/safe"', '"Incorrect. This was phishing/safe"'],
        ['Explanation',   'Brief confirmation of red flags',  'Detailed explanation of every missed red flag'],
        ['Learning Tip',  'Optional bonus info',              'Always provided to reinforce correct identification'],
    ],
    caption='Table 10.3: Quiz Feedback Components',
)

add_heading('10.3 Training Module Interface', level=2)
add_para(
    'Two delivery methods are offered: a video player with post-video comprehension quiz, '
    'or an interactive lesson built around Saudi-specific scenarios (fake Nafath SMS, '
    'spoofed Absher emails). Both require a minimum 80% score to mark the module complete.'
)
add_figure('Figure 10.10', 'Training Module Page - video player, learning objectives, and Start button')
add_figure('Figure 10.11', 'Training Mode Selection - Video Training vs Interactive Learning cards')
add_figure('Figure 10.12', 'Video Player - controls, closed captions in AR/EN, auto progress tracking')
add_figure('Figure 10.13', 'Interactive Learning - Nafath SMS fake vs real comparison with click-to-reveal')
add_figure('Figure 10.14', 'Post-Training Quiz - multiple choice questions unlocked after watching video')
add_figure('Figure 10.15', 'Training Completion Certificate - downloadable PDF with score and date')
add_table(
    headers=['Component', 'Function', 'Completion Criteria'],
    rows=[
        ['Video Player',       'Primary content delivery',      'Must watch at least 90% of video'],
        ['Interactive Lesson', 'Alternative delivery method',   'Complete all scenes and scene quiz'],
        ['Comprehension Quiz', 'Knowledge assessment',          'Score 80% or above (4/5 correct minimum)'],
        ['Certificate',        'Completion proof for records',  'Auto-generated as PDF on passing quiz'],
        ['Progress Tracker',   'Resume capability',             'Position saved every 30 seconds automatically'],
    ],
    caption='Table 10.4: Training Module Components',
)

add_heading('10.4 Gamification Features', level=2)
add_figure('Figure 10.16', 'Badges Showcase - earned (colour) and locked (greyscale) badges in grid layout')
add_figure('Figure 10.17', 'Leaderboard View - ranked table with medal icons for top 3 and department column')
add_figure('Figure 10.18', 'Badge Unlock Animation - shine effect, points earned, and confetti celebration')
add_figure('Figure 10.19', 'Leaderboard Top 3 Podium - gold/silver/bronze with animated confetti for first')
add_table(
    headers=['Badge', 'Requirement', 'Points Reward'],
    rows=[
        ['First Steps',         'Complete first training module',           '50 pts'],
        ['Quick Learner',       'Finish any training in under 10 minutes',  '25 pts'],
        ['Perfect Score',       'Achieve 100% on any quiz',                 '50 pts'],
        ['Phish Detector',      'Correctly identify 10 phishing emails',    '100 pts'],
        ['Reporter',            'Report your first real phishing attempt',  '75 pts'],
        ['Streak Master',       'Complete training 5 consecutive weeks',    '150 pts'],
        ['Department Champion', 'Highest score in your department',         '200 pts'],
    ],
    caption='Table 10.5: Available Badges',
)

add_heading('10.5 Public Portal Homepage', level=2)
add_figure('Figure 10.20', 'Public Homepage Hero - headline, CTA button, and AR/EN language toggle')
add_figure('Figure 10.21', 'Public Homepage Features - Free Training, Saudi Context, Bilingual cards')
add_figure('Figure 10.22', 'Public Homepage Training Topics - Phishing, Vishing, Smishing category cards')
add_figure('Figure 10.23', 'Public Homepage Statistics - user count, company count, satisfaction rate')
add_table(
    headers=['Section', 'Content', 'Purpose'],
    rows=[
        ['Hero',         'Main headline, CTA button, background illustration', 'Capture attention, drive to training'],
        ['Features',     '3 key benefits with icons and descriptions',         'Communicate value proposition'],
        ['Topics',       'Phishing, Vishing, Smishing cards with counts',      'Navigate directly to content'],
        ['Statistics',   'User count, company count, satisfaction rate',       'Build trust and social proof'],
        ['How It Works', '3-step process: Choose, Learn, Test',               'Explain user journey simply'],
        ['Footer',       'Links, contact info, social media icons',            'Navigation and brand presence'],
    ],
    caption='Table 10.6: Public Homepage Sections',
)

add_heading('10.6 Awareness Content Pages', level=2)
add_figure('Figure 10.24', 'Phishing Awareness Page - English with Nafath and Absher examples annotated')
add_figure('Figure 10.25', 'Phishing Awareness Page - Arabic RTL with identical structure mirrored')
add_figure('Figure 10.26', 'Interactive Nafath SMS Comparison - fake vs real message click-to-reveal')
add_figure('Figure 10.27', 'Vishing Awareness Page - audio scam call player with transcript')
add_figure('Figure 10.28', 'Smishing Awareness Page - annotated SMS screenshots with red flag callouts')
add_table(
    headers=['Page', 'Interactive Elements', 'Saudi-Specific Examples'],
    rows=[
        ['Phishing', 'Click-to-reveal red flags, embedded quiz',    'Nafath SMS, Absher email, bank notifications'],
        ['Vishing',  'Audio scam call player, Spot Red Flags game', 'Fake ministry calls, lottery scams, family emergencies'],
        ['Smishing', 'Screenshot comparison tool, embedded quiz',   'Package delivery SMS, SADAD payment alerts'],
    ],
    caption='Table 10.7: Awareness Content Structure',
)

add_heading('10.7 Public Quiz Interface', level=2)
add_figure('Figure 10.29', 'Public Quiz Landing - no login required, difficulty selection screen')
add_figure('Figure 10.30', 'Public Quiz Question - multiple choice with progress bar')
add_figure('Figure 10.31', 'Public Quiz Results - score breakdown by category and registration CTA')
add_table(
    headers=['Feature', 'Public Users (No Login)', 'Registered Employees'],
    rows=[
        ['Authentication',    'Not required',              'JWT token required'],
        ['Score Persistence', 'Session only, not saved',   'Saved permanently to profile'],
        ['Recommendations',   'Generic suggestions',       'Personalised based on risk profile'],
        ['Gamification',      'None',                      'Badges, leaderboard points, achievement'],
        ['Certificate',       'No',                        'Yes - downloadable PDF on completion'],
        ['Progress Tracking', 'No history',                'Full attempt history and improvement chart'],
    ],
    caption='Table 10.8: Public vs Employee Quiz Experience',
)

# =============================================================================
#  11. RESPONSIVE DESIGN & ACCESSIBILITY
# =============================================================================
doc.add_page_break()
add_heading('11. Responsive Design & Accessibility', level=1)

add_para(
    'All interfaces were designed and tested across four breakpoints, ensuring a '
    'consistent experience from mobile phones to large desktop monitors. Accessibility '
    'compliance follows WCAG 2.1 AA standards throughout.'
)
add_table(
    headers=['Breakpoint', 'Screen Width', 'Target Devices', 'Layout Changes'],
    rows=[
        ['Mobile (sm)',  '< 640px',     'Smartphones',           'Single column, hamburger menu, stacked cards'],
        ['Tablet (md)',  '640-1024px',  'Tablets, landscape',    '2-column grid, collapsible sidebar'],
        ['Desktop (lg)', '1024-1280px', 'Laptops',               '3-column grid, full sidebar visible'],
        ['Wide (xl)',    '> 1280px',    'Large monitors',         '4-column grid, expanded dashboard widgets'],
    ],
    caption='Table 11.1: Responsive Breakpoints',
)
add_figure('Figure 11.1', 'Responsive Layout Comparison - same admin dashboard at mobile, tablet, desktop')
add_table(
    headers=['Accessibility Feature', 'Implementation', 'Standard Met'],
    rows=[
        ['Colour Contrast',  'Minimum 4.5:1 ratio for all normal text',       'WCAG 2.1 AA'],
        ['Keyboard Nav',     'All interactive elements fully tab-accessible',  'WCAG 2.1 A'],
        ['ARIA Labels',      'All icon-only buttons have aria-label',          'WCAG 2.1 A'],
        ['Focus Indicators', 'Visible blue focus ring on all focusable items', 'WCAG 2.1 AA'],
        ['Alt Text',         'All images have descriptive alt text',           'WCAG 2.1 A'],
        ['Error Messages',   'Icon + text communication (not colour only)',    'WCAG 2.1 AA'],
        ['RTL Support',      'Complete mirroring for Arabic language',         'W3C i18n'],
    ],
    caption='Table 11.2: Accessibility Features',
)

# =============================================================================
#  SECTION DIVIDER – TESTING
# =============================================================================
doc.add_page_break()
div3 = doc.add_paragraph()
div3.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr3 = div3.add_run('TESTING & CONCLUSION')
dr3.font.size = Pt(18); dr3.font.bold = True; dr3.font.color.rgb = DARK_BLUE
doc.add_paragraph()

# =============================================================================
#  12. BACKEND TESTING DOCUMENTATION
# =============================================================================
add_heading('12. Backend Testing Documentation', level=1)

add_heading('12.1 Authentication & Authorization', level=2)
add_para(
    'All authentication endpoints were tested for: successful registration and login, '
    'token refresh and expiry, email verification flow, password change, and rejection '
    'of malformed or expired tokens. Role enforcement was verified by attempting '
    'cross-role access on every protected endpoint.'
)

add_heading('12.2 Company Management', level=2)
add_para(
    'Company CRUD operations were tested including multi-tenant isolation (ensuring '
    'Company A admins cannot access Company B data), company size and industry '
    'validation, and super-admin cross-company visibility.'
)

add_heading('12.3 Employee Management', level=2)
add_para(
    'The invitation system was tested end-to-end: invite email delivery, 7-day token '
    'expiry, duplicate invitation prevention, accept with profile completion, and '
    'status transitions (PENDING -> ACCEPTED / EXPIRED).'
)

add_heading('12.4 Training Module System', level=2)
add_para(
    'Training assignment, completion, quiz submission, and auto-remediation triggers '
    'were tested. Edge cases covered: force_assign with pending warnings, bulk '
    'assignment to empty departments, and quiz retry after failure.'
)

add_heading('12.5 Phishing Simulation System', level=2)
add_para(
    'Simulation send, tracking pixel fire, lure link click, "caught" redirect, and '
    'TrackingEvent creation were all validated. Concurrent click events on the same '
    'token were tested to confirm idempotency. Stats denormalisation was verified '
    'after each event type.'
)

add_heading('12.6 Security Testing', level=2)
add_para('Security testing covered OWASP Top 10 vulnerabilities relevant to the platform:')
for vuln, result in [
    ('SQL Injection via API parameters', 'PASS - DRF ORM parameterises all queries'),
    ('JWT token tampering',              'PASS - Signature validation rejects modified tokens'),
    ('IDOR (Insecure Direct Object Reference)', 'PASS - All querysets filtered by company FK'),
    ('XSS in template content',         'PASS - HTML sanitised before storage and render'),
    ('Brute-force login',               'PASS - Rate limiting via django-ratelimit on auth endpoints'),
    ('Privilege escalation',             'PASS - Role checked on every ViewSet action'),
]:
    add_bullet(f'{result}', bold_prefix=vuln)

add_heading('12.7 Results Summary', level=2)
add_table(
    headers=['Metric', 'Value'],
    rows=[
        ['Total API endpoints tested',    '104'],
        ['Total test cases executed',     '234'],
        ['Test cases passed',             '217'],
        ['Defects found',                 '17'],
        ['Defects resolved',              '17 (100%)'],
        ['Critical defects',              '0 remaining'],
        ['Test coverage (endpoints)',     '100%'],
        ['Automated tests (pytest)',      '89 unit + integration tests'],
    ],
    caption='Table 12.1: Backend Testing Results Summary',
)

# =============================================================================
#  13. CONCLUSION
# =============================================================================
doc.add_page_break()
add_heading('13. Conclusion', level=1)

add_para(
    'PhishAware successfully delivers a production-ready, bilingual cybersecurity '
    'awareness platform tailored for Saudi Arabian and MENA region organisations. '
    'The system integrates a robust Django REST backend with a modern React frontend '
    'to provide a seamless end-to-end security training experience.'
)

add_heading('13.1 Key Achievements', level=2)
for item in [
    'Complete RESTful API with 104 endpoints covering all platform features',
    'Secure JWT authentication with four-tier RBAC enforced on every endpoint',
    'Automated phishing simulation pipeline with per-employee tracking and analytics',
    'Signal-driven risk scoring engine updating in real time on every user interaction',
    'Gamification layer (badges, points, leaderboard) proven to drive engagement',
    'Fully bilingual interface (Arabic/English) with automatic RTL layout transformation',
    'Culturally relevant content using Nafath, Absher, and Saudi-specific examples',
    'Responsive design verified from 375px mobile to 1440px wide desktop',
    'WCAG 2.1 AA accessibility compliance throughout all three portals',
    '234 test cases executed with 100% defect resolution before submission',
]:
    add_bullet(item)

add_heading('13.2 Future Enhancements', level=2)
for item in [
    'Mobile application (React Native) for on-the-go training and notifications',
    'Advanced AI red-flag scoring with GPT integration for dynamic template generation',
    'Integration with Saudi national identity systems (Nafath API) for employee verification',
    'Expanded content library: deepfake video detection, QR code phishing, social media threats',
    'Organisation-level benchmarking and industry peer comparison analytics',
]:
    add_bullet(item)

# =============================================================================
#  APPENDIX – SCREENSHOT INDEX
# =============================================================================
doc.add_page_break()
add_heading('Appendix A: Complete Screenshot Index', level=1)
add_para('All figure placeholders use Word\'s Caption style. After inserting screenshots, '
         'right-click the Table of Figures and select "Update Field" to populate page numbers.',
         italic=True, color=GRAY_700)

all_figures = [
    ('Figure 7.1',  'PhishAware Color Palette',                         'Section 7.2'),
    ('Figure 7.2',  'Typography Hierarchy - EN and AR comparison',      'Section 7.3'),
    ('Figure 7.3',  'PhishAware Logo System - all variants',            'Section 7.4'),
    ('Figure 8.1',  'Project Folder Structure',                         'Section 8.1'),
    ('Figure 8.2',  'Language Switcher Component',                      'Section 8.2'),
    ('Figure 8.3',  'RTL Layout Transformation',                        'Section 8.2'),
    ('Figure 8.4',  'Axios API Configuration Code',                     'Section 8.3'),
    ('Figure 8.5',  'Login Page - English Version',                     'Section 8.4'),
    ('Figure 8.6',  'Login Page - Arabic Version (RTL)',                'Section 8.4'),
    ('Figure 8.7',  'Login Error States',                               'Section 8.4'),
    ('Figure 8.8',  'Company Registration - Section 1',                 'Section 8.5'),
    ('Figure 8.9',  'Company Registration - Section 2',                 'Section 8.5'),
    ('Figure 8.10', 'Password Strength Indicator',                      'Section 8.5'),
    ('Figure 8.11', 'Email Verification Success Page',                  'Section 8.5'),
    ('Figure 8.12', 'Route Guard Flow Diagram',                         'Section 8.6'),
    ('Figure 9.1',  'Admin Dashboard Layout - Desktop',                 'Section 9.1'),
    ('Figure 9.2',  'Admin Sidebar Navigation',                         'Section 9.1'),
    ('Figure 9.3',  'Admin Header Components',                          'Section 9.1'),
    ('Figure 9.4',  'Campaigns List View',                              'Section 9.2'),
    ('Figure 9.5',  'Create Campaign Modal',                            'Section 9.2'),
    ('Figure 9.6',  'Template Selection Grid',                          'Section 9.2'),
    ('Figure 9.7',  'Recipient Selection Interface',                    'Section 9.2'),
    ('Figure 9.8',  'Campaign Send Confirmation',                       'Section 9.2'),
    ('Figure 9.9',  'Template Library Grid View',                       'Section 9.3'),
    ('Figure 9.10', 'Upload Template Modal',                            'Section 9.3'),
    ('Figure 9.11', 'Template Preview Modal',                           'Section 9.3'),
    ('Figure 9.12', 'Admin Analytics Dashboard Overview',               'Section 9.4'),
    ('Figure 9.13', 'KPI Summary Cards',                                'Section 9.4'),
    ('Figure 9.14', 'Campaign Click Rates Trend Chart',                 'Section 9.4'),
    ('Figure 9.15', 'Department Comparison Chart',                      'Section 9.4'),
    ('Figure 9.16', 'Risk Distribution Pie Chart',                      'Section 9.4'),
    ('Figure 9.17', 'Employee List View',                               'Section 9.5'),
    ('Figure 9.18', 'Invite Employee Modal',                            'Section 9.5'),
    ('Figure 9.19', 'Employee Detail View',                             'Section 9.5'),
    ('Figure 9.20', 'Risk Score Breakdown',                             'Section 9.5'),
    ('Figure 9.21', 'AI Generation Parameters',                         'Section 9.6'),
    ('Figure 9.22', 'AI-Generated Email Preview',                       'Section 9.6'),
    ('Figure 10.1', 'Employee Dashboard Overview',                      'Section 10.1'),
    ('Figure 10.2', 'Risk Score Gauge Widget',                          'Section 10.1'),
    ('Figure 10.3', 'Assigned Training Cards',                          'Section 10.1'),
    ('Figure 10.4', 'Employee Dashboard - Mobile View',                 'Section 10.1'),
    ('Figure 10.5', 'Quiz Start Screen',                                'Section 10.2'),
    ('Figure 10.6', 'Quiz Question - Email Display',                    'Section 10.2'),
    ('Figure 10.7', 'Quiz Feedback - Correct Answer',                   'Section 10.2'),
    ('Figure 10.8', 'Quiz Feedback - Incorrect Answer',                 'Section 10.2'),
    ('Figure 10.9', 'Quiz Results Screen',                              'Section 10.2'),
    ('Figure 10.10','Training Module Page',                             'Section 10.3'),
    ('Figure 10.11','Training Mode Selection',                          'Section 10.3'),
    ('Figure 10.12','Video Player Interface',                           'Section 10.3'),
    ('Figure 10.13','Interactive Learning - Nafath Example',            'Section 10.3'),
    ('Figure 10.14','Post-Training Quiz',                               'Section 10.3'),
    ('Figure 10.15','Training Completion Certificate',                  'Section 10.3'),
    ('Figure 10.16','Badges Showcase',                                  'Section 10.4'),
    ('Figure 10.17','Leaderboard View',                                 'Section 10.4'),
    ('Figure 10.18','Badge Unlock Animation',                           'Section 10.4'),
    ('Figure 10.19','Leaderboard Top 3 Podium',                        'Section 10.4'),
    ('Figure 10.20','Public Homepage Hero Section',                     'Section 10.5'),
    ('Figure 10.21','Public Homepage Features Section',                 'Section 10.5'),
    ('Figure 10.22','Public Homepage Training Topics',                  'Section 10.5'),
    ('Figure 10.23','Public Homepage Statistics Banner',                'Section 10.5'),
    ('Figure 10.24','Phishing Awareness Page - English',                'Section 10.6'),
    ('Figure 10.25','Phishing Awareness Page - Arabic RTL',             'Section 10.6'),
    ('Figure 10.26','Interactive Nafath SMS Comparison',                'Section 10.6'),
    ('Figure 10.27','Vishing Awareness Page',                           'Section 10.6'),
    ('Figure 10.28','Smishing Awareness Page',                         'Section 10.6'),
    ('Figure 10.29','Public Quiz Landing Page',                         'Section 10.7'),
    ('Figure 10.30','Public Quiz Question',                             'Section 10.7'),
    ('Figure 10.31','Public Quiz Results Screen',                       'Section 10.7'),
    ('Figure 11.1', 'Responsive Layout Comparison',                     'Section 11'),
]
add_table(
    headers=['Figure #', 'Description', 'Document Section'],
    rows=all_figures,
    caption='Table A.1: Complete Screenshot Index (69 figures - Chapters 7 to 11)',
)

# =============================================================================
#  SAVE
# =============================================================================
doc.save(OUTPUT_DOC)
print(f'Saved: {OUTPUT_DOC}')
print(f'Chapters : 1 Introduction + Weeks 1-5 Backend + Weeks 9-11 Frontend + Testing + Conclusion')
print(f'Figures  : 69 placeholders (Figure 7.1 to Figure 11.1)')
print(f'Tables   : 51 (Table 1.1 to Table A.1)')
print(f'ToC      : static, 13 chapters')
print(f'ToF      : Word auto-update field (Caption style)')
print(f'ToT      : static list of all 51 tables')
