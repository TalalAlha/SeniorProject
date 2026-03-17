"""
Generate PhishAware IEEE Complete Report as DOCX
"""
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ── paths ──
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS_DIR = os.path.join(SCRIPT_DIR, "screenshots")
OUTPUT_PATH = os.path.join(SCRIPT_DIR, "PhishAware_IEEE_Complete_Report.docx")

# ── constants ──
FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(10)
CAPTION_SIZE = Pt(9)
HEADING1_SIZE = Pt(12)
HEADING2_SIZE = Pt(11)
HEADING3_SIZE = Pt(10)
CODE_FONT = "Courier New"
CODE_SIZE = Pt(8)

# ── figure mapping: Fig N → filename ──
FIGURE_MAP = {
    1: "Figure_7.1.png", 2: "Figure_7.3.png",
    3: "Figure_8.1.png", 4: "Figure_8.2.png",
    5: "Figure_8.3a.png", 6: "Figure_8.3b.png",
    7: "Figure_8.4.png", 8: "Figure_8.5.png",
    9: "Figure_8.6.png", 10: "Figure_8.7.png",
    11: "Figure_8.8.png", 12: "Figure_8.9.png",
    13: "Figure_9.1.png", 14: "Figure_9.2.png",
    15: "Figure_9.3.png", 16: "Figure_9.4.png",
    17: "Figure_9.5.png", 18: "Figure_9.6.png",
    19: "Figure_9.7.png", 20: "Figure_9.8.png",
    21: "Figure_9.9.png", 22: "Figure_9.10.png",
    23: "Figure_9.11.png", 24: "Figure_9.12.png",
    25: "Figure_9.13.png", 26: "Figure_9.14.png",
    27: "Figure_9.15.png", 28: "Figure_9.16.png",
    29: "Figure_9.17.png", 30: "Figure_9.18.png",
    31: "Figure_9.19.png", 32: "Figure_9.20.png",
    33: "Figure_10.1.png", 34: "Figure_10.2.png",
    35: "Figure_10.3.png", 36: "Figure_10.4.png",
    37: "Figure_10.5.png", 38: "Figure_10.6.png",
    39: "Figure_10.7.png", 40: "Figure_10.8.png",
    41: "Figure_10.9.png", 42: "Figure_10.10.png",
    43: "Figure_10.11.png", 44: "Figure_10.12.png",
    45: "Figure_10.13.png", 46: "Figure_10.14.png",
    47: "Figure_10.15.png", 48: "Figure_10.16.png",
    49: "Figure_10.17.png", 50: "Figure_10.18.png",
}

FIGURE_CAPTIONS = {
    1: "Fig. 1. Color palette showing 10 color swatches with hex codes and usage labels for PhishAware platform design system.",
    2: "Fig. 2. PhishAware logo featuring shield icon with wordmark displayed in the navigation bar.",
    3: "Fig. 3. VS Code file explorer showing Frontend/src/ folder structure with organized component directories.",
    4: "Fig. 4. AR/EN language switcher toggle button in the navigation bar for bilingual interface control.",
    5: "Fig. 5. Application page in English demonstrating left-to-right (LTR) layout orientation.",
    6: "Fig. 6. Same application page in Arabic demonstrating right-to-left (RTL) mirrored layout.",
    7: "Fig. 7. Login page in English showing PhishAware logo, email input, password field, and Sign In button.",
    8: "Fig. 8. Login page in Arabic displaying full RTL transformation with mirrored layout elements.",
    9: "Fig. 9. Login page showing validation errors with red borders and descriptive error messages below fields.",
    10: "Fig. 10. Company registration form displaying Company Name, Website URL, and admin account creation fields.",
    11: "Fig. 11. Password field validation showing browser message \"lengthen to 8 characters\" for security requirements.",
    12: "Fig. 12. Email verification success page with green checkmark icon and Login button for account access.",
    13: "Fig. 13. Admin portal full layout showing collapsible sidebar, header with notifications, and main dashboard content area.",
    14: "Fig. 14. Admin sidebar in expanded state displaying all menu items with icons, active item highlighted in blue.",
    15: "Fig. 15. Admin header components including notification bell with badge count, AR/EN language toggle, and profile dropdown menu.",
    16: "Fig. 16. Campaigns management page showing one campaign entry in Draft status with action buttons.",
    17: "Fig. 17. Campaign creation interface with AR/EN language ratio slider and phishing scenario difficulty ratio controls.",
    18: "Fig. 18. Employee assignment interface with selection checkboxes, department filter dropdown, and selected employee count display.",
    19: "Fig. 19. Simulations tab displaying list of available pre-built phishing email templates.",
    20: "Fig. 20. Simulation creation Step 1 showing name input, description field, and pre-built template selection.",
    21: "Fig. 21. Simulation creation Step 2 displaying employee selection interface with department filtering options.",
    22: "Fig. 22. Simulation creation Step 3 showing campaign summary review and confirm sending interface.",
    23: "Fig. 23. Simulation creation Step 4 displaying loading screen with progress indicator while sending simulation emails.",
    24: "Fig. 24. Analytics dashboard upper section with 4 KPI cards, Average Risk Score Over Time line chart, and Risk Level Distribution pie chart (empty state shown).",
    25: "Fig. 25. Analytics dashboard lower section showing Security Training Activity counters and Training Completion Rates breakdown by training type.",
    26: "Fig. 26. Employee management table displaying name, email, department, risk badge, campaign count, and action buttons.",
    27: "Fig. 27. Invite employee modal with input fields for first name, last name, email address, and department selection.",
    28: "Fig. 28. Employee detail view showing profile information, current risk score, and tabbed sections for training history and campaign results.",
    29: "Fig. 29. Admin training page displaying 3 training modules: Email Phishing Awareness, SMS Scams Prevention, and Voice Call Scams Detection.",
    30: "Fig. 30. Training module View Details modal showing module description, learning objectives, and quiz information.",
    31: "Fig. 31. Training module assignment interface for selecting target employees with department and risk level filters.",
    32: "Fig. 32. Admin profile page displaying personal information fields and account settings options.",
    33: "Fig. 33. Employee dashboard showing personalized welcome message, risk score gauge, training progress indicators, and earned badges display.",
    34: "Fig. 34. Employee dashboard mobile view (375px width) with stacked vertical layout and hamburger menu navigation.",
    35: "Fig. 35. Quiz start screen displaying quiz title, description text, difficulty level indicator, and Start Quiz button.",
    36: "Fig. 36. Quiz question interface showing phishing email scenario with Safe and Phishing answer buttons and progress indicator dots.",
    37: "Fig. 37. Quiz results page displaying final score, performance breakdown by category, and personalized recommendations.",
    38: "Fig. 38. Training module detail page showing module title, description, learning objectives, and training mode selection.",
    39: "Fig. 39. Training mode selection interface with two cards: Video Training and Interactive Learning, each listing key features.",
    40: "Fig. 40. Badges showcase page displaying earned badges in full color and locked badges in greyscale grid layout.",
    41: "Fig. 41. Leaderboard interface showing \"Your Position\" section with Weekly/Monthly/All Time tabs, statistics cards, and \"How Points Work\" scoring formula.",
    42: "Fig. 42. Employee profile and settings page displaying personal information and account configuration options.",
    43: "Fig. 43. Public homepage hero section with headline \"Protect Yourself from Cyber Threats\", tagline, and dual CTA buttons (Start Learning, Take Free Quiz).",
    44: "Fig. 44. Public homepage \"Learn About Cyber Threats\" section featuring 3 topic cards (Email in blue, SMS in green, Voice in purple) with Saudi-specific examples.",
    45: "Fig. 45. Public homepage statistics banner displaying platform adoption metrics and success indicators.",
    46: "Fig. 46. Public training page in English showing 3 training module cards with descriptions and \"Learn More\" action buttons.",
    47: "Fig. 47. Public training page in Arabic displaying identical content with complete RTL layout and Arabic text.",
    48: "Fig. 48. Public quiz landing page \"Test Your Knowledge\" section with 3 quiz options: Email Phishing (10 questions), SMS Scams (10 questions), and Voice Scams (10 questions), all marked as Beginner difficulty.",
    49: "Fig. 49. Public quiz question interface with multiple choice options, progress bar indicator, and Submit Answer button.",
    50: "Fig. 50. Public quiz results page showing achieved score, category performance breakdown, and registration call-to-action for full platform access.",
}


# ── helper functions ──

def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    """Add borders to table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


from lxml import etree

XML_SPACE_NS = '{http://www.w3.org/XML/1998/namespace}space'


def _append_field_to_paragraph(p_elem, fld_type, dirty=False):
    """Append a <w:r><w:fldChar w:fldCharType="..."/></w:r> to paragraph element."""
    r = etree.SubElement(p_elem, qn('w:r'))
    fld = etree.SubElement(r, qn('w:fldChar'))
    fld.set(qn('w:fldCharType'), fld_type)
    if dirty:
        fld.set(qn('w:dirty'), 'true')
    return r


def _append_instr_to_paragraph(p_elem, instruction):
    """Append a <w:r><w:instrText>...</w:instrText></w:r> to paragraph element."""
    r = etree.SubElement(p_elem, qn('w:r'))
    instr = etree.SubElement(r, qn('w:instrText'))
    instr.set(XML_SPACE_NS, 'preserve')
    instr.text = f' {instruction} '
    return r


def _append_text_to_paragraph(p_elem, text):
    """Append a <w:r><w:t>text</w:t></w:r> to paragraph element."""
    r = etree.SubElement(p_elem, qn('w:r'))
    t = etree.SubElement(r, qn('w:t'))
    t.set(XML_SPACE_NS, 'preserve')
    t.text = text
    return r


def add_seq_field(paragraph, seq_name, number=0, run_font_name=FONT_NAME, run_font_size=CAPTION_SIZE):
    """Insert a SEQ field (auto-number) into a paragraph using lxml SubElement."""
    p_elem = paragraph._element

    # begin
    _append_field_to_paragraph(p_elem, 'begin')
    # instruction
    _append_instr_to_paragraph(p_elem, f'SEQ {seq_name} \\* ARABIC')
    # separate
    _append_field_to_paragraph(p_elem, 'separate')
    # pre-filled number (correct value even before Word updates)
    _append_text_to_paragraph(p_elem, str(number))
    # end
    _append_field_to_paragraph(p_elem, 'end')


def add_toc_field(doc, field_instruction, heading_text=None):
    """Insert a TOC/TOF field into the document using lxml SubElement."""
    if heading_text:
        doc.add_heading(heading_text, level=1)
    p = doc.add_paragraph()
    p_elem = p._element

    # begin
    _append_field_to_paragraph(p_elem, 'begin')
    # instruction
    _append_instr_to_paragraph(p_elem, field_instruction)
    # separate
    _append_field_to_paragraph(p_elem, 'separate')

    # Placeholder text
    run4 = p.add_run("Right-click and select 'Update Field' to populate this list.")
    run4.font.name = FONT_NAME
    run4.font.size = BODY_SIZE
    run4.font.italic = True
    run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # end
    _append_field_to_paragraph(p_elem, 'end')
    return p


class IEEEDocGenerator:
    def __init__(self):
        self.doc = Document()
        self.table_counter = 0
        self.figure_counter = 0
        self._setup_styles()
        self._setup_margins()

    def _setup_margins(self):
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = FONT_NAME
        font.size = BODY_SIZE
        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15

        # Heading 1
        h1 = self.doc.styles['Heading 1']
        h1.font.name = FONT_NAME
        h1.font.size = HEADING1_SIZE
        h1.font.bold = True
        h1.font.all_caps = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(12)
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Heading 2
        h2 = self.doc.styles['Heading 2']
        h2.font.name = FONT_NAME
        h2.font.size = HEADING2_SIZE
        h2.font.bold = True
        h2.font.italic = False
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(8)

        # Heading 3
        h3 = self.doc.styles['Heading 3']
        h3.font.name = FONT_NAME
        h3.font.size = HEADING3_SIZE
        h3.font.bold = True
        h3.font.italic = True
        h3.font.color.rgb = RGBColor(0, 0, 0)
        h3.paragraph_format.space_before = Pt(10)
        h3.paragraph_format.space_after = Pt(6)

        # Caption style (used by Word's Table of Figures / Table of Tables)
        caption_style = self.doc.styles['Caption']
        caption_style.font.name = FONT_NAME
        caption_style.font.size = CAPTION_SIZE
        caption_style.font.italic = True
        caption_style.font.color.rgb = RGBColor(0, 0, 0)
        caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_style.paragraph_format.space_before = Pt(4)
        caption_style.paragraph_format.space_after = Pt(10)

    def add_paragraph(self, text, bold=False, italic=False, size=None, align=None, space_after=None):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = size or BODY_SIZE
        run.font.bold = bold
        run.font.italic = italic
        if align:
            p.alignment = align
        if space_after is not None:
            p.paragraph_format.space_after = space_after
        return p

    def add_code_block(self, code_text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(code_text)
        run.font.name = CODE_FONT
        run.font.size = CODE_SIZE
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        # Add light gray background via shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        run._element.get_or_add_rPr().append(shading)
        return p

    def add_figure(self, fig_num, width=None):
        """Insert a figure image + caption with SEQ field for Table of Figures."""
        filename = FIGURE_MAP.get(fig_num)
        # Get the caption text after "Fig. N. "
        full_caption = FIGURE_CAPTIONS.get(fig_num, f"Fig. {fig_num}.")
        # Extract just the description part (after "Fig. N. ")
        caption_desc = full_caption.split(". ", 2)[-1] if ". " in full_caption else full_caption
        if not filename:
            return
        filepath = os.path.join(SCREENSHOTS_DIR, filename)
        if not os.path.exists(filepath):
            self.add_paragraph(f"[Image not found: {filename}]", italic=True)
            return

        # Image paragraph
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(filepath, width=width or Inches(5.5))

        # Caption paragraph using Caption style with SEQ field
        cap_p = self.doc.add_paragraph(style='Caption')
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add "Fig. " prefix
        run_prefix = cap_p.add_run("Fig. ")
        run_prefix.font.name = FONT_NAME
        run_prefix.font.size = CAPTION_SIZE
        run_prefix.font.italic = True

        # Add SEQ Figure field (auto-number)
        self.figure_counter += 1
        add_seq_field(cap_p, "Figure", number=self.figure_counter)

        # Add ". " separator and description
        run_desc = cap_p.add_run(". " + caption_desc)
        run_desc.font.name = FONT_NAME
        run_desc.font.size = CAPTION_SIZE
        run_desc.font.italic = True
        cap_p.paragraph_format.space_after = Pt(10)

    def add_table(self, headers, rows, col_widths=None, caption=None):
        """Add a table with optional caption. If caption is None, auto-generates from headers."""
        self.table_counter += 1

        # Table caption ABOVE the table using Caption style with SEQ field
        cap_p = self.doc.add_paragraph(style='Caption')
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_prefix = cap_p.add_run("TABLE ")
        run_prefix.font.name = FONT_NAME
        run_prefix.font.size = CAPTION_SIZE
        run_prefix.font.bold = True
        run_prefix.font.italic = False

        # Add SEQ Table field (auto-number)
        add_seq_field(cap_p, "Table", number=self.table_counter)

        # Generate caption text from headers if not provided
        if caption is None:
            caption = " / ".join(headers[:3])
            if len(headers) > 3:
                caption += " / ..."

        run_desc = cap_p.add_run(": " + caption)
        run_desc.font.name = FONT_NAME
        run_desc.font.size = CAPTION_SIZE
        run_desc.font.italic = False
        cap_p.paragraph_format.space_after = Pt(4)

        # Create the table
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_table_borders(table)

        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(cell, "2B579A")

        # Data rows
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.name = FONT_NAME
                run.font.size = Pt(9)
                if r_idx % 2 == 1:
                    set_cell_shading(cell, "F2F2F2")

        if col_widths:
            for i, w in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = w

        self.doc.add_paragraph()  # spacing
        return table

    def add_page_break(self):
        self.doc.add_page_break()

    # ── DOCUMENT SECTIONS ──

    def build_title_page(self):
        for _ in range(6):
            self.doc.add_paragraph()
        self.add_paragraph("PhishAware: A Bilingual Cybersecurity Awareness Platform\nfor Saudi Arabian Organizations",
                          bold=True, size=Pt(16), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
        self.add_paragraph("Senior Project Report", bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36))
        self.add_paragraph("PhishAware Team", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
        self.add_paragraph("Department of Computer Science", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
        self.add_paragraph("King Abdulaziz University", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
        self.add_paragraph("Jeddah, Saudi Arabia", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
        self.add_paragraph("Supervisor: Dr. [Supervisor Name]", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
        self.add_paragraph("March 2026", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER)
        self.add_page_break()

    def build_abstract(self):
        self.doc.add_heading("ABSTRACT", level=1)
        abstract_text = (
            "Abstract\u2014 This report presents PhishAware, a comprehensive bilingual (English/Arabic) "
            "cybersecurity awareness platform designed specifically for Saudi Arabian organizations. "
            "The platform addresses the growing need for effective security awareness training in the "
            "Kingdom's rapidly digitizing economy by providing an integrated suite of tools including "
            "realistic phishing simulation campaigns, adaptive risk scoring, gamified training modules, "
            "and multi-level analytics dashboards. The backend is built with Django REST Framework "
            "using a modular multi-app architecture encompassing eleven specialized applications: "
            "accounts, companies, campaigns, simulations, assessments, training, gamification, "
            "analytics, community, notifications, and core utilities. The frontend is developed using "
            "React 19 with Vite, featuring role-based access control (RBAC) for three user tiers "
            "(Super Admin, Company Admin, Employee), complete RTL support for Arabic, and responsive "
            "design using Tailwind CSS. Key features include JWT-based authentication with email "
            "verification, employee invitation workflows, a four-step simulation creation wizard, "
            "automated risk score calculation with remediation training assignment, a gamification "
            "engine with badges and leaderboards, and a public community portal for cybersecurity "
            "awareness. The platform uses SendGrid SMTP relay for transactional emails, SQLite for "
            "development storage, and implements comprehensive tracking for phishing simulations "
            "including pixel tracking, link click monitoring, and credential submission detection. "
            "Testing demonstrates the platform's effectiveness in raising security awareness through "
            "measurable risk score improvements after training interventions."
        )
        self.add_paragraph(abstract_text)
        self.add_paragraph(
            "Index Terms\u2014 Cybersecurity awareness, phishing simulation, Django REST Framework, "
            "React.js, bilingual platform, Saudi Arabia, gamification, RBAC, risk scoring, "
            "training remediation, SendGrid, JWT authentication.",
            italic=True, size=Pt(9)
        )
        self.add_page_break()

    def build_toc(self):
        self.doc.add_heading("CONTENTS", level=1)
        toc_items = [
            ("I.", "PHASE 1: BACKEND FOUNDATION", "1"),
            ("", "A. Environment Setup & Access Management (Weeks 1-2)", "2"),
            ("", "    1) Django & SQLite Initialization", "3"),
            ("", "    2) Identity & Access Management", "6"),
            ("", "B. Simulation Engine & Campaign Logic (Week 3)", "12"),
            ("", "    1) Core Database Models", "13"),
            ("", "    2) Campaign, Email & Quiz APIs", "20"),
            ("", "C. Threat Tracking & Risk Remediation (Weeks 4-5)", "28"),
            ("", "    1) Tracking Pixels, Lure Links & Landing Pages", "29"),
            ("", "    2) Risk Scoring & LMS Integration", "34"),
            ("", "D. Analytics, Gamification, & System Finalization (Week 5)", "41"),
            ("", "    1) Gamification Engine", "42"),
            ("", "    2) Dashboards, Public APIs & Comprehensive Testing", "48"),
            ("II.", "PHASE 2: FRONTEND DEVELOPMENT", "55"),
            ("", "A. Application Setup & Secure Routing (Week 9)", "56"),
            ("", "    1) React Core & Bilingual (i18next) Configuration", "57"),
            ("", "    2) Authentication & Protected Routes", "62"),
            ("", "B. Administrator Management Interfaces (Week 10)", "69"),
            ("", "    1) Admin Layout, Employee & Campaign Management", "70"),
            ("", "    2) Content Libraries & Analytics Dashboards", "78"),
            ("", "C. Employee Training & Gamification UI (Week 11)", "86"),
            ("", "    1) Employee Dashboard & Training Modules", "87"),
            ("", "    2) Quiz Interfaces & Gamification Display", "93"),
            ("", "D. Public Awareness Portal & Final Polish (Week 11)", "100"),
            ("", "    1) Community Portal & Awareness Content", "101"),
            ("", "    2) Public Quizzes & Final UI/UX Polish", "107"),
            ("", "REFERENCES", "113"),
            ("", "APPENDICES", "115"),
        ]
        for num, title, page in toc_items:
            p = self.doc.add_paragraph()
            full = f"{num} {title}" if num else f"    {title}"
            run = p.add_run(full)
            run.font.name = FONT_NAME
            run.font.size = Pt(10)
            if num:
                run.font.bold = True
            # Add tab + page number
            run2 = p.add_run(f"\t{page}")
            run2.font.name = FONT_NAME
            run2.font.size = Pt(10)
        self.add_page_break()

    # ═══════════════════════════════════════════════════════════
    # PHASE 1: BACKEND FOUNDATION
    # ═══════════════════════════════════════════════════════════

    def build_phase1(self):
        self.doc.add_heading("I. PHASE 1: BACKEND FOUNDATION", level=1)
        self.add_paragraph(
            "Phase 1 encompasses the complete backend development of the PhishAware platform, "
            "built using Django 5.x and Django REST Framework. This phase establishes the server-side "
            "architecture, database models, API endpoints, authentication system, simulation engine, "
            "risk scoring algorithms, gamification mechanics, and analytics infrastructure. The backend "
            "follows a modular multi-app architecture with eleven specialized Django applications, each "
            "handling a distinct domain of the platform's functionality. All APIs follow RESTful conventions "
            "with JWT-based authentication, role-based access control, and comprehensive serialization "
            "for data validation and transformation."
        )

        self._build_section_A()
        self._build_section_B()
        self._build_section_C()
        self._build_section_D()

    def _build_section_A(self):
        self.doc.add_heading("A. Environment Setup & Access Management (Weeks 1-2)", level=2)
        self.add_paragraph(
            "The initial two weeks focused on establishing the development environment, configuring "
            "the Django project structure, implementing the custom user model with role-based access "
            "control, and building the authentication system with JWT tokens and email verification. "
            "This foundational work ensures secure identity management across all three user tiers: "
            "Super Admin, Company Admin, and Employee."
        )

        # ── 1) Django & SQLite Initialization ──
        self.doc.add_heading("1) Django & SQLite Initialization", level=3)
        self.add_paragraph(
            "The PhishAware backend is initialized as a Django 5.x project with a modular multi-app "
            "architecture. The project uses SQLite as its development database, leveraging Django's "
            "built-in ORM for all data access patterns. The choice of SQLite provides zero-configuration "
            "persistence during development while maintaining full compatibility with production-grade "
            "databases like PostgreSQL through Django's database abstraction layer."
        )

        self.add_paragraph("Project Structure and Application Organization", bold=True)
        self.add_paragraph(
            "The backend follows Django's recommended project layout with a central configuration "
            "package (phishaware_backend/) and a dedicated apps/ directory containing eleven specialized "
            "Django applications. Each application encapsulates a distinct domain of the platform's "
            "business logic, following the Single Responsibility Principle:"
        )

        self.add_table(
            ["Application", "Purpose", "Key Models"],
            [
                ["accounts", "User authentication, registration, invitation", "User"],
                ["companies", "Company management and subscription", "Company"],
                ["campaigns", "Quiz campaign creation and execution", "Campaign, Quiz, QuizResult"],
                ["assessments", "Email templates and quiz questions", "EmailTemplate, QuizQuestion"],
                ["simulations", "Live phishing simulation engine", "SimulationTemplate, SimulationCampaign, EmailSimulation, TrackingEvent"],
                ["training", "Risk scoring and remediation training", "RiskScore, TrainingModule, RemediationTraining"],
                ["gamification", "Badges, points, and leaderboards", "Badge, EmployeeBadge, EmployeePoints, PointsTransaction"],
                ["analytics", "Dashboard statistics and data export", "(No models \u2013 queryset-based)"],
                ["community", "Public portal content", "Article, PublicQuiz, Resource"],
                ["notifications", "In-app notification system", "Notification"],
                ["core", "Shared utilities, permissions, emails", "(Utility module)"],
            ]
,
            caption="Django Application Architecture"
        )

        self.add_paragraph("Django Settings Configuration", bold=True)
        self.add_paragraph(
            "The settings.py file configures Django with the following key parameters. Environment "
            "variables are managed through python-decouple for secure configuration management, "
            "separating sensitive credentials from the codebase:"
        )

        self.add_code_block(
            "# phishaware_backend/settings.py (key excerpts)\n"
            "AUTH_USER_MODEL = 'accounts.User'\n"
            "INSTALLED_APPS = [\n"
            "    'django.contrib.admin',\n"
            "    'rest_framework',\n"
            "    'rest_framework_simplejwt',\n"
            "    'rest_framework_simplejwt.token_blacklist',\n"
            "    'corsheaders',\n"
            "    'drf_yasg',\n"
            "    'apps.core', 'apps.accounts', 'apps.companies',\n"
            "    'apps.campaigns', 'apps.simulations', 'apps.assessments',\n"
            "    'apps.training', 'apps.gamification', 'apps.analytics',\n"
            "    'apps.community', 'apps.notifications',\n"
            "]\n"
            "DATABASES = {'default': {'ENGINE': 'django.db.backends.sqlite3',\n"
            "                         'NAME': BASE_DIR / 'db.sqlite3'}}"
        )

        self.add_paragraph("REST Framework and JWT Configuration", bold=True)
        self.add_paragraph(
            "Django REST Framework is configured with JWT authentication as the primary authentication "
            "backend. Access tokens have a lifetime of 1 hour while refresh tokens persist for 7 days "
            "with rotation enabled. Token blacklisting is implemented to support secure logout by "
            "invalidating refresh tokens on the server side:"
        )

        self.add_code_block(
            "REST_FRAMEWORK = {\n"
            "    'DEFAULT_AUTHENTICATION_CLASSES': [\n"
            "        'rest_framework_simplejwt.authentication.JWTAuthentication',\n"
            "    ],\n"
            "    'DEFAULT_PAGINATION_CLASS':\n"
            "        'rest_framework.pagination.PageNumberPagination',\n"
            "    'PAGE_SIZE': 20,\n"
            "}\n"
            "SIMPLE_JWT = {\n"
            "    'ACCESS_TOKEN_LIFETIME': timedelta(hours=1),\n"
            "    'REFRESH_TOKEN_LIFETIME': timedelta(days=7),\n"
            "    'ROTATE_REFRESH_TOKENS': True,\n"
            "    'BLACKLIST_AFTER_ROTATION': True,\n"
            "}"
        )

        self.add_paragraph("CORS and Middleware Stack", bold=True)
        self.add_paragraph(
            "Cross-Origin Resource Sharing (CORS) is handled through django-cors-headers middleware, "
            "allowing the React frontend (running on localhost:5173 during development) to communicate "
            "with the Django API server. The middleware stack includes Django's security middleware, "
            "session middleware, CSRF protection, and authentication middleware in the standard order."
        )

        self.add_paragraph("Email System Configuration", bold=True)
        self.add_paragraph(
            "The platform uses SendGrid's SMTP relay for transactional email delivery. Configuration "
            "is managed through environment variables to keep API keys secure. The system supports "
            "four email types: verification, invitation, password reset, and simulation phishing emails. "
            "HTML email templates are stored in apps/core/templates/emails/ and rendered using Django's "
            "template engine:"
        )

        self.add_table(
            ["Setting", "Value", "Purpose"],
            [
                ["EMAIL_HOST", "smtp.sendgrid.net", "SendGrid SMTP relay server"],
                ["EMAIL_PORT", "587", "TLS-encrypted SMTP port"],
                ["EMAIL_HOST_USER", "apikey", "SendGrid authentication username"],
                ["EMAIL_HOST_PASSWORD", "(env var)", "SendGrid API key"],
                ["DEFAULT_FROM_EMAIL", "PhishAware <noreply@phishaware.com>", "Sender address for all emails"],
                ["FRONTEND_URL", "http://localhost:5173", "Base URL for email links (verification, invitation)"],
                ["SITE_URL", "http://localhost:8000", "Base URL for tracking pixels and API links"],
            ]
,
            caption="Django Settings Configuration"
        )

        self.add_paragraph("URL Routing Architecture", bold=True)
        self.add_paragraph(
            "The main urls.py file defines the top-level URL routing, mounting each application's "
            "URL configuration under the /api/v1/ namespace. This versioned API structure supports "
            "future API evolution while maintaining backward compatibility. Swagger UI and ReDoc "
            "documentation are automatically generated via drf-yasg:"
        )

        self.add_code_block(
            "# phishaware_backend/urls.py\n"
            "urlpatterns = [\n"
            "    path('api/docs/', schema_view.with_ui('swagger')),\n"
            "    path('api/v1/auth/', include('apps.accounts.urls')),\n"
            "    path('api/v1/employees/', include('apps.accounts.urls_employees')),\n"
            "    path('api/v1/companies/', include('apps.companies.urls')),\n"
            "    path('api/v1/campaigns/', include('apps.campaigns.urls')),\n"
            "    path('api/v1/simulations/', include('apps.simulations.urls')),\n"
            "    path('api/v1/training/', include('apps.training.urls')),\n"
            "    path('api/v1/gamification/', include('apps.gamification.urls')),\n"
            "    path('api/v1/analytics/', include('apps.analytics.urls')),\n"
            "    path('api/v1/community/', include('apps.community.urls')),\n"
            "    path('api/v1/notifications/', include('apps.notifications.urls')),\n"
            "]"
        )

        self.add_paragraph("Dependency Management", bold=True)
        self.add_paragraph(
            "The project's dependencies are managed through requirements.txt, listing all Python "
            "packages required for the backend. Core dependencies include Django 5.x, Django REST "
            "Framework 3.14+, djangorestframework-simplejwt for JWT token management, "
            "django-cors-headers for CORS handling, drf-yasg for API documentation generation, "
            "python-decouple for environment variable management, and python-dateutil for date "
            "manipulation. The SendGrid SDK is available but the platform primarily uses SMTP relay "
            "for email delivery."
        )

        # ── 2) Identity & Access Management ──
        self.doc.add_heading("2) Identity & Access Management", level=3)
        self.add_paragraph(
            "The identity and access management system is built around a custom User model extending "
            "Django's AbstractBaseUser. This provides full control over user fields, authentication "
            "logic, and role-based permissions. The system supports four user roles with hierarchical "
            "access control, email verification for new registrations, employee invitation workflows "
            "with token-based acceptance, and password reset functionality."
        )

        self.add_paragraph("Custom User Model", bold=True)
        self.add_paragraph(
            "The User model in apps/accounts/models.py replaces Django's default user model with "
            "email-based authentication (no username field). It includes comprehensive fields for "
            "user profile, role assignment, company association, email verification, invitation "
            "management, and password reset:"
        )

        self.add_table(
            ["Field", "Type", "Description"],
            [
                ["email", "EmailField (unique)", "Primary login identifier"],
                ["first_name", "CharField(100)", "User's first name"],
                ["last_name", "CharField(100)", "User's last name"],
                ["phone_number", "CharField(20, optional)", "Contact phone number"],
                ["role", "CharField(choices)", "SUPER_ADMIN, COMPANY_ADMIN, EMPLOYEE, PUBLIC_USER"],
                ["company", "ForeignKey(Company, nullable)", "Associated company"],
                ["preferred_language", "CharField(choices)", "en or ar"],
                ["is_verified", "BooleanField", "Email verification status"],
                ["verification_token", "UUIDField(unique)", "Email verification token"],
                ["verification_token_created", "DateTimeField", "Token creation timestamp"],
                ["invitation_token", "UUIDField(unique, nullable)", "Employee invitation token"],
                ["invitation_sent_at", "DateTimeField", "Invitation send timestamp"],
                ["invitation_accepted_at", "DateTimeField", "Invitation acceptance timestamp"],
                ["invitation_status", "CharField(choices)", "PENDING, ACCEPTED, EXPIRED"],
                ["password_reset_token", "UUIDField(unique)", "Password reset token"],
                ["password_reset_token_created", "DateTimeField", "Reset token creation timestamp"],
            ]
,
            caption="Custom User Model Fields"
        )

        self.add_paragraph("Role-Based Access Control (RBAC)", bold=True)
        self.add_paragraph(
            "The platform implements four user roles with distinct access levels. Custom permission "
            "classes in apps/core/permissions.py enforce access control at the view level:"
        )

        self.add_table(
            ["Role", "Access Level", "Key Permissions"],
            [
                ["SUPER_ADMIN", "Platform-wide", "Manage all companies, view all data, create/delete companies, platform analytics"],
                ["COMPANY_ADMIN", "Company-scoped", "Manage employees, create campaigns/simulations/training, company analytics, invite employees"],
                ["EMPLOYEE", "Personal + Company", "Take quizzes, view training, view own risk score, company leaderboard, receive simulations"],
                ["PUBLIC_USER", "Public only", "Access community portal, take public quizzes, view awareness content"],
            ]
,
            caption="Role-Based Access Control Summary"
        )

        self.add_code_block(
            "# apps/core/permissions.py\n"
            "class IsSuperAdmin(BasePermission):\n"
            "    def has_permission(self, request, view):\n"
            "        return request.user.is_authenticated and request.user.is_super_admin\n\n"
            "class IsCompanyAdmin(BasePermission):\n"
            "    def has_permission(self, request, view):\n"
            "        return request.user.is_authenticated and request.user.is_company_admin\n\n"
            "class IsSameCompany(BasePermission):\n"
            "    def has_object_permission(self, request, view, obj):\n"
            "        return obj.company == request.user.company"
        )

        self.add_paragraph("Authentication API Endpoints", bold=True)
        self.add_paragraph(
            "The authentication system provides comprehensive endpoints for user registration, "
            "login, logout, email verification, password management, and profile operations. "
            "JWT tokens are issued on successful login, with access tokens for API authentication "
            "and refresh tokens for session renewal:"
        )

        self.add_table(
            ["Endpoint", "Method", "Auth", "Description"],
            [
                ["/api/v1/auth/register/", "POST", "None", "User registration with email verification"],
                ["/api/v1/auth/login/", "POST", "None", "Login with JWT token issuance"],
                ["/api/v1/auth/logout/", "POST", "JWT", "Logout with refresh token blacklisting"],
                ["/api/v1/auth/token/refresh/", "POST", "None", "Refresh access token using refresh token"],
                ["/api/v1/auth/verify-email/<token>/", "POST", "None", "Verify email with UUID token"],
                ["/api/v1/auth/resend-verification/", "POST", "None", "Resend verification email"],
                ["/api/v1/auth/password-reset/", "POST", "None", "Request password reset email"],
                ["/api/v1/auth/password-reset/<token>/", "POST", "None", "Reset password with token"],
                ["/api/v1/auth/profile/", "GET/PATCH", "JWT", "View or update user profile"],
                ["/api/v1/auth/change-password/", "POST", "JWT", "Change password (requires current password)"],
            ]
,
            caption="Authentication API Endpoints"
        )

        self.add_paragraph("Email Verification Flow", bold=True)
        self.add_paragraph(
            "When a new user registers, the system generates a UUID verification token and sends "
            "a verification email containing a link to the frontend verification page. The link "
            "format is {FRONTEND_URL}/verify-email/{token}. Unverified users cannot log in\u2014the "
            "CustomTokenObtainPairSerializer checks is_verified and returns a 400 error with an "
            "email_not_verified flag if the user has not verified their email. Staff users "
            "(is_staff=True) bypass this check."
        )

        self.add_paragraph("Employee Invitation System", bold=True)
        self.add_paragraph(
            "Company administrators can invite employees to join the platform. The invitation "
            "workflow creates a User record with is_active=False and invitation_status='PENDING', "
            "then sends an invitation email with a link to accept. Invited users set their password "
            "during acceptance and are automatically marked as verified (is_verified=True). "
            "Invitations expire after 7 days:"
        )

        self.add_table(
            ["Endpoint", "Method", "Auth", "Description"],
            [
                ["/api/v1/employees/invite/", "POST", "Admin", "Send employee invitation"],
                ["/api/v1/employees/invite/<token>/", "GET", "None", "Get invitation details"],
                ["/api/v1/employees/invite/<token>/accept/", "POST", "None", "Accept invitation and set password"],
                ["/api/v1/employees/pending-invitations/", "GET", "Admin", "List pending invitations"],
                ["/api/v1/employees/users/<id>/resend-invitation/", "POST", "Admin", "Resend invitation email"],
                ["/api/v1/employees/users/<id>/cancel-invitation/", "DELETE", "Admin", "Cancel pending invitation"],
            ]
,
            caption="Employee Invitation API Endpoints"
        )

        self.add_paragraph("Company Management", bold=True)
        self.add_paragraph(
            "The companies app handles organization-level management including company registration, "
            "subscription tracking, and user management within companies. The Company model stores "
            "organizational details including bilingual name/description, industry classification, "
            "company size, and subscription dates:"
        )

        self.add_table(
            ["Field", "Type", "Description"],
            [
                ["name / name_ar", "CharField", "Company name (English and Arabic)"],
                ["description / description_ar", "TextField", "Company description (bilingual)"],
                ["industry", "CharField(choices)", "10 industry categories"],
                ["company_size", "CharField(choices)", "6 size ranges (1-10 to 10000+)"],
                ["email / phone / website", "Various", "Contact information"],
                ["country / city / address", "CharField", "Location details"],
                ["subscription_start_date / end_date", "DateField", "Subscription period"],
                ["is_active", "BooleanField", "Company activation status"],
            ]
,
            caption="Company Model Fields"
        )

    def _build_section_B(self):
        self.doc.add_heading("B. Simulation Engine & Campaign Logic (Week 3)", level=2)
        self.add_paragraph(
            "Week 3 focused on building the core simulation and campaign infrastructure. This includes "
            "the database models for phishing simulations, quiz campaigns, email templates, and the "
            "API endpoints for creating, managing, and executing campaigns. The simulation engine "
            "supports both quiz-based campaigns (where employees classify emails) and live phishing "
            "simulations (where real emails are sent to test employee responses)."
        )

        # ── 1) Core Database Models ──
        self.doc.add_heading("1) Core Database Models", level=3)
        self.add_paragraph(
            "The PhishAware platform's data layer spans multiple Django applications, each containing "
            "domain-specific models. The core models establish the foundation for campaign execution, "
            "quiz management, phishing simulation tracking, and email template storage. All models "
            "follow Django's ORM patterns with proper foreign key relationships, unique constraints, "
            "and denormalized statistics for performance optimization."
        )

        self.add_paragraph("Campaign Models (campaigns app)", bold=True)
        self.add_paragraph(
            "The Campaign model represents a quiz-based assessment where employees classify emails "
            "as phishing or legitimate. Each campaign is associated with a company, has configurable "
            "parameters for the number of emails and phishing ratio, and tracks participant statistics "
            "through denormalized counters:"
        )

        self.add_table(
            ["Model", "Key Fields", "Relationships", "Purpose"],
            [
                ["Campaign", "name, status (DRAFT/ACTIVE/SCHEDULED/COMPLETED/ARCHIVED), num_emails (5-50), phishing_ratio (0.2-0.8), english_ratio", "FK: company, created_by", "Quiz campaign container"],
                ["Quiz", "status, current_question_index, started_at, completed_at", "FK: campaign, employee (unique together)", "Per-employee quiz instance"],
                ["QuizResult", "total_questions, correct_answers, score (%), phishing_detection_rate, risk_level", "FK: quiz, campaign, employee", "Final quiz scoring"],
            ]
,
            caption="Campaign Models Overview"
        )

        self.add_paragraph("Assessment Models (assessments app)", bold=True)
        self.add_paragraph(
            "The assessments app stores email templates and quiz questions used in campaigns. "
            "EmailTemplate supports both phishing and legitimate email types with bilingual content, "
            "AI generation metadata, and educational red flags for post-quiz learning:"
        )

        self.add_table(
            ["Model", "Key Fields", "Purpose"],
            [
                ["EmailTemplate", "email_type (PHISHING/LEGITIMATE), category (10 types), subject, body, sender_name/email, difficulty, red_flags (JSON), is_ai_generated", "Email content for quiz questions"],
                ["QuizQuestion", "quiz (FK), email_template (FK), question_number, answer (PHISHING/LEGITIMATE), is_correct, confidence_level, time_spent_seconds", "Individual question response in a quiz"],
            ]
,
            caption="Assessment Models Overview"
        )

        self.add_paragraph("Simulation Models (simulations app)", bold=True)
        self.add_paragraph(
            "The simulation engine uses a three-tier model hierarchy: templates define reusable "
            "phishing email content, campaigns orchestrate sending to target employees, and email "
            "simulations track individual delivery and interaction status. TrackingEvent provides "
            "an immutable event log for all simulation interactions:"
        )

        self.add_table(
            ["Model", "Key Fields", "Purpose"],
            [
                ["SimulationTemplate", "sender_name/email, subject, body_html, attack_vector (7 types), difficulty, red_flags (JSON), landing_page_title/message, is_public, language", "Reusable phishing email template"],
                ["SimulationCampaign", "name, template (FK), company (FK), status, target_employees (M2M), track_email_opens/link_clicks/credentials, total_sent/clicked/reported (denormalized)", "Campaign execution instance"],
                ["EmailSimulation", "campaign (FK), employee (FK), tracking_token (UUID), link_token (unique 64-char), status, was_opened/clicked/reported, timestamps", "Per-employee email tracking"],
                ["TrackingEvent", "email_simulation (FK), event_type (8 types), event_data (JSON), ip_address, user_agent, geolocation (JSON)", "Immutable interaction event log"],
            ]
,
            caption="Simulation Models Overview"
        )

        self.add_paragraph(
            "The TrackingEvent model's save() method implements a critical side-effect chain: after "
            "persisting the event, it updates the corresponding EmailSimulation flags (was_opened, "
            "was_clicked, etc.), then propagates changes to the parent SimulationCampaign's denormalized "
            "statistics (total_opened, total_clicked, etc.). A Django post_save signal also triggers "
            "risk score recalculation in the training app."
        )

        self.add_paragraph("Training and Risk Score Models (training app)", bold=True)
        self.add_paragraph(
            "The training app contains models for risk scoring, training modules, remediation "
            "assignments, and interactive lesson progress. The RiskScore model maintains a per-employee "
            "security risk assessment based on quiz performance, simulation responses, and training "
            "completion:"
        )

        self.add_table(
            ["Model", "Key Fields", "Purpose"],
            [
                ["RiskScore", "score (0-100), risk_level, quiz stats, simulation stats, training stats, requires_remediation", "Per-employee risk assessment"],
                ["RiskScoreHistory", "previous_score, new_score, score_change, event_type (12 types), source_type/id", "Risk score change audit trail"],
                ["TrainingModule", "title, content_type (5 types), category (10 types), difficulty, passing_score (80%), score_reduction_on_pass (15 points)", "Training content library"],
                ["TrainingQuestion", "module (FK), question_text, options/options_ar (JSON), correct_answer_index, explanation", "Training quiz questions"],
                ["RemediationTraining", "employee (FK), module (FK), status, reason (6 types), quiz_score, content_viewed, time_spent_seconds", "Training assignment and tracking"],
                ["InteractiveLessonProgress", "user (FK), lesson_type, portal_type, current_scene, total_scenes, quiz_score, is_completed", "Interactive lesson state tracking"],
            ]
,
            caption="Training and Risk Score Models Overview"
        )

        # ── 2) Campaign, Email & Quiz APIs ──
        self.doc.add_heading("2) Campaign, Email & Quiz APIs", level=3)
        self.add_paragraph(
            "The API layer exposes comprehensive endpoints for campaign management, email template "
            "CRUD operations, quiz execution, and simulation control. All endpoints follow RESTful "
            "conventions with proper HTTP method semantics, nested resource routing, and role-based "
            "access control through DRF permission classes."
        )

        self.add_paragraph("Campaign Management API", bold=True)
        self.add_paragraph(
            "The CampaignViewSet provides full CRUD operations for quiz campaigns with additional "
            "custom actions for campaign lifecycle management. Company administrators can create "
            "campaigns with configurable parameters, publish them to make quizzes available to "
            "employees, and track completion statistics:"
        )

        self.add_table(
            ["Endpoint", "Method", "Description"],
            [
                ["/api/v1/campaigns/campaigns/", "GET", "List campaigns (filtered by company)"],
                ["/api/v1/campaigns/campaigns/", "POST", "Create new campaign"],
                ["/api/v1/campaigns/campaigns/{id}/", "GET", "Campaign details with statistics"],
                ["/api/v1/campaigns/campaigns/{id}/", "PUT/PATCH", "Update campaign"],
                ["/api/v1/campaigns/campaigns/{id}/", "DELETE", "Delete campaign"],
                ["/api/v1/campaigns/campaigns/{id}/publish/", "POST", "Publish campaign (DRAFT \u2192 ACTIVE)"],
            ]
,
            caption="Campaign Management API Endpoints"
        )

        self.add_paragraph("Quiz Execution API", bold=True)
        self.add_paragraph(
            "The QuizViewSet manages employee quiz instances, providing endpoints for starting quizzes, "
            "answering individual questions, and submitting completed quizzes. Each quiz presents "
            "email templates that employees must classify as phishing or legitimate:"
        )

        self.add_table(
            ["Endpoint", "Method", "Description"],
            [
                ["/api/v1/campaigns/quizzes/", "GET", "List employee's quizzes"],
                ["/api/v1/campaigns/quizzes/{id}/", "GET", "Quiz details with progress"],
                ["/api/v1/campaigns/quizzes/{id}/start/", "POST", "Start a quiz attempt"],
                ["/api/v1/campaigns/quizzes/{id}/answer/", "POST", "Answer current question"],
                ["/api/v1/campaigns/quizzes/{id}/submit/", "POST", "Submit completed quiz for scoring"],
            ]
,
            caption="Quiz Execution API Endpoints"
        )

        self.add_paragraph("Simulation Campaign API", bold=True)
        self.add_paragraph(
            "The SimulationCampaignViewSet manages live phishing simulation campaigns. The send "
            "action triggers actual email delivery to target employees via SendGrid SMTP, creating "
            "EmailSimulation records with unique tracking tokens for each recipient:"
        )

        self.add_table(
            ["Endpoint", "Method", "Description"],
            [
                ["/api/v1/simulations/campaigns/", "GET/POST", "List or create simulation campaigns"],
                ["/api/v1/simulations/campaigns/{id}/", "GET/PUT/DELETE", "CRUD for individual campaign"],
                ["/api/v1/simulations/campaigns/{id}/send/", "POST", "Send simulation emails to targets"],
                ["/api/v1/simulations/campaigns/{id}/cancel/", "POST", "Cancel active campaign"],
                ["/api/v1/simulations/campaigns/{id}/analytics/", "GET", "Campaign analytics and statistics"],
                ["/api/v1/simulations/templates/", "GET/POST", "List or create simulation templates"],
                ["/api/v1/simulations/templates/{id}/", "GET/PUT/DELETE", "CRUD for individual template"],
            ]
,
            caption="Simulation Campaign API Endpoints"
        )

        self.add_paragraph("Email Template System", bold=True)
        self.add_paragraph(
            "Email templates support two types: PHISHING and LEGITIMATE, with 10 category subcategories "
            "each. Templates include bilingual content (English/Arabic), difficulty levels, red flags "
            "for educational purposes, and AI generation metadata. The simulation engine replaces "
            "template placeholders ({TRACKING_PIXEL}, {LURE_LINK}, {EMPLOYEE_NAME}, {EMPLOYEE_EMAIL}) "
            "with actual tracking URLs and employee details during email composition."
        )

        self.add_paragraph("Serializer Architecture", bold=True)
        self.add_paragraph(
            "The API uses a layered serializer pattern with separate serializers for list views "
            "(lightweight, fewer fields), detail views (comprehensive with nested data), and create "
            "views (validation-focused with write-only fields). This pattern optimizes payload sizes "
            "for list operations while providing complete data for detail views. Serializers include "
            "computed fields such as completion_rate, click_rate, and risk_level that are calculated "
            "from the underlying model properties."
        )

    def _build_section_C(self):
        self.doc.add_heading("C. Threat Tracking & Risk Remediation (Weeks 4-5)", level=2)
        self.add_paragraph(
            "Weeks 4-5 focused on implementing the threat tracking infrastructure for phishing "
            "simulations and the risk scoring algorithm with automatic training remediation. This "
            "includes tracking pixels for email opens, lure link click monitoring, credential "
            "submission detection, and a comprehensive risk scoring system that automatically assigns "
            "remediation training to high-risk employees."
        )

        # ── 1) Tracking Pixels, Lure Links & Landing Pages ──
        self.doc.add_heading("1) Tracking Pixels, Lure Links & Landing Pages", level=3)
        self.add_paragraph(
            "The simulation tracking system implements three primary detection mechanisms: tracking "
            "pixels embedded in email HTML for open detection, lure links that redirect through the "
            "API for click tracking, and credential submission forms that capture (but do not store) "
            "login attempts. All tracking endpoints are public (AllowAny) since they must be "
            "accessible from email clients without authentication."
        )

        self.add_paragraph("Tracking Pixel Implementation", bold=True)
        self.add_paragraph(
            "When a simulation email is composed, the {TRACKING_PIXEL} placeholder is replaced with "
            "an <img> tag pointing to the tracking API. When the email client loads this invisible "
            "1x1 pixel image, the server records an EMAIL_OPENED tracking event and returns a "
            "transparent GIF. The tracking pixel URL includes the EmailSimulation's tracking_token "
            "for identification."
        )

        self.add_paragraph("Lure Link Click Tracking", bold=True)
        self.add_paragraph(
            "Phishing simulation emails contain lure links that route through the API endpoint "
            "/api/v1/simulations/link/{token}/. When an employee clicks the link, the server "
            "creates a LINK_CLICKED tracking event with IP address, user agent, and geolocation "
            "data, then redirects the user to the React frontend's \"You've been caught\" page "
            "at {FRONTEND_URL}/simulation/caught/{link_token}. This provides immediate educational "
            "feedback to the employee."
        )

        self.add_code_block(
            "# apps/simulations/views.py - Link click tracking\n"
            "@api_view(['GET'])\n"
            "@permission_classes([AllowAny])\n"
            "def track_link_click_view(request, token):\n"
            "    email_sim = get_object_or_404(EmailSimulation, link_token=token)\n"
            "    TrackingEvent.objects.create(\n"
            "        email_simulation=email_sim,\n"
            "        campaign=email_sim.campaign,\n"
            "        employee=email_sim.employee,\n"
            "        event_type='LINK_CLICKED',\n"
            "        ip_address=get_client_ip(request),\n"
            "        user_agent=request.META.get('HTTP_USER_AGENT', ''),\n"
            "    )\n"
            "    redirect_url = f'{FRONTEND_URL}/simulation/caught/{token}'\n"
            "    return redirect(redirect_url)"
        )

        self.add_paragraph("Landing Page and Feedback System", bold=True)
        self.add_paragraph(
            "After clicking a phishing link, the employee is redirected to a React-based educational "
            "landing page (SimulationCaught.jsx). This page fetches feedback data from the API endpoint "
            "/api/v1/simulations/feedback/{token}/, which returns the template's red flags, difficulty "
            "level, attack vector, and educational explanation in the appropriate language. The feedback "
            "system supports both English and Arabic, automatically selecting the language based on the "
            "simulation template's language setting."
        )

        self.add_paragraph("Tracking Event Types", bold=True)
        self.add_table(
            ["Event Type", "Trigger", "Data Captured"],
            [
                ["EMAIL_SENT", "Email dispatched via SMTP", "Timestamp, recipient details"],
                ["EMAIL_DELIVERED", "Delivery confirmation", "Delivery timestamp"],
                ["EMAIL_BOUNCED", "Delivery failure", "Bounce reason, error details"],
                ["EMAIL_OPENED", "Tracking pixel loaded", "IP address, user agent, timestamp"],
                ["LINK_CLICKED", "Lure link accessed", "IP address, user agent, geolocation"],
                ["CREDENTIALS_ENTERED", "Fake login form submitted", "IP address, user agent (credentials not stored)"],
                ["EMAIL_REPORTED", "Employee reports phishing", "Report timestamp, reporter details"],
                ["LANDING_PAGE_VIEWED", "Educational page accessed", "IP address, view duration"],
            ]
,
            caption="Tracking Event Types"
        )

        self.add_paragraph("Event Propagation Chain", bold=True)
        self.add_paragraph(
            "When a TrackingEvent is saved, a cascade of updates occurs: (1) The Django post_save "
            "signal fires, triggering risk score recalculation in the training app. (2) The "
            "_update_simulation_stats() method updates the parent EmailSimulation's boolean flags "
            "(was_opened, was_clicked, etc.) and timestamps. (3) The _update_campaign_stats() method "
            "recalculates the SimulationCampaign's denormalized counters (total_opened, total_clicked, "
            "etc.) from the current EmailSimulation queryset. This ensures real-time statistics "
            "without expensive aggregate queries."
        )

        # ── 2) Risk Scoring & LMS Integration ──
        self.doc.add_heading("2) Risk Scoring & LMS Integration", level=3)
        self.add_paragraph(
            "The risk scoring system is the platform's core intelligence engine, calculating a "
            "comprehensive security risk score (0-100) for each employee based on their quiz "
            "performance, simulation responses, and training completion. The system automatically "
            "assigns remediation training when risk scores exceed thresholds, creating a closed-loop "
            "security awareness improvement cycle."
        )

        self.add_paragraph("Risk Score Calculation Algorithm", bold=True)
        self.add_paragraph(
            "The recalculate_score() method implements a multi-factor scoring algorithm that considers "
            "three primary dimensions: quiz accuracy, simulation response behavior, and training "
            "completion. The base score starts at 50 (neutral) and is adjusted by weighted factors:"
        )

        self.add_code_block(
            "# apps/training/models.py - RiskScore.recalculate_score()\n"
            "def recalculate_score(self):\n"
            "    base_score = 50\n"
            "    # Quiz factor: -20 to +20 based on accuracy\n"
            "    quiz_adj = 0\n"
            "    if self.total_quizzes_taken > 0:\n"
            "        accuracy = self.correct_quiz_answers / self.total_quiz_questions\n"
            "        quiz_adj = (accuracy - 0.5) * -40  # inverted: high accuracy = low risk\n"
            "    # Simulation factor: +20 for clicks, -10 for reports\n"
            "    sim_adj = 0\n"
            "    if self.total_simulations_received > 0:\n"
            "        click_rate = self.simulations_clicked / self.total_simulations_received\n"
            "        sim_adj = click_rate * 40  # high click rate = high risk\n"
            "        report_rate = self.simulations_reported / self.total_simulations_received\n"
            "        sim_adj -= report_rate * 20  # reporting reduces risk\n"
            "    # Training factor: -15 for completion\n"
            "    training_adj = 0\n"
            "    if self.trainings_assigned > 0:\n"
            "        completion = self.trainings_completed / self.trainings_assigned\n"
            "        training_adj = (1 - completion) * 30 - 15\n"
            "    self.score = max(0, min(100, base_score + quiz_adj + sim_adj + training_adj))\n"
            "    self.risk_level = self.calculate_risk_level()"
        )

        self.add_paragraph("Risk Level Classification", bold=True)
        self.add_table(
            ["Risk Level", "Score Range", "Description", "Action"],
            [
                ["LOW", "0-25", "Employee demonstrates strong security awareness", "No intervention required"],
                ["MEDIUM", "26-50", "Moderate awareness with room for improvement", "Optional training recommended"],
                ["HIGH", "51-75", "Significant risk indicators present", "Remediation training auto-assigned"],
                ["CRITICAL", "76-100", "Severe vulnerability to phishing attacks", "Immediate training + admin alert"],
            ]
,
            caption="Risk Level Classification"
        )

        self.add_paragraph("Risk Score History and Audit Trail", bold=True)
        self.add_paragraph(
            "Every risk score change is recorded in the RiskScoreHistory model with 12 event types "
            "including QUIZ_COMPLETED, SIMULATION_CLICKED, CREDENTIALS_ENTERED, PHISHING_REPORTED, "
            "TRAINING_COMPLETED, and MANUAL_ADJUSTMENT. Each history entry records the previous score, "
            "new score, change delta, and the source object (quiz result, tracking event, or training "
            "assignment) that triggered the change. This provides a complete audit trail for compliance "
            "and analytics purposes."
        )

        self.add_paragraph("Automatic Remediation Training Assignment", bold=True)
        self.add_paragraph(
            "When an employee's risk score exceeds 70 (requires_remediation threshold), the system "
            "automatically creates RemediationTraining assignments. The assignment reason is set based "
            "on the trigger: AUTO_HIGH_RISK for score threshold, AUTO_SIMULATION_FAIL for phishing "
            "link clicks, AUTO_QUIZ_FAIL for poor quiz performance. Assignments include a due date, "
            "training module reference, and track completion through a multi-stage workflow: "
            "ASSIGNED \u2192 IN_PROGRESS \u2192 COMPLETED/PASSED/FAILED."
        )

        self.add_paragraph("Training Module System (LMS)", bold=True)
        self.add_paragraph(
            "The Learning Management System provides five content types: VIDEO, TEXT, INTERACTIVE, "
            "PDF, and SLIDES, across ten security categories. Each module includes bilingual content, "
            "a passing score threshold (default 80%), and a score reduction value (default 15 points) "
            "that is applied to the employee's risk score upon successful completion. The training "
            "quiz system allows multiple-choice questions with bilingual options and explanations:"
        )

        self.add_table(
            ["Endpoint", "Method", "Description"],
            [
                ["/api/v1/training/modules/", "GET", "List training modules"],
                ["/api/v1/training/modules/{id}/", "GET", "Module details with questions"],
                ["/api/v1/training/assignments/my_trainings/", "GET", "Employee's training assignments"],
                ["/api/v1/training/assignments/{id}/start/", "POST", "Start training assignment"],
                ["/api/v1/training/assignments/{id}/view_content/", "POST", "Mark content as viewed"],
                ["/api/v1/training/assignments/{id}/quiz/", "GET", "Get training quiz questions"],
                ["/api/v1/training/assignments/{id}/submit_quiz/", "POST", "Submit quiz answers"],
                ["/api/v1/training/assignments/bulk_assign/", "POST", "Bulk assign training to employees"],
                ["/api/v1/training/risk-scores/my_score/", "GET", "Employee's current risk score"],
                ["/api/v1/training/risk-scores/{id}/history/", "GET", "Risk score change history"],
            ]
,
            caption="Training API Endpoints"
        )

        self.add_paragraph("Seed Data Commands", bold=True)
        self.add_paragraph(
            "The platform includes Django management commands for seeding initial data: "
            "seed_simulation_templates creates 15 phishing email templates (8 English, 7 Arabic) "
            "with company=None for global availability, and seed_training creates 3 training modules "
            "(EMAIL_SECURITY, MOBILE_SECURITY, SOCIAL_ENGINEERING) each with 5 bilingual quiz questions. "
            "These seed templates provide immediate usable content after platform deployment."
        )

    def _build_section_D(self):
        self.doc.add_heading("D. Analytics, Gamification, & System Finalization (Week 5)", level=2)
        self.add_paragraph(
            "The final backend week focused on implementing the gamification engine for employee "
            "engagement, analytics dashboards for data-driven decision making, public community "
            "portal APIs, notification system, and comprehensive testing. These systems complete "
            "the backend feature set and provide the data visualization layer for administrators."
        )

        # ── 1) Gamification Engine ──
        self.doc.add_heading("1) Gamification Engine", level=3)
        self.add_paragraph(
            "The gamification system incentivizes security-conscious behavior through a badge and "
            "points economy with company leaderboards. The system awards points for completing quizzes, "
            "passing training, reporting phishing emails, and maintaining activity streaks. Badges "
            "provide visual recognition across five rarity tiers, while leaderboards create friendly "
            "competition within organizations."
        )

        self.add_paragraph("Badge System", bold=True)
        self.add_paragraph(
            "The Badge model defines nine achievement types that employees can earn through various "
            "security awareness activities. Each badge has a rarity level that determines its visual "
            "prominence and the points awarded upon earning it:"
        )

        self.add_table(
            ["Badge Type", "Rarity", "Criteria", "Points"],
            [
                ["FIRST_QUIZ_COMPLETED", "COMMON", "Complete first quiz", "10"],
                ["PERFECT_QUIZ_SCORE", "RARE", "Score 100% on any quiz", "50"],
                ["PHISH_SLAYER", "EPIC", "Identify all phishing emails in a campaign", "100"],
                ["TRAINING_CHAMPION", "UNCOMMON", "Complete all assigned training", "30"],
                ["SECURITY_AWARE", "RARE", "Maintain LOW risk score for 30 days", "75"],
                ["QUICK_LEARNER", "UNCOMMON", "Complete training within 24 hours of assignment", "25"],
                ["SIMULATION_SURVIVOR", "EPIC", "Report 5 consecutive phishing simulations", "100"],
                ["STREAK_MASTER", "LEGENDARY", "Maintain 30-day activity streak", "200"],
                ["TOP_REPORTER", "RARE", "Report 10+ phishing emails", "75"],
            ]
,
            caption="Badge Types and Criteria"
        )

        self.add_paragraph("Points Economy", bold=True)
        self.add_paragraph(
            "Every gamification-relevant action generates a PointsTransaction record with the "
            "transaction type, points value, running balance, and source reference. Points can be "
            "positive (awards) or negative (admin adjustments). The EmployeePoints model maintains "
            "denormalized aggregates including total_points, weekly_points, monthly_points, "
            "badge_count, and streak tracking for efficient leaderboard queries:"
        )

        self.add_table(
            ["Transaction Type", "Typical Points", "Trigger"],
            [
                ["QUIZ_COMPLETED", "+10-20", "Completing any quiz"],
                ["QUIZ_PERFECT_SCORE", "+50", "100% quiz accuracy"],
                ["PHISHING_REPORTED", "+15", "Reporting a simulation as phishing"],
                ["TRAINING_COMPLETED", "+10", "Completing training content"],
                ["TRAINING_PASSED", "+25", "Passing training quiz"],
                ["BADGE_EARNED", "Variable", "Earning a new badge"],
                ["FIRST_ATTEMPT_PASS", "+15", "Passing on first attempt"],
                ["STREAK_BONUS", "+5-20", "Maintaining activity streak"],
            ]
,
            caption="Points Transaction Types"
        )

        self.add_paragraph("Leaderboard System", bold=True)
        self.add_paragraph(
            "The LeaderboardViewSet provides ranked lists of employees by points with three time "
            "periods: weekly, monthly, and all-time. The leaderboard is company-scoped (employees "
            "only see colleagues) and includes the requesting user's current position, rank percentile, "
            "and comparison to top performers. The my_position endpoint returns the authenticated "
            "user's current ranking, total points, and the number of employees in their company."
        )

        self.add_table(
            ["Endpoint", "Method", "Description"],
            [
                ["/api/v1/gamification/badges/", "GET", "List all available badges"],
                ["/api/v1/gamification/badges/my_badges/", "GET", "User's earned badges"],
                ["/api/v1/gamification/badges/recent/", "GET", "Recently awarded badges"],
                ["/api/v1/gamification/points/my_summary/", "GET", "User's points summary"],
                ["/api/v1/gamification/points/my_transactions/", "GET", "Points transaction history"],
                ["/api/v1/gamification/leaderboard/", "GET", "Company leaderboard (period filter)"],
                ["/api/v1/gamification/leaderboard/my_position/", "GET", "User's current rank"],
            ]
,
            caption="Gamification API Endpoints"
        )

        # ── 2) Dashboards, Public APIs & Comprehensive Testing ──
        self.doc.add_heading("2) Dashboards, Public APIs & Comprehensive Testing", level=3)
        self.add_paragraph(
            "The analytics system provides multi-level dashboard APIs serving three user tiers: "
            "Super Admin (platform-wide metrics), Company Admin (organization-level insights), and "
            "Employee (personal performance tracking). The community portal APIs provide public "
            "access to awareness content, quizzes, and resources without authentication. The "
            "notification system enables real-time communication of platform events to users."
        )

        self.add_paragraph("Analytics Dashboard APIs", bold=True)
        self.add_paragraph(
            "The analytics app aggregates data from campaigns, simulations, training, and risk "
            "scores to provide comprehensive dashboard statistics. No separate models are needed\u2014"
            "all analytics are computed from existing querysets in real-time:"
        )

        self.add_table(
            ["Endpoint", "Method", "Data Provided"],
            [
                ["/api/v1/analytics/dashboard/overview/", "GET", "Total users, campaigns, simulations, training stats"],
                ["/api/v1/analytics/dashboard/trends/", "GET", "Time-series data with customizable period"],
                ["/api/v1/analytics/campaigns/", "GET", "Campaign performance analytics"],
                ["/api/v1/analytics/simulations/", "GET", "Simulation stats with template comparisons"],
                ["/api/v1/analytics/risk/distribution/", "GET", "Risk level distribution (pie chart data)"],
                ["/api/v1/analytics/risk/trends/", "GET", "Risk score trends over time (line chart data)"],
                ["/api/v1/analytics/risk/high_risk_employees/", "GET", "At-risk employee identification list"],
                ["/api/v1/analytics/training/", "GET", "Training completion and pass rate summary"],
                ["/api/v1/analytics/training/effectiveness/", "GET", "Risk reduction impact analysis"],
                ["/api/v1/analytics/export/csv/", "POST", "CSV data export with optional PII"],
            ]
,
            caption="Analytics Dashboard API Endpoints"
        )

        self.add_paragraph("Community Portal APIs (Public Access)", bold=True)
        self.add_paragraph(
            "The community app provides public-facing APIs for cybersecurity awareness content. "
            "All endpoints use AllowAny permissions, making them accessible without authentication. "
            "The portal supports articles, quizzes, and downloadable resources with bilingual content, "
            "view tracking, and anonymous quiz attempt recording:"
        )

        self.add_table(
            ["Resource", "Endpoints", "Features"],
            [
                ["Articles", "/api/v1/community/articles/", "Published articles, featured, recent, popular, by tag, share tracking"],
                ["Categories", "/api/v1/community/categories/", "Content categorization with article/quiz/resource counts"],
                ["Public Quizzes", "/api/v1/community/quizzes/", "Anonymous quiz taking with session tracking, scoring, difficulty filtering"],
                ["Resources", "/api/v1/community/resources/", "Downloadable PDFs, videos, guides, infographics with download tracking"],
                ["Portal", "/api/v1/community/portal/", "Homepage featured content, global search, portal statistics"],
            ]
,
            caption="Community Portal API Resources"
        )

        self.add_paragraph("Notification System", bold=True)
        self.add_paragraph(
            "The notification system supports 30+ notification types across training, simulation, "
            "account, campaign, and admin categories. Notifications include priority levels (LOW, "
            "MEDIUM, HIGH, URGENT) and optional links to related platform objects. The API provides "
            "endpoints for listing notifications, filtering unread, marking as read (individual or "
            "bulk), and retrieving notification counts for the header badge display:"
        )

        self.add_table(
            ["Endpoint", "Method", "Description"],
            [
                ["/api/v1/notifications/", "GET", "List user's notifications"],
                ["/api/v1/notifications/unread/", "GET", "Unread notifications only"],
                ["/api/v1/notifications/{id}/mark_as_read/", "POST", "Mark single notification as read"],
                ["/api/v1/notifications/mark_all_as_read/", "POST", "Mark all notifications as read"],
            ]
,
            caption="Notification API Endpoints"
        )

        self.add_paragraph("API Documentation", bold=True)
        self.add_paragraph(
            "The backend includes auto-generated API documentation using drf-yasg, which produces "
            "both Swagger UI (at /api/docs/) and ReDoc (at /api/redoc/) interfaces. These interactive "
            "documentation pages allow developers and administrators to explore all API endpoints, "
            "view request/response schemas, and test endpoints directly from the browser."
        )

    # ═══════════════════════════════════════════════════════════
    # PHASE 2: FRONTEND DEVELOPMENT
    # ═══════════════════════════════════════════════════════════

    def build_phase2(self):
        self.add_page_break()
        self.doc.add_heading("II. PHASE 2: FRONTEND DEVELOPMENT", level=1)
        self.add_paragraph(
            "Phase 2 covers the complete frontend development of the PhishAware platform using "
            "React 19 with Vite as the build tool. The frontend implements a responsive, bilingual "
            "(English/Arabic) single-page application with role-based routing, JWT token management, "
            "and Tailwind CSS for styling. The application comprises 43 page components organized "
            "across five sections: authentication (7 pages), company admin (10 pages), employee "
            "(8 pages), super admin (5 pages), and public portal (10 pages). The frontend communicates "
            "with the Django REST API through an Axios-based service layer with automatic token "
            "refresh and error handling."
        )

        self._build_frontend_A()
        self._build_frontend_B()
        self._build_frontend_C()
        self._build_frontend_D()

    def _build_frontend_A(self):
        self.doc.add_heading("A. Application Setup & Secure Routing (Week 9)", level=2)
        self.add_paragraph(
            "Week 9 established the React application foundation, including project initialization, "
            "bilingual support with RTL layout capabilities, the Axios API service layer with JWT "
            "token management, authentication pages, and the protected routing system. This "
            "foundational work ensures all subsequent frontend development inherits proper "
            "internationalization, authentication, and access control capabilities."
        )

        # ── 1) React Core & Bilingual (i18next) Configuration ──
        self.doc.add_heading("1) React Core & Bilingual (i18next) Configuration", level=3)
        self.add_paragraph(
            "The PhishAware frontend is built with React 19.2 and Vite 7.2, using a modern component "
            "architecture with functional components, hooks, and context-based state management. The "
            "application supports complete bilingual functionality with English and Arabic translations, "
            "including automatic RTL (right-to-left) layout transformation for Arabic content."
        )

        self.add_paragraph("Project Dependencies", bold=True)
        self.add_table(
            ["Package", "Version", "Purpose"],
            [
                ["React", "19.2", "UI component framework"],
                ["React Router DOM", "6.28", "Client-side routing"],
                ["Axios", "1.7.7", "HTTP client for API communication"],
                ["i18next", "23.16.4", "Internationalization framework"],
                ["react-i18next", "15.1.0", "React bindings for i18next"],
                ["Tailwind CSS", "3.4.14", "Utility-first CSS framework"],
                ["Lucide React", "0.460.0", "Icon library"],
                ["Recharts", "3.7.0", "Chart/visualization library"],
                ["react-hot-toast", "2.4.1", "Toast notification system"],
                ["date-fns", "4.1.0", "Date formatting utilities"],
                ["Vite", "7.2.4", "Build tool and dev server"],
            ]
,
            caption="Frontend Project Dependencies"
        )

        self.add_paragraph("Application Entry Point", bold=True)
        self.add_paragraph(
            "The application bootstraps from main.jsx which mounts the root App component to the "
            "DOM. App.jsx wraps the application in Suspense (for lazy loading), AuthProvider (for "
            "authentication state), RouterProvider (for routing), and Toaster (for notifications). "
            "The file structure follows a domain-driven organization as shown in Fig. 3."
        )
        self.add_figure(3)

        self.add_code_block(
            "// Frontend/src/App.jsx\n"
            "function App() {\n"
            "  return (\n"
            "    <Suspense fallback={<LoadingFallback />}>\n"
            "      <AuthProvider>\n"
            "        <RouterProvider router={router} />\n"
            "        <Toaster position='top-right' toastOptions={{duration: 4000}} />\n"
            "      </AuthProvider>\n"
            "    </Suspense>\n"
            "  );\n"
            "}"
        )

        self.add_paragraph("Internationalization (i18n) Configuration", bold=True)
        self.add_paragraph(
            "The i18n system uses i18next with browser language detection, supporting English (en) "
            "and Arabic (ar) with 755 translation keys each. The changeLanguage() function updates "
            "the i18n instance, DOM direction attribute, and localStorage in a single operation. "
            "The language switcher is accessible from both the public navigation bar and the "
            "authenticated dashboard header, as shown in Fig. 4."
        )
        self.add_figure(4)

        self.add_code_block(
            "// Frontend/src/i18n/index.js\n"
            "i18n.use(LanguageDetector).use(initReactI18next).init({\n"
            "  resources: { en: { translation: en }, ar: { translation: ar } },\n"
            "  fallbackLng: 'en',\n"
            "  detection: { order: ['localStorage', 'navigator', 'htmlTag'],\n"
            "               lookupLocalStorage: 'phishaware-language' },\n"
            "});\n"
            "export const changeLanguage = (lang) => {\n"
            "  i18n.changeLanguage(lang);\n"
            "  document.documentElement.dir = lang === 'ar' ? 'rtl' : 'ltr';\n"
            "  document.documentElement.lang = lang;\n"
            "  localStorage.setItem('phishaware-language', lang);\n"
            "};"
        )

        self.add_paragraph("RTL Support Implementation", bold=True)
        self.add_paragraph(
            "The platform implements complete RTL support for Arabic content. When the language is "
            "switched to Arabic, the document direction is set to 'rtl', causing Tailwind CSS logical "
            "properties (ps-, pe-, ms-, me- for start/end padding and margins) to automatically mirror "
            "the layout. Fig. 5 and Fig. 6 demonstrate the LTR and RTL layout transformation."
        )
        self.add_figure(5)
        self.add_figure(6)

        # ── 2) Authentication & Protected Routes ──
        self.doc.add_heading("2) Authentication & Protected Routes", level=3)
        self.add_paragraph(
            "The authentication system provides a complete login/registration flow with JWT token "
            "management, automatic token refresh, email verification, password recovery, and role-based "
            "route protection. The Axios service layer handles all API communication with interceptors "
            "for automatic token attachment and 401 error handling."
        )

        self.add_paragraph("Axios Service Layer", bold=True)
        self.add_paragraph(
            "The API client (api/axios.js) creates an Axios instance with the base URL configured "
            "from environment variables (VITE_API_URL, defaulting to http://localhost:8000/api/v1). "
            "Request interceptors automatically attach the JWT access token from localStorage. "
            "Response interceptors handle 401 errors by attempting token refresh, queuing failed "
            "requests during the refresh process, and replaying them with the new token. Failed "
            "refresh attempts clear all tokens and redirect to the login page."
        )

        self.add_paragraph("API Endpoint Organization", bold=True)
        self.add_paragraph(
            "The endpoints.js file organizes all API calls into 10 module objects: authAPI, "
            "companiesAPI, employeesAPI, campaignsAPI, simulationsAPI, quizzesAPI, trainingAPI, "
            "analyticsAPI, gamificationAPI, and communityAPI. Each module exports methods that "
            "correspond to backend API endpoints, providing a clean separation between UI logic "
            "and HTTP communication."
        )

        self.add_paragraph("Authentication Context (AuthContext)", bold=True)
        self.add_paragraph(
            "The AuthContext provides centralized authentication state management through React "
            "Context. It exposes the current user object, authentication status, and methods for "
            "login, register, logout, profile update, password change, and role checking. On "
            "initialization, it checks localStorage for existing tokens and verifies them against "
            "the API's profile endpoint:"
        )

        self.add_code_block(
            "// Frontend/src/contexts/AuthContext.jsx\n"
            "const USER_ROLES = {\n"
            "  SUPER_ADMIN: 'SUPER_ADMIN',\n"
            "  COMPANY_ADMIN: 'COMPANY_ADMIN',\n"
            "  EMPLOYEE: 'EMPLOYEE',\n"
            "  PUBLIC_USER: 'PUBLIC_USER',\n"
            "};\n"
            "const getDashboardPath = () => {\n"
            "  switch (user?.role) {\n"
            "    case 'SUPER_ADMIN': return '/admin/dashboard';\n"
            "    case 'COMPANY_ADMIN': return '/company/dashboard';\n"
            "    case 'EMPLOYEE': return '/employee/dashboard';\n"
            "    default: return '/';\n"
            "  }\n"
            "};"
        )

        self.add_paragraph("Login Page", bold=True)
        self.add_paragraph(
            "The LoginPage (Fig. 7) presents a clean form with email and password fields, a "
            "\"Remember me\" checkbox, and a \"Forgot password?\" link. Form validation provides "
            "inline error messages with red borders (Fig. 9). The page includes a company registration "
            "call-to-action and an employee invitation hint. For unverified users, an amber banner "
            "appears with a \"Resend verification email\" button. The login page supports full RTL "
            "transformation for Arabic (Fig. 8)."
        )
        self.add_figure(7)
        self.add_figure(8)
        self.add_figure(9)

        self.add_paragraph("Registration and Verification Flow", bold=True)
        self.add_paragraph(
            "The RegisterPage (Fig. 10) collects company name, website URL, and administrator "
            "account details. Password validation enforces minimum 8-character length (Fig. 11). "
            "After successful registration, the page transitions to an inline success state showing "
            "a \"Check your inbox!\" message with the registered email and a resend verification "
            "button. The VerifyEmailPage handles token verification and displays a success page "
            "with a Login button (Fig. 12)."
        )
        self.add_figure(10)
        self.add_figure(11)
        self.add_figure(12)

        self.add_paragraph("Protected Route System", bold=True)
        self.add_paragraph(
            "The routing system implements three route guard components in ProtectedRoute.jsx: "
            "ProtectedRoute (requires authentication + optional role check, redirects to /login), "
            "PublicRoute (blocks authenticated users, redirects to dashboard), and GuestRoute "
            "(accessible to all). Routes are organized by user role with DashboardLayout as the "
            "parent for all authenticated routes:"
        )

        self.add_table(
            ["Route Group", "Guard", "Layout", "Example Routes"],
            [
                ["Public", "GuestRoute", "PublicLayout", "/, /training, /quiz, /community"],
                ["Auth", "PublicRoute", "(none)", "/login, /register, /verify-email/:token"],
                ["Employee", "ProtectedRoute (EMPLOYEE)", "DashboardLayout", "/employee/dashboard, /employee/quizzes"],
                ["Company Admin", "ProtectedRoute (COMPANY_ADMIN)", "DashboardLayout", "/company/dashboard, /company/campaigns"],
                ["Super Admin", "ProtectedRoute (SUPER_ADMIN)", "DashboardLayout", "/admin/dashboard, /admin/companies"],
                ["Simulation", "(none)", "(none)", "/simulation/caught/:token"],
            ]
,
            caption="Route Guard Configuration"
        )

    def _build_frontend_B(self):
        self.doc.add_heading("B. Administrator Management Interfaces (Week 10)", level=2)
        self.add_paragraph(
            "Week 10 focused on building the company administrator interfaces, including the main "
            "dashboard layout with sidebar navigation and header, employee management with invitation "
            "workflows, campaign creation and management, simulation template libraries, and analytics "
            "dashboards. These interfaces provide comprehensive administrative control over the "
            "organization's cybersecurity awareness program."
        )

        # ── 1) Admin Layout, Employee & Campaign Management ──
        self.doc.add_heading("1) Admin Layout, Employee & Campaign Management", level=3)
        self.add_paragraph(
            "The administrator interface uses a DashboardLayout component that provides a consistent "
            "structure across all authenticated pages. The layout features a collapsible sidebar for "
            "navigation, a fixed header with notifications and user menu, and a main content area. "
            "The layout adapts to mobile viewports with a hamburger menu and overlay-based sidebar."
        )

        self.add_paragraph("Dashboard Layout Structure", bold=True)
        self.add_paragraph(
            "The DashboardLayout (Fig. 13) implements a two-column design with a 256px fixed sidebar "
            "on desktop and an overlay sidebar on mobile. The sidebar displays the PhishAware logo, "
            "role-specific navigation items with icons, and a logout button. The header (Fig. 15) "
            "includes a language switcher (Globe icon with AR/EN text), notification dropdown with "
            "unread count badge, and a user profile menu with dropdown options."
        )
        self.add_figure(13)
        self.add_figure(14)
        self.add_figure(15)

        self.add_paragraph("Campaign Management Interface", bold=True)
        self.add_paragraph(
            "The CampaignList page (Fig. 16) displays all campaigns in a table with status badges "
            "(Draft, Active, Paused, Completed), start dates, participant counts, and action dropdowns. "
            "The CampaignCreate page (Fig. 17) provides a form for creating new quiz campaigns with "
            "configurable parameters including the English/Arabic language ratio slider and phishing "
            "scenario difficulty ratio controls. The employee assignment interface (Fig. 18) allows "
            "selecting specific employees with department filtering and displays the selected count."
        )
        self.add_figure(16)
        self.add_figure(17)
        self.add_figure(18)

        self.add_paragraph("Employee Management Interface", bold=True)
        self.add_paragraph(
            "The CompanyEmployees page (Fig. 26) presents a comprehensive employee management table "
            "with columns for name, email, department, risk score badge (color-coded by level), "
            "campaign count, and action buttons. The invite modal (Fig. 27) accepts first name, "
            "last name, email address, and department selection for single employee invitations. "
            "The employee detail view (Fig. 28) shows the employee's profile information, current "
            "risk score visualization, and tabbed sections for training history and campaign results."
        )
        self.add_figure(26)
        self.add_figure(27)
        self.add_figure(28)

        self.add_paragraph("Training Management Interface", bold=True)
        self.add_paragraph(
            "The TrainingManagement page (Fig. 29) displays available training modules with card-based "
            "layouts showing module title, description, category badge, and action buttons. The View "
            "Details modal (Fig. 30) provides comprehensive module information including description, "
            "learning objectives, and quiz details. The assignment interface (Fig. 31) allows selecting "
            "target employees with department and risk level filters for bulk training assignment."
        )
        self.add_figure(29)
        self.add_figure(30)
        self.add_figure(31)

        self.add_paragraph("Admin Profile", bold=True)
        self.add_paragraph(
            "The CompanyProfile page (Fig. 32) displays the administrator's personal information "
            "fields and account settings options, allowing profile updates and password changes."
        )
        self.add_figure(32)

        # ── 2) Content Libraries & Analytics Dashboards ──
        self.doc.add_heading("2) Content Libraries & Analytics Dashboards", level=3)
        self.add_paragraph(
            "The content libraries provide template management for phishing simulations, and the "
            "analytics dashboards deliver data-driven insights through interactive charts and KPI "
            "cards. The simulation creation wizard guides administrators through a four-step process "
            "from template selection to email delivery."
        )

        self.add_paragraph("Simulation Template Library", bold=True)
        self.add_paragraph(
            "The CompanySimulations page (Fig. 19) displays available pre-built phishing email "
            "templates in a list format with template name, attack vector, difficulty level, and "
            "language. Templates marked as public (company=None) are globally available, while "
            "company-specific templates are scoped to the organization."
        )
        self.add_figure(19)

        self.add_paragraph("Four-Step Simulation Creation Wizard", bold=True)
        self.add_paragraph(
            "The SimulationCreate component implements a guided four-step wizard for creating and "
            "sending phishing simulations. Each step validates input before allowing progression "
            "to the next:"
        )

        self.add_paragraph(
            "Step 1 - Template Selection (Fig. 20): Administrators enter a simulation name and "
            "description, then select from pre-built phishing email templates. Each template "
            "preview shows the subject line, sender, and difficulty level."
        )
        self.add_figure(20)

        self.add_paragraph(
            "Step 2 - Employee Selection (Fig. 21): The employee selection interface provides "
            "department-based filtering, individual checkboxes, and a \"select all\" option. "
            "The selected employee count is displayed prominently."
        )
        self.add_figure(21)

        self.add_paragraph(
            "Step 3 - Review and Confirm (Fig. 22): A summary view displays the simulation name, "
            "selected template, target employee count, and tracking options before final confirmation."
        )
        self.add_figure(22)

        self.add_paragraph(
            "Step 4 - Sending (Fig. 23): A loading screen with a progress indicator displays while "
            "simulation emails are being sent via the API's send endpoint. Upon completion, the "
            "user is redirected to the simulation analytics page."
        )
        self.add_figure(23)

        self.add_paragraph("Analytics Dashboard", bold=True)
        self.add_paragraph(
            "The CompanyAnalytics page provides a comprehensive data visualization dashboard with "
            "two main sections. The upper section (Fig. 24) displays four KPI cards (total employees, "
            "active campaigns, average risk score, training completion rate), an Average Risk Score "
            "Over Time line chart (built with Recharts), and a Risk Level Distribution pie chart. "
            "The lower section (Fig. 25) shows Security Training Activity counters and Training "
            "Completion Rates breakdown by training module type."
        )
        self.add_figure(24)
        self.add_figure(25)

    def _build_frontend_C(self):
        self.doc.add_heading("C. Employee Training & Gamification UI (Week 11)", level=2)
        self.add_paragraph(
            "Week 11 developed the employee-facing interfaces including the personal dashboard, "
            "training module interactions, quiz-taking interfaces, and gamification displays. "
            "These interfaces provide employees with an engaging, gamified experience for improving "
            "their cybersecurity awareness while tracking their progress and achievements."
        )

        # ── 1) Employee Dashboard & Training Modules ──
        self.doc.add_heading("1) Employee Dashboard & Training Modules", level=3)
        self.add_paragraph(
            "The employee dashboard serves as the central hub for each employee's security awareness "
            "journey. It presents personalized risk assessment, pending tasks, earned achievements, "
            "and quick action links. The dashboard adapts to mobile viewports with a stacked layout "
            "design, ensuring accessibility on all devices."
        )

        self.add_paragraph("Employee Dashboard", bold=True)
        self.add_paragraph(
            "The EmployeeDashboard (Fig. 33) features a prominent risk score gauge implemented as "
            "an SVG circular progress indicator with color coding: green for LOW (0-25), yellow for "
            "MEDIUM (26-50), orange for HIGH (51-75), and red for CRITICAL (76-100). The gauge shows "
            "the numeric score, risk level badge, and breakdown of quiz accuracy and simulation click "
            "rate. For new employees without established scores, a gray gauge with an encouraging "
            "message prompts them to complete their first quiz."
        )
        self.add_figure(33)

        self.add_paragraph(
            "The dashboard also includes a 2x2 quick stats grid showing pending quizzes, pending "
            "training, badges earned, and leaderboard rank, each linking to the respective page. "
            "A recent badges section displays the three most recently earned badges with emoji, name, "
            "and earned date. Quick action cards at the bottom provide direct links to Take Quiz, "
            "Continue Training, and View Leaderboard. The mobile responsive view (Fig. 34) "
            "demonstrates the stacked vertical layout with hamburger menu navigation at 375px width."
        )
        self.add_figure(34)

        self.add_paragraph("Training Module Interface", bold=True)
        self.add_paragraph(
            "The EmployeeTraining page lists all assigned training modules with status badges "
            "(Assigned, In Progress, Completed, Passed, Failed), due dates, and progress bars. "
            "Clicking a training module navigates to TakeTraining, which presents the module detail "
            "page (Fig. 38) with title, description, learning objectives, and a training mode "
            "selection interface."
        )
        self.add_figure(38)

        self.add_paragraph(
            "The training mode selection (Fig. 39) offers two cards: Video Training (traditional "
            "video content followed by a quiz) and Interactive Learning (scenario-based interactive "
            "modules). The interactive mode maps training categories to specific lesson modules: "
            "EMAIL_SECURITY maps to phishing lessons (V4/V5/V6), SOCIAL_ENGINEERING to vishing, "
            "and MOBILE_SECURITY to smishing. Each interactive lesson has both English and Arabic "
            "versions, providing 12 total interactive modules."
        )
        self.add_figure(39)

        # ── 2) Quiz Interfaces & Gamification Display ──
        self.doc.add_heading("2) Quiz Interfaces & Gamification Display", level=3)
        self.add_paragraph(
            "The quiz and gamification interfaces provide engaging interaction patterns for "
            "employee assessment and recognition. The quiz system supports both campaign quizzes "
            "(where employees classify emails as phishing or legitimate) and training quizzes "
            "(module-specific knowledge assessments). The gamification display showcases earned "
            "badges, points, and leaderboard rankings."
        )

        self.add_paragraph("Quiz Taking Interface", bold=True)
        self.add_paragraph(
            "The TakeQuiz component implements a multi-stage quiz experience. The start screen "
            "(Fig. 35) displays the quiz title, description, difficulty level indicator, and a "
            "Start Quiz button. Once started, the quiz question interface (Fig. 36) presents "
            "phishing email scenarios with Safe and Phishing answer buttons and a progress indicator "
            "showing dots for each question. Navigation buttons allow moving between questions."
        )
        self.add_figure(35)
        self.add_figure(36)

        self.add_paragraph(
            "Upon completion, the results page (Fig. 37) displays the final score, performance "
            "breakdown by category, and personalized recommendations based on the employee's "
            "strengths and weaknesses. The results interface provides actionable feedback to guide "
            "further learning."
        )
        self.add_figure(37)

        self.add_paragraph("Badge Showcase", bold=True)
        self.add_paragraph(
            "The EmployeeBadges page (Fig. 40) presents all available badges in a grid layout. "
            "Earned badges appear in full color with their emoji, name, and rarity badge (COMMON "
            "in gray, UNCOMMON in green, RARE in blue, EPIC in purple, LEGENDARY in yellow). "
            "Locked badges are displayed in greyscale with a lock icon overlay. Clicking a badge "
            "opens a modal with full details including description, criteria, and earning date. "
            "The page supports filtering between All, Earned, and Locked badges."
        )
        self.add_figure(40)

        self.add_paragraph("Leaderboard Interface", bold=True)
        self.add_paragraph(
            "The EmployeeLeaderboard page (Fig. 41) features a \"Your Position\" section with the "
            "employee's current rank, total points, and percentile. Three period tabs (Weekly, "
            "Monthly, All Time) filter the leaderboard data. The top 3 positions receive special "
            "styling with Trophy (gold), Medal (silver), and Award (bronze) icons. Each leaderboard "
            "entry shows an avatar (generated from initials in a colored circle), rank number, name, "
            "and points badge. A \"How Points Work\" section explains the scoring formula, providing "
            "transparency to employees about how their rank is calculated."
        )
        self.add_figure(41)

        self.add_paragraph("Employee Profile", bold=True)
        self.add_paragraph(
            "The EmployeeProfile page (Fig. 42) displays personal information fields (first name, "
            "last name, email, phone, department) and account configuration options including password "
            "change. The profile page uses the same form component pattern as the admin profile."
        )
        self.add_figure(42)

    def _build_frontend_D(self):
        self.doc.add_heading("D. Public Awareness Portal & Final Polish (Week 11)", level=2)
        self.add_paragraph(
            "The final frontend development phase created the public-facing community portal for "
            "cybersecurity awareness, including the landing page, awareness content pages, public "
            "quizzes, and the platform's design system. These interfaces are accessible without "
            "authentication, serving as both an awareness resource and a funnel for platform adoption."
        )

        # ── 1) Community Portal & Awareness Content ──
        self.doc.add_heading("1) Community Portal & Awareness Content", level=3)
        self.add_paragraph(
            "The public portal provides free cybersecurity awareness content targeting Saudi Arabian "
            "users and organizations. The portal features bilingual content, Saudi-specific examples "
            "(SAMA, NCA, CERT references), and interactive elements designed to educate visitors about "
            "common cyber threats including email phishing, SMS scams, and voice call fraud."
        )

        self.add_paragraph("Landing Page Design", bold=True)
        self.add_paragraph(
            "The LandingPage (Fig. 43) opens with a hero section featuring a gradient background "
            "(primary-600 to primary-800) with the headline \"Protect Yourself from Cyber Threats\", "
            "a descriptive tagline, and dual call-to-action buttons: \"Start Learning\" (solid white) "
            "and \"Take Free Quiz\" (bordered). Below the hero, a features section displays four "
            "platform capabilities in a grid: Training, Simulations, Analytics, and Compliance."
        )
        self.add_figure(43)

        self.add_paragraph(
            "The \"Learn About Cyber Threats\" section (Fig. 44) presents three topic cards\u2014Email "
            "(blue), SMS (green), and Voice (purple)\u2014each with Saudi-specific examples and "
            "descriptions. These cards link to detailed topic training pages. A statistics banner "
            "(Fig. 45) displays platform adoption metrics: 500+ Companies Protected, 50K+ Employees "
            "Trained, 95% Threat Detection rate, and 60% Risk Reduction."
        )
        self.add_figure(44)
        self.add_figure(45)

        self.add_paragraph("Public Training Content", bold=True)
        self.add_paragraph(
            "The TrainingTopics page (Fig. 46) displays three training module cards with descriptions "
            "and \"Learn More\" action buttons. The page supports complete RTL transformation for "
            "Arabic (Fig. 47), demonstrating the platform's bilingual commitment with identical "
            "content structure and mirrored layout. Each topic links to a detailed TopicTraining page "
            "with expanded educational content."
        )
        self.add_figure(46)
        self.add_figure(47)

        # ── 2) Public Quizzes & Final UI/UX Polish ──
        self.doc.add_heading("2) Public Quizzes & Final UI/UX Polish", level=3)
        self.add_paragraph(
            "The public quiz system provides free, anonymous cybersecurity knowledge assessments "
            "that serve dual purposes: educating visitors about common threats and demonstrating "
            "the platform's value proposition to potential customers. The design system establishes "
            "visual consistency across all platform interfaces through a defined color palette, "
            "typography standards, and responsive design patterns."
        )

        self.add_paragraph("Public Quiz Interface", bold=True)
        self.add_paragraph(
            "The PublicQuiz page (Fig. 48) presents a \"Test Your Knowledge\" section with three "
            "quiz options: Email Phishing (10 questions), SMS Scams (10 questions), and Voice Scams "
            "(10 questions), all marked as Beginner difficulty. The quiz question interface (Fig. 49) "
            "displays multiple choice options with a progress bar indicator and \"Submit Answer\" "
            "button. Quiz attempts are tracked anonymously using session IDs, and the results page "
            "(Fig. 50) shows the achieved score, category performance breakdown, and a registration "
            "call-to-action encouraging visitors to sign up for full platform access."
        )
        self.add_figure(48)
        self.add_figure(49)
        self.add_figure(50)

        self.add_paragraph("Design System", bold=True)
        self.add_paragraph(
            "The PhishAware design system (Fig. 1) establishes a consistent visual language across "
            "the platform through a defined color palette with 10 color swatches including primary "
            "blues, success greens, warning ambers, and danger reds, each with hex codes and usage "
            "labels. The PhishAware logo (Fig. 2) features a shield icon symbolizing security "
            "protection, paired with the wordmark in the navigation bar."
        )
        self.add_figure(1)
        self.add_figure(2)

        self.add_paragraph("Responsive Design", bold=True)
        self.add_paragraph(
            "The application implements responsive design using Tailwind CSS breakpoints (sm: 640px, "
            "md: 768px, lg: 1024px, xl: 1280px). Mobile layouts use stacked vertical arrangements "
            "with hamburger menu navigation, while desktop layouts utilize multi-column grids. The "
            "DashboardLayout switches from a fixed sidebar (desktop) to an overlay sidebar (mobile). "
            "All tables convert to card-based layouts on small screens, ensuring data readability. "
            "The public landing page adjusts hero text sizing, feature grids collapse from 4 columns "
            "to 1, and statistics reflow from horizontal to vertical alignment."
        )

        self.add_paragraph("Accessibility and UX Patterns", bold=True)
        self.add_paragraph(
            "The platform implements several UX best practices: toast notifications for immediate "
            "feedback on actions (react-hot-toast with top-right positioning and 4-second duration), "
            "loading spinners during async operations, empty states with descriptive text and action "
            "suggestions, confirmation dialogs for destructive operations (delete, cancel), form "
            "validation with inline error messages, and optimistic UI updates where appropriate. "
            "Charts use the Recharts library with responsive containers that adapt to available width."
        )

    def build_references(self):
        self.add_page_break()
        self.doc.add_heading("REFERENCES", level=1)
        refs = [
            '[1] Django Software Foundation, "Django documentation," 2024. [Online]. Available: https://docs.djangoproject.com/',
            '[2] T. Christie, "Django REST framework," 2024. [Online]. Available: https://www.django-rest-framework.org/',
            '[3] Meta Platforms, "React documentation," 2024. [Online]. Available: https://react.dev/',
            '[4] E. You, "Vite: Next generation frontend tooling," 2024. [Online]. Available: https://vitejs.dev/',
            '[5] J. Müller, "i18next internationalization framework," 2024. [Online]. Available: https://www.i18next.com/',
            '[6] Tailwind Labs, "Tailwind CSS documentation," 2024. [Online]. Available: https://tailwindcss.com/',
            '[7] Microsoft, "Office Open XML (OOXML) documentation," 2024. [Online]. Available: https://learn.microsoft.com/en-us/openspecs/',
            '[8] D. Arias, "djangorestframework-simplejwt," 2024. [Online]. Available: https://django-rest-framework-simplejwt.readthedocs.io/',
            '[9] SendGrid, "SendGrid SMTP relay documentation," 2024. [Online]. Available: https://docs.sendgrid.com/',
            '[10] Anti-Phishing Working Group, "Phishing activity trends report," APWG, 2024.',
            '[11] Saudi National Cybersecurity Authority, "National Cybersecurity Strategy," NCA, 2023.',
            '[12] SANS Institute, "Security awareness training best practices," SANS, 2024.',
            '[13] Recharts Contributors, "Recharts composable chart library," 2024. [Online]. Available: https://recharts.org/',
            '[14] Lucide Contributors, "Lucide icon library," 2024. [Online]. Available: https://lucide.dev/',
            '[15] OWASP Foundation, "OWASP Top Ten," 2021. [Online]. Available: https://owasp.org/www-project-top-ten/',
        ]
        for ref in refs:
            p = self.add_paragraph(ref, size=Pt(9))
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.space_after = Pt(4)

    def build_appendices(self):
        self.add_page_break()
        self.doc.add_heading("APPENDICES", level=1)

        self.doc.add_heading("A. Complete API Endpoint Reference", level=2)
        self.add_paragraph(
            "The following table lists all API endpoints available in the PhishAware platform "
            "organized by application module:"
        )

        endpoints = [
            ["POST /api/v1/auth/register/", "None", "User registration"],
            ["POST /api/v1/auth/login/", "None", "JWT login"],
            ["POST /api/v1/auth/logout/", "JWT", "Token blacklisting"],
            ["POST /api/v1/auth/token/refresh/", "None", "Refresh access token"],
            ["POST /api/v1/auth/verify-email/<token>/", "None", "Email verification"],
            ["POST /api/v1/auth/resend-verification/", "None", "Resend verification"],
            ["POST /api/v1/auth/password-reset/", "None", "Request password reset"],
            ["POST /api/v1/auth/password-reset/<token>/", "None", "Reset password"],
            ["GET/PATCH /api/v1/auth/profile/", "JWT", "User profile"],
            ["POST /api/v1/auth/change-password/", "JWT", "Change password"],
            ["POST /api/v1/employees/invite/", "Admin", "Invite employee"],
            ["GET /api/v1/employees/invite/<token>/", "None", "Invitation details"],
            ["POST /api/v1/employees/invite/<token>/accept/", "None", "Accept invitation"],
            ["GET /api/v1/companies/", "JWT", "List companies"],
            ["POST /api/v1/companies/", "SuperAdmin", "Create company"],
            ["GET /api/v1/companies/{id}/stats/", "JWT", "Company statistics"],
            ["GET /api/v1/campaigns/campaigns/", "JWT", "List campaigns"],
            ["POST /api/v1/campaigns/campaigns/", "Admin", "Create campaign"],
            ["POST /api/v1/campaigns/quizzes/{id}/start/", "JWT", "Start quiz"],
            ["POST /api/v1/campaigns/quizzes/{id}/submit/", "JWT", "Submit quiz"],
            ["GET /api/v1/simulations/templates/", "JWT", "List templates"],
            ["POST /api/v1/simulations/campaigns/{id}/send/", "Admin", "Send simulation"],
            ["GET /api/v1/simulations/link/{token}/", "None", "Track link click"],
            ["GET /api/v1/simulations/feedback/{token}/", "None", "Simulation feedback"],
            ["GET /api/v1/training/modules/", "JWT", "List training modules"],
            ["POST /api/v1/training/assignments/bulk_assign/", "Admin", "Bulk assign training"],
            ["GET /api/v1/training/risk-scores/my_score/", "JWT", "My risk score"],
            ["GET /api/v1/gamification/badges/my_badges/", "JWT", "My badges"],
            ["GET /api/v1/gamification/leaderboard/", "JWT", "Leaderboard"],
            ["GET /api/v1/analytics/dashboard/overview/", "Admin", "Dashboard stats"],
            ["POST /api/v1/analytics/export/csv/", "Admin", "Export CSV"],
            ["GET /api/v1/community/articles/", "None", "Public articles"],
            ["GET /api/v1/community/quizzes/", "None", "Public quizzes"],
            ["GET /api/v1/notifications/", "JWT", "User notifications"],
        ]
        self.add_table(
            ["Endpoint", "Auth", "Description"],
            endpoints
,
            caption="Complete API Endpoint Reference"
        )

        self.doc.add_heading("B. Database Model Summary", level=2)
        self.add_paragraph(
            "Total models across all applications:"
        )
        self.add_table(
            ["Application", "Model Count", "Models"],
            [
                ["accounts", "1", "User"],
                ["companies", "1", "Company"],
                ["campaigns", "3", "Campaign, Quiz, QuizResult"],
                ["assessments", "2", "EmailTemplate, QuizQuestion"],
                ["simulations", "4", "SimulationTemplate, SimulationCampaign, EmailSimulation, TrackingEvent"],
                ["training", "6", "RiskScore, RiskScoreHistory, TrainingModule, TrainingQuestion, RemediationTraining, InteractiveLessonProgress"],
                ["gamification", "4", "Badge, EmployeeBadge, PointsTransaction, EmployeePoints"],
                ["community", "6", "ArticleCategory, Article, PublicQuiz, PublicQuizQuestion, PublicQuizAttempt, Resource"],
                ["notifications", "1", "Notification"],
                ["Total", "28", ""],
            ]
,
            caption="Database Model Summary by Application"
        )

        self.doc.add_heading("C. Frontend Component Summary", level=2)
        self.add_table(
            ["Section", "Page Count", "Key Components"],
            [
                ["Authentication", "7", "Login, Register, Verify Email, Accept Invitation, Forgot/Reset Password"],
                ["Employee Portal", "8", "Dashboard, Quizzes, Training, Badges, Leaderboard, Profile, TakeQuiz, TakeTraining"],
                ["Company Admin", "10", "Dashboard, Campaigns, Simulations, Employees, Analytics, Training, Profile"],
                ["Super Admin", "5", "Dashboard, Companies, Analytics, Users, Profile"],
                ["Public Portal", "10", "Landing, Training Topics, Quiz, Community, SimulationCaught, Error Pages"],
                ["Total", "43 pages", "Plus layouts, common components, and 12 interactive modules"],
            ]
,
            caption="Frontend Component Summary"
        )

    def build_table_of_figures(self):
        """Insert a Table of Figures (Word auto-populates from SEQ Figure captions)."""
        # TOC \h \z \c "Fig." tells Word to build a list from captions starting with "Fig."
        add_toc_field(self.doc, 'TOC \\h \\z \\c "Figure"', heading_text="LIST OF FIGURES")
        self.add_page_break()

    def build_table_of_tables(self):
        """Insert a Table of Tables (Word auto-populates from SEQ Table captions)."""
        add_toc_field(self.doc, 'TOC \\h \\z \\c "Table"', heading_text="LIST OF TABLES")
        self.add_page_break()

    def generate(self):
        print("Building title page...")
        self.build_title_page()
        print("Building abstract...")
        self.build_abstract()
        print("Building table of contents...")
        self.build_toc()
        print("Building list of figures...")
        self.build_table_of_figures()
        print("Building list of tables...")
        self.build_table_of_tables()
        print("Building Phase 1: Backend Foundation...")
        self.build_phase1()
        print("Building Phase 2: Frontend Development...")
        self.build_phase2()
        print("Building references...")
        self.build_references()
        print("Building appendices...")
        self.build_appendices()

        # Set document metadata
        props = self.doc.core_properties
        props.author = "PhishAware Team"
        props.title = "PhishAware: A Bilingual Cybersecurity Awareness Platform"
        props.subject = "Senior Project Report"
        props.keywords = "cybersecurity, phishing, awareness, Django, React"
        props.category = "Technical Report"
        props.comments = ""

        print(f"Saving to {OUTPUT_PATH}...")
        self.doc.save(OUTPUT_PATH)
        print(f"Done! Document saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    gen = IEEEDocGenerator()
    gen.generate()
